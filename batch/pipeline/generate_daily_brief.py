"""
generate_daily_brief.py
=======================
MLB Betting Model — Daily Brief Generator
Reads from mlb_stats.db and outputs the formatted betting brief.

CHANGE LOG (latest first)
──────────────────────────
2026-04-22  Session windows: removed clock-derived ``closing`` (18:45–20:15 → ``primary``); run
            closing only with ``--session closing``. Word ``build_docx_from_text`` bolds Signal stack
            / Core / Support / Minor lines under the pick card.
2026-04-22  Primary / additional brief pick rows: ``format_bet_block`` (ASCII card, bucket BET
            label, grouped signals, model score, stake); ALT/INFO below card. Word SIGNAL column
            uses ``signal_summary_for_doc`` (same Core/Support/Minor grouping when ``_scored_game``
            present).
2026-04-22  Brief helpers: ``score_to_confidence``, ``tier_label``, ``SIGNAL_LABELS`` /
            ``humanize_signal``, ``group_signals``, ``format_bet_block``; ``evaluate_signals``
            attaches ``_scored_game`` (in-process). ``format_bet_block`` uses ``best_aggregate_score``
            on ``ScoredGame`` (not ``data_flags`` ``BEST …`` parsing).
2026-04-22  Customer-facing action briefs: removed all “bets to avoid” output (text + Word).
            Avoid-only games roll into the no-signal slate; ``save_signal_state`` /
            ``brief_picks`` no longer persist AVOID rows from these runs. Confidence scores
            on picks are unchanged.
2026-04-22  Default ``now``: non-``prior`` briefs no longer require ``--as-of`` / ``--as-of-time``;
            wall-clock America/New_York is used. ``as_of`` / ``as_of_time`` are for replays and
            the pipeline runner. ``as_of_for_slate`` stays ``None`` for those live runs so
            ``load_games`` still excludes completed games in SQL.
2026-04-22  ``already_ran``: treat ``brief_log`` as duplicate only when ``output_file`` exists on
            disk, so a logged-but-missing file no longer blocks without ``--force``.
2026-04-22  ``--game-group-id`` + ``brief_log.game_group_id``: pipeline ``group_brief`` can run once
            per group without tripping the duplicate guard; bet_ledger T−30 materialization
            is no longer skipped for later groups. Filenames get ``_gN`` when set.
2026-04-21  Morning session (pipeline ``early_peek``): single "Today's Slate" section
            only — matchup, venue, start time, factual weather, starters when known,
            ML / O/U / runline. No signal evaluation, no picks, streak monitor, or avoids;
            no ``signal_state`` writes from morning. ``weather_line(..., wind_signal_hints=False)``
            for this run; ``odds_summary_line`` shows both sides' runline when present.
2026-04-20  Retractable roof removed as a model/brief factor (no AvoidFinding, no
            confidence penalty, no weather-line or prior-report retractable copy).
2026-04-19  Default brief filenames: ``brief-{slate_date}_{ET_run_stamp}_ET[ _prior].txt`` (Windows-safe);
            ``brief_log.output_file`` stores that path; legacy ``{date}_{session}.txt`` names may be archived.
2026-04-19  ``--as-of-time`` accepts full ``YYYY-MM-DD HH:MM`` (ET) or time-only ``HH:MM`` with ``--date``.
2026-04-19  Hybrid session CLI: ``--as-of-time HH:MM`` (ET) or full ``--as-of``;
            non-``prior`` sessions resolve from ``SESSION_WINDOWS`` (``SESSION_PULL_WINDOW``
            unchanged). ``--session`` ignored when a clock is supplied unless ``prior``.
2026-04-19  Fix: closing brief now skips games already past ``game_start_utc``.
            Active signal display and CLV confirmation only shown for unplayed games.
2026-04-19  Fix: removed H3b cross-reference from MV-B reason string.
            H3b is monitor-only and should not appear in other signals'
            reason text. MV-B reason now self-contained (``score_game._eval_mv_b``).
2026-04-17  brief_picks: skip inserts after game start (vs ``now``); cross-session
            dedupe on (game_date, game_pk, signal, market); ``model_version``
            column (default ``legacy``); ``ensure_brief_picks`` before save.
2026-04-17  CLI: --debug-wind prints per-game wind debug (DB row vs dressed
            GameEnvironment) during brief / prior; pair with --verbose for tiers.
2026-04-17  Venue wind suppression no longer emits an AvoidFinding (tier/env gates +
            data_flags only). Legacy avoid=True only for hard avoids or no-signal
            avoid rows; Word brief avoids duplicate AVOID banner when picks exist.
2026-04-17  Prior report: removed duplicate FURTHER SIGNALS block (ranks #7+ only in
            RETROACTIVE); ENV/venue avoids spell out “class” of bet skipped (not one ticket).
2026-04-17  Prior report: NEXT picks use the same ledger box as TOP; AVOID section
            spells out bet-to-skip + counterfactual verdict; bet_ledger unique key
            includes signal_at_time so AVOID rows (stake 0) materialize from
            signal_state; backfill_bet_ledger_from_signal_state runs on prior;
            grade_bet_ledger scores avoids as good_avoid / bad_avoid / push_avoid.
2026-04-17  Scoring function implemented. score_game() replaces evaluate_signals().
            All signal logic consolidated into per-signal evaluator functions in
            batch/pipeline/score_game.py. Three-dimension tier decision: signal_strength
            × env_ceiling × data_completeness. LHP_FADE signal added (replaces NF4 as
            primary formulation; NF4 retained as backward-compat alias). Hostile
            environment detection added. Avoid evaluation separated from signal evaluation.
            ScoredGame output dataclass added. enrich_game(conn, game, starters) is the
            integration point (starter inject + dress); evaluate_signals maps ScoredGame
            to the legacy dict and attaches output_tier / tier_basis / stake_multiplier.
2026-04-17  Ops: --sync-bet-ledger-only (no brief) + run_pipeline job_type bet_ledger_sync
            for recurring T−30 bet_ledger materialization without re-running briefs.
2026-04-17  Research: persist AVOID calls into brief_picks (pick_rank=0, signal='AVOID')
            so they can be tracked/grated later without generating bets.
2026-04-17  Fix starters column name: players.throws (was throw_hand). Update
            load_starters() SELECT + mapping, and enrich_game_with_starters().
2026-04-20  Run line analysis complete (OW era 2022-2025).
            LHP_FADE with ERA gate: RL cover 66.1% vs 60.0% BE,
            ROI +0.102. Added as active supplementary pick (score -1
            vs ML). S1H2: RL at breakeven (-0.2pp), shown as
            informational only. All other signals: ML primary only.
            MarketSnapshot extended with rl_available, away_rl_odds,
            home_rl_odds. LHP_FADE_RL signal added to evaluator.
2026-04-16  Finding 5 validated: NF4 implied prob gate confirmed as 60–67%.
            55–60% band shows only 5.5pp edge (not significant). Gate constants
            unchanged (already 0.60/0.67). Comments, signal reason, and constant
            block updated to document the band validation and independence from
            MV-F (wind-in + strong LHP co-occur too rarely to stack).
2026-04-16  NF4 signal added — Home Fav vs Strong LHP. Monitoring status,
            half-stake. Constants: NF4_HOME_IMP_LOW/HIGH, NF4_SP_ERA_MAX,
            NF4_OPS_MIN, NF4_MONTHS_OK. load_starters() extended to pull
            throws and team_rolling_ops. enrich_game_with_starters()
            helper injects these into game dict before signal evaluation.
            Priority order: S1+H2=1, MV-F=2, NF4=3, MV-B=4, S1=5, H3b=6.
2026-04-13 22:15 ET  Default DB from get_db_path(); repo root on sys.path for core.* imports.
2026-04-13 16:24 ET  Refactor: route sqlite3.connect() calls through core.db.connection.connect().

2026-04-11 19:00 ET  Feature: ⚡ MOVEMENT ALERT added to every pick in action
                     session briefs (early/afternoon/primary/late). When total
                     or ML moves ≥0.5 pts/10¢ since the earliest prior session
                     today, the alert is shown inline with the pick showing
                     direction and magnitude. brief_picks schema extended with
                     total_line column. ensure_brief_picks() auto-migrates.

2026-04-11 19:00 ET  Step 2: Updated Team_Strength_Analysis doc section 1.1
                     to correct prior analysis — MV-F WIN on Apr 10 prior
                     report was a fabricated retroactive pick, not a legitimate
                     graded pick from Apr 9. Apr 9 primary shows no signals.

2026-04-11 15:30 ET  Bug fix: prior report retroactive fallback now shows
                     [RETROACTIVE — not counted] label and NO graded result
                     box when brief_picks is empty for a date. Previously the
                     fallback rendered a full WIN/LOSS result using post-game
                     actual wind, fabricating a pick that was never shown
                     before the game was played. (Bug confirmed: Apr 10 prior
                     report showed MV-F WIN on PIT@CHC from Apr 9 — Apr 9
                     primary clearly showed "No confirmed signals fire".)

2026-04-11 15:30 ET  Feature: "Generated: YYYY-MM-DD HH:MM ET" timestamp
                     added to the first line of every session banner. Enables
                     validation of when each report was produced vs game times.

USAGE
-----
Morning brief   (after 9:00 AM odds pull):
    python generate_daily_brief.py --session morning

Primary brief   (after 5:00 PM odds pull)  ← act on these picks:
    python generate_daily_brief.py --session primary

Closing brief   (after 6:30 PM odds pull):
    python generate_daily_brief.py --session closing

Late West Coast games (after ~8:00 PM late-games odds pull):
    python load_odds.py --late-games --markets game
    python check_odds_ready.py
    python generate_daily_brief.py --session late

RECOVERY / RERUN FLAGS
----------------------
--date YYYY-MM-DD       Target a specific game date (default: today).
                        Use for reruns or catching up after a missed session.

--force                 Rerun even if a brief for this session/date already
                        exists in the brief_log table. Without this flag the
                        script warns and exits to prevent duplicate output.

--dry-run               Print the brief to the console but do NOT write to
                        brief_log or any output file.  Safe to run anytime.

--output FILE           Write the brief to FILE instead of (or in addition to)
                        the console.  Appends if file exists.

--no-file               Suppress file output; console only.  Overrides --output.

--warn-missing          Downgrade missing-data errors to warnings and continue.
                        By default the script exits if critical fields (ML odds,
                        totals) are absent.  Use after a partial odds pull.

--check-prereqs         Validate that the required odds pull for this session
                        was completed before generating the brief.  Exits with
                        a clear message if the pull is missing.  Recommended
                        for automated/scheduled runs.

--verbose               Print extra diagnostic rows from the DB (line counts,
                        data freshness, signal evaluation details).

--debug-wind            For each game, print wind classification debug to stdout:
                        raw DB wind_* / roof / park vs dressed wind_dir_label,
                        wind_in/out, suppression, env_ceiling (see evaluate_signals).

DOCX OUTPUT (DEFAULT)
--------------------
This script attempts to write a formatted Word (.docx) brief alongside the .txt
file by default (no flag required). Requires: pip install python-docx
File saved under outputs/briefs/ (default name: brief-SLATE-DATE_RUN-STAMP_ET.docx).

SESSION FILTERING MODEL
-----------------------
early / afternoon / primary all show the FULL unplayed slate for the
date — every game that hasn't started yet at the time you run the brief.
Session = when you pull odds and act, not which games start in that hour.

  · An all-afternoon slate correctly appears in --session primary because
    those games haven't started yet when you run it at 5:30 PM.
  · West Coast 10 PM ET games appear in --session afternoon (run 3:45 PM)
    because they haven't started yet.
  · A 1 PM game is dropped from --session primary (run 5:30 PM) because
    it already started. A 10-minute grace buffer prevents instant drop.

morning / closing show all games regardless of start time.

SIGNALS APPLIED (in priority order)
------------------------------------
1. MV-F   Wind in ≥10 mph + home fav −130 to −160  → fade ML (bet away)
          [CLV gate: only bet when CLV ≥ +0.5pp vs morning open — Mar 2026]
2. S1     Home team on W5+ win streak               → fade ML (bet away)
3. MV-B   Wind out ≥15 mph + home dog 35–42% impl.  → OVER  [band tightened Mar 2026]
4. S1+H2  W5+ home streak AND priced −130 to −160   → highest-priority stack
5. H3b    Wind out ≥10 mph + home dog (totals ctx)  → OVER (with MV-B only)

HOME TEAM INDICATOR
-------------------
Every matchup is shown as:  AWAY TEAM  vs  HOME TEAM (h)
"""

import argparse
import datetime
import os
import sqlite3
import sys
import textwrap
from pathlib import Path

_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
_REPO_ROOT = os.path.abspath(os.path.join(_SCRIPT_DIR, "..", ".."))
if _REPO_ROOT not in sys.path:
    sys.path.insert(0, _REPO_ROOT)

from core.db.connection import connect as db_connect, get_db_path

# ── Optional .env support ────────────────────────────────────────────────────
try:
    from dotenv import load_dotenv
    # Load in a stable order; avoids relying on "current working directory".
    load_dotenv(os.path.join(_REPO_ROOT, "config", ".env"), override=False)
    load_dotenv(os.path.join(_REPO_ROOT, ".env"), override=False)
    load_dotenv(override=False)  # fallback: cwd / parent-chain
except ImportError:
    pass

# ── ET timezone helper ────────────────────────────────────────────────────
try:
    from zoneinfo import ZoneInfo as _ZI
    _ET = _ZI("America/New_York")
except Exception:
    # Windows tzdata fallback + Python < 3.9 — fixed UTC-4 (EDT, covers MLB Apr-Oct season)
    _ET = datetime.timezone(datetime.timedelta(hours=-4))


def _now_et(override: datetime.datetime | None = None) -> datetime.datetime:
    """Current datetime in US/Eastern. Use for all user-facing timestamps."""
    if override is not None:
        return override
    return datetime.datetime.now(tz=_ET)


def _game_start_utc_dt(g: dict) -> datetime.datetime | None:
    """Parse game_start_utc to naive UTC datetime. None if missing or unparseable."""
    raw = g.get("game_start_utc") or ""
    if "T" not in raw:
        return None
    try:
        return datetime.datetime.fromisoformat(raw.rstrip("Z"))
    except ValueError:
        return None


def _game_started_or_in_progress_for_closing(game: dict, now: datetime.datetime) -> bool:
    """
    True if first pitch time is known and ``now`` is at or after start (UTC).
    Closing brief skips these games so active-signal lines are not shown for
    games already underway.
    """
    gst = _game_start_utc_dt(game)
    if gst is None:
        return False
    gst_utc = (
        gst.replace(tzinfo=datetime.timezone.utc)
        if gst.tzinfo is None
        else gst.astimezone(datetime.timezone.utc)
    )
    now_utc = now.astimezone(datetime.timezone.utc)
    return now_utc >= gst_utc


def _game_already_started_for_brief_picks(game: dict, now: datetime.datetime) -> bool:
    """
    True if game_start_utc is known and ``now`` is after first pitch (UTC).
    Used to skip brief_picks rows for retroactive/late-session writes.
    """
    gst = _game_start_utc_dt(game)
    if gst is None:
        return False
    gst_utc = (
        gst.replace(tzinfo=datetime.timezone.utc)
        if gst.tzinfo is None
        else gst.astimezone(datetime.timezone.utc)
    )
    now_utc = now.astimezone(datetime.timezone.utc)
    return now_utc > gst_utc


def _brief_pick_exists_cross_session(
    conn: sqlite3.Connection,
    game_date: str,
    game_pk: int,
    signal: str,
    market: str,
) -> bool:
    """True if any session already recorded this game/signal/market for game_date."""
    try:
        row = conn.execute(
            """
            SELECT 1 FROM brief_picks
            WHERE game_pk = ? AND game_date = ? AND signal = ? AND market = ?
            LIMIT 1
            """,
            (game_pk, game_date, signal, market),
        ).fetchone()
        return row is not None
    except Exception:
        return False


def _game_start_et(game: dict) -> str:
    """Return game start time as an ET string e.g. '7:10 PM ET'.
    Returns empty string if game_start_utc is missing or unparseable.
    Cross-platform: avoids %-I which fails on Windows.
    """
    raw = game.get("game_start_utc") or ""
    if not raw or "T" not in raw:
        return ""
    try:
        utc_dt = datetime.datetime.fromisoformat(
            raw.rstrip("Z")).replace(tzinfo=datetime.timezone.utc)
        et_dt  = utc_dt.astimezone(_ET)
        # %I gives zero-padded hour; lstrip removes leading zero cross-platform
        t = et_dt.strftime("%I:%M %p").lstrip("0") or et_dt.strftime("%I:%M %p")
        return f"{t} ET"
    except (ValueError, AttributeError):
        return ""

# ── S6 pitcher streak monitoring signal ───────────────────────────────────
# Loaded conditionally — non-fatal if file not yet present.
try:
    from s6_pitcher_streak import check_s6_pitcher_streak, log_s6_fire_to_db
    S6_AVAILABLE = True
except ImportError:
    S6_AVAILABLE = False

# ── Optional Word output (python-docx) ────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def build_docx_from_text(session: str, game_date: str, brief_text: str) -> "Document":
    """
    Render the exact same content as the .txt brief into a Word document.
    This keeps brief logic single-sourced in the text builders, while restoring
    the high-level Word formatting (title + section headers + matchup emphasis).
    """
    doc = Document()
    text = brief_text or ""
    # Strip ANSI escape codes (terminal-only) from Word output.
    try:
        import re

        text = re.sub(r"\x1b\[[0-9;]*m", "", text)
    except Exception:
        pass

    # Base style (match older docs: default font, compact spacing)
    try:
        style = doc.styles["Normal"]
        style.font.size = Pt(9)
    except Exception:
        pass

    def _add_line(line: str, *, size_pt: float = 9, bold: bool | None = None, align=None, color_hex: str | None = None) -> None:
        p = doc.add_paragraph()
        if align is not None:
            try:
                p.alignment = align
            except Exception:
                pass
        r = p.add_run(line)
        try:
            r.font.size = Pt(size_pt)
        except Exception:
            pass
        if bold is not None:
            try:
                r.bold = bool(bold)
            except Exception:
                pass
        if color_hex:
            try:
                r.font.color.rgb = RGBColor(int(color_hex[0:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16))
            except Exception:
                pass
        try:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
        except Exception:
            pass

    # Heuristic formatting based on the existing text layout.
    title_written = False
    for raw in text.splitlines():
        line = raw.rstrip("\n")
        s = line.strip()

        # Keep blank lines as blank paragraphs (for spacing consistency).
        if not s:
            doc.add_paragraph("")
            continue

        # Banner / separators: keep but de-emphasize.
        if set(s) <= {"═"} or set(s) <= {"─"}:
            _add_line(line, size_pt=8, color_hex="94A3B8")  # slate-300
            continue

        # Title line: first non-separator line.
        if not title_written and ("MLB" in s and ("BRIEF" in s or "REPORT" in s)):
            title_written = True
            _add_line(s, size_pt=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            continue

        # Section headings (Top Pick / No Signal / Ledger Summary etc.)
        if any(
            k in s.upper()
            for k in (
                "TOP PICK",
                "ADDITIONAL MODEL SELECTIONS",
                "NO SIGNAL",
                "BET LEDGER SUMMARY",
                "SIGNAL TRACKER",
                "HOT PITCHER",
                "PITCHER STREAK",
                "TODAY'S SLATE",
                "TODAYS SLATE",
            )
        ):
            _add_line(s, size_pt=12, bold=True)
            continue

        # Ledger / pick boxes (same visual weight as TOP pick)
        if s.startswith("┌") or s.startswith("│") or s.startswith("└"):
            _add_line(s, size_pt=9, bold=True)
            continue

        # Signal stack lines under the ASCII pick card
        if (
            s.startswith("Signal stack")
            or s.startswith("Core:")
            or s.startswith("Support:")
            or s.startswith("Minor:")
        ):
            _add_line(s, size_pt=9, bold=True)
            continue

        # Matchup lines: "XXX  vs  YYY (h)".
        if " vs " in s and "[" in s and "]" in s and "ET" in s:
            _add_line(s, size_pt=11, bold=True)
            continue

        # Default line.
        _add_line(s, size_pt=9)

    return doc


def strip_ansi(text: str) -> str:
    """Remove ANSI escape sequences (terminal-only formatting)."""
    try:
        import re

        return re.sub(r"\x1b\[[0-9;]*m", "", text or "")
    except Exception:
        return text or ""

# ── DB location (env / config/.env / cwd fallback via get_db_path) ───────
DB_PATH = Path(get_db_path())

# When False, any persistence helpers (signal_state / bet_ledger) must no-op.
PERSIST_WRITES = True

# ── Output directory for saved briefs ──────────────────────────────────────
# Standard location across the repo: <repo root>/outputs/briefs
OUTPUT_DIR = Path(_REPO_ROOT) / "outputs" / "briefs"


def _default_brief_file_stem(
    game_date: str,
    now_et: datetime.datetime,
    *,
    is_prior: bool,
    game_group_id: int | None = None,
) -> str:
    """
    Windows-safe default filename stem (no extension).
    Pattern: brief-{slate_date}_{run_yyyymmdd_hhmmss}_ET[_gN][_prior]
    - slate_date: YYYY-MM-DD the brief is for (``today`` for forward briefs, ``yesterday`` for prior report).
    - run_*: generation instant in America/New_York (unique across intraday runs).
    - _gN: optional ``pipeline_jobs.game_group_id`` (per-group group_brief) so same-minute runs
      do not overwrite each other.
    Stored in brief_log.output_file as the full path to the .txt file.
    """
    wall = now_et.astimezone(_ET) if now_et.tzinfo else now_et.replace(tzinfo=_ET)
    wall = wall.replace(microsecond=0)
    run_stamp = wall.strftime("%Y%m%d_%H%M%S")
    base = f"brief-{game_date}_{run_stamp}_ET"
    if game_group_id is not None and int(game_group_id) > 0:
        base = f"{base}_g{int(game_group_id)}"
    if is_prior:
        return f"{base}_prior"
    return base


# ── Signal thresholds ──────────────────────────────────────────────────────
WIND_OUT_MIN_MPH   = 10     # H3b standalone OVER threshold (unchanged)
WIND_IN_MIN_MPH    = 10
WIND_OUT_MVB_MPH   = 15     # MV-B raised to 15 mph for higher signal confidence
HOME_FAV_MV_F_LOW  = -130   # e.g. -130
HOME_FAV_MV_F_HIGH = -160   # e.g. -160  (more negative = bigger fav)
MV_F_CLV_GATE      = 0.5    # CLV gate (pp): only bet MV-F when opening away ML
                             # implied prob is ≥0.5pp lower than closing implied.
                             # Mar 2026 backtest (N=137): CLV≥+0.5pp → SBRO +24.0%
                             # OW +10.6%. CLV<+0.5pp → SBRO −9.4%. Both eras pass.
                             # This is a LIVE EXECUTION filter — compare current
                             # away ML to the morning opening line at bet time.
S1_PRICE_LOW       = -105   # S1 standalone lower price bound
S1_PRICE_HIGH      = -170   # S1 standalone upper price bound
DOG_IMPL_LOW       = 0.35   # 35%
DOG_IMPL_HIGH      = 0.42   # 42% — tightened from 45% (Mar 2026 CLV study)
                             # 42.5-45.0% bucket: -24.1% ROI on 18 non-Oracle games.
                             # ≤42% bucket: +18.5% ROI on 26 games.
                             # Near-even dogs (42-45% impl) are priced efficiently;
                             # the wind-out edge does not survive the vig at that range.
STREAK_THRESHOLD   = 5      # W5+  — used by S1+H2 stack
S1_STANDALONE_MIN  = 6      # standalone S1 requires W6+
S6_WIN_STREAK_MIN  = 7      # S6 pitcher fade threshold (backtest: +25% ROI, 7/8 seasons positive)

# ── Finding 4 / Finding 5: Home Fav vs Strong LHP (NF4) ─────────────────
# 2025 full-season regression (n=4,856 enriched game rows).
# Finding 4: high-OPS home teams (rolling OPS ≥ 0.736) vs strong LHP
#   (rolling 5-start ERA ≤ 3.04) win only 36.4% while priced at 62.6%.
#   Edge = +26.2pp, z = −3.11, p<.01.
# Finding 5 (band validation): the 60–67% implied band is where the signal
#   lives. z = −3.33, p<.01, n=34, win 35.3% vs 62.9% implied, edge +27.6pp.
#   The 55–60% band showed only 5.5pp edge (not significant) — confirmed
#   that moderate home favs do NOT show meaningful overpricing here.
#   The 60–67% gate (approximately −150 to −200 ML) is the validated range.
# Signal fires: fade home team (bet away ML).
# Independent of MV-F — wind-in and strong LHP co-occur too rarely to stack.
# Monitoring status — half-stake until N ≥ 50 live fires.
# September excluded: only month with inverted edge (66.7% win rate).
NF4_HOME_IMP_LOW   = 0.60   # validated lower bound (60%) — 55-60% not significant
NF4_HOME_IMP_HIGH  = 0.67   # validated upper bound (67%) — approx -150 to -200 ML
NF4_SP_ERA_MAX     = 3.04   # strong SP: rolling 5-start ERA ≤ p33 (2025 tertile)
NF4_OPS_MIN        = 0.736  # high-OPS home team: rolling OPS ≥ median (Finding 4)
NF4_MONTHS_OK      = {4, 5, 6, 7, 8}   # April–August only; September excluded

# H3b park factor gate — only fire at venues with PF ≥ this value.
# Derived from 2024-2025 data: wind-out OVER effect is strongest at
# neutral-to-hitter-friendly parks. Pitcher-friendly parks (PF < 98)
# suppress the signal even with wind-out because the run environment
# is already suppressed by the park itself.
H3B_MIN_PARK_FACTOR = 98

# H3b late-season flag — Aug/Sep performance is structurally weaker.
# CLV study (Mar 2026, n=316 H3b-eligible games 2023-2025):
#   Aug 2024: 9 games, 22.2% over rate.  Sep 2024: 8 games, 37.5% over rate.
#   Apr-Jul 2024: 50-54% over rate — completely normal.
# Root cause: 2024 second-half run environment was unusually suppressed
# league-wide (not a signal flaw). Flag informs the user without suppressing
# the signal. Combined with --late-season-stake this is sufficient defence.
H3B_LATE_SEASON_MONTHS = {8, 9}   # August and September

# ── Paper trading account ────────────────────────────────────────────────
PAPER_BANK_START    = 500.00   # Starting bankroll in dollars
PAPER_STAKE_TOP     =  10.00   # Top pick stake ($)
PAPER_STAKE_ADD     =   5.00   # Additional picks #2-#6 stake ($)
PAPER_STAKE_REST    =   0.00   # 7th pick onward — not wagered
PAPER_LATE_FACTOR   =   0.50   # Half-stake multiplier for late-season H3b (Aug/Sep)

# July OVER signal constants — data shows 52.3% OVER rate in July
# (p=0.0006, n=875) vs 46.1% for all other months. Market consistently
# under-sets totals in July due to seasonal lag in line-setting.
JULY_OVER_MONTHS       = {7}           # calendar month numbers
JULY_OVER_MIN_PF       = 100           # only at neutral/hitter-friendly parks
JULY_OVER_MIN_TOTAL    = 7.5           # skip very low totals (already adjusted)
JULY_OVER_MAX_TOTAL    = 11.0          # skip extreme totals (park-specific outliers)

# ── Legal caveat — appended to every brief output (txt and docx) ──────────
# Displayed once at the foot of every session file.
# Identical wording appears in both .txt and .docx formats.
CAVEAT = (
    "\n"
    + "\u2550" * 72 + "\n"
    "  EDUCATIONAL USE ONLY \u2014 NOT FINANCIAL ADVICE\n"
    + "\u2500" * 72 + "\n"
    "  This brief is produced by a personal statistical model for\n"
    "  research and educational purposes only. Nothing in this output\n"
    "  constitutes financial advice, investment advice, or a recommendation\n"
    "  to place any wager of any kind. Sports betting carries substantial\n"
    "  financial risk and is not appropriate for all individuals.\n"
    "\n"
    "  Past signal performance does not guarantee future results. Odds,\n"
    "  lines, and conditions can change materially between the time this\n"
    "  brief is generated and game time. Always verify every line\n"
    "  independently with your bookmaker before placing any bet.\n"
    "\n"
    "  You are solely responsible for all betting decisions and any\n"
    "  resulting financial outcomes. Never bet more than you can afford\n"
    "  to lose. If gambling is causing financial or personal harm,\n"
    "  contact the National Problem Gambling Helpline: 1-800-522-4700\n"
    "  or visit ncpgambling.org.\n"
    + "\u2550" * 72 + "\n"
)

# H3b park whitelist — wind-out OVER only fires at venues where the effect
# is historically strongest and wind readings are reliable open-air.
# Oracle (SF), Tropicana (TB), loanDepot (MIA) are already SUPPRESSED.
# This list excludes modern enclosed or semi-enclosed parks where the
# wind_mph reading does not reflect in-play conditions reliably.
H3B_PARK_WHITELIST = {
    "Wrigley Field",           # Chicago Cubs  — canonical wind-out park
    "Coors Field",             # Colorado Rockies — altitude + wind
    "Kauffman Stadium",        # Kansas City Royals — open, wind-prone
    "Globe Life Field",        # Texas Rangers — open retractable (open config)
    "Progressive Field",       # Cleveland Guardians — open, Lake Erie wind
    "PNC Park",                # Pittsburgh Pirates — open, river gap wind
    "Comerica Park",           # Detroit Tigers — open, Great Lakes wind
    "American Family Field",   # Milwaukee Brewers — open config days
    "Great American Ball Park",# Cincinnati Reds — open, Ohio River wind
    "Nationals Park",          # Washington Nationals — open
    "Camden Yards",            # Baltimore Orioles — open
    "Fenway Park",             # Boston Red Sox — open
    "Yankee Stadium",          # New York Yankees — open (wind tunnel effect)
    "Citi Field",              # New York Mets — open
    "Oakland Coliseum",        # Athletics — very open, bay wind
    # Added 2026-04-17: confirmed HIGH wind_effect in venues table,
    # PF >= 98, appeared in 2026 wind-out qualifying games with no H3b fire
    "Citizens Bank Park",      # Philadelphia Phillies — HIGH, PF 104
    "Target Field",            # Minnesota Twins — HIGH, PF 99
    "Rate Field",              # Chicago White Sox — HIGH, PF 102
}

# ── Session → expected pull window (for --check-prereqs) ──────────────────
SESSION_PULL_WINDOW = {
    "morning": ("06:00", "09:30"),   # after 9 AM odds pull
    "primary": ("09:30", "17:30"),   # after 5 PM odds pull
    "closing": ("17:30", "23:59"),   # after 6:30 PM odds pull
    "late":    ("20:00", "23:59"),   # after ~8 PM late-games pull
}

# ET wall-clock → session for hybrid mode (--as-of-time / full --as-of).
# Half-open minute ranges [start, end) on a 0–1440 minute-of-day axis. Does not alter
# SESSION_PULL_WINDOW (used by --check-prereqs). Boundaries align with simulate_day /
# operator schedule for early/afternoon/primary/late.
#
# ``closing`` is NOT derived from the clock — only ``--session closing`` (see ``main``).
# The former 18:45–20:15 window maps here as ``primary`` so late-day runs still get the
# full pick card / docx layout until ``late`` (west-coast window).
_SESSION_WINDOW_RANGES_ET: tuple[tuple[int, int, str], ...] = (
    (0, 570, "morning"),       # 00:00–09:30
    (570, 735, "early"),       # 09:30–12:15
    (735, 945, "afternoon"),   # 12:15–15:45
    (945, 1215, "primary"),    # 15:45–20:15 (was split at 18:45 into primary+closing)
    (1215, 1440, "late"),      # 20:15–24:00
)
SESSION_WINDOWS = tuple(
    (f"{a // 60:02d}:{a % 60:02d}", f"{b // 60:02d}:{b % 60:02d}", s)
    for a, b, s in _SESSION_WINDOW_RANGES_ET
)


def _minute_of_day_et(dt: datetime.datetime) -> int:
    """Minute index 0..1439 in America/New_York."""
    if dt.tzinfo is None:
        dt = dt.replace(tzinfo=_ET)
    else:
        dt = dt.astimezone(_ET)
    return dt.hour * 60 + dt.minute


def _session_from_et_datetime(dt: datetime.datetime) -> str:
    """Map an ET instant to a forward session using SESSION_WINDOWS ranges."""
    m = _minute_of_day_et(dt)
    for lo, hi, sess in _SESSION_WINDOW_RANGES_ET:
        if lo <= m < hi:
            return sess
    return "late"


# ═══════════════════════════════════════════════════════════════════════════
# Utility helpers
# ═══════════════════════════════════════════════════════════════════════════

def american_to_implied(odds: int) -> float:
    """Convert American odds integer to implied probability (0–1)."""
    if odds is None:
        return None
    if odds > 0:
        return 100 / (odds + 100)
    else:
        return abs(odds) / (abs(odds) + 100)


def implied_to_american(prob: float) -> str:
    """Convert implied probability back to American odds string."""
    if prob is None or prob <= 0 or prob >= 1:
        return "N/A"
    if prob >= 0.5:
        odds = -(prob / (1 - prob)) * 100
        return f"{int(round(odds))}"
    else:
        odds = ((1 - prob) / prob) * 100
        return f"+{int(round(odds))}"


def wind_direction_label(direction: str) -> str:
    """Return OUT / IN / CROSS / CALM given a wind_direction string."""
    if not direction:
        return "UNKNOWN"
    d = direction.upper()
    # Treat 'OUT*', 'BLOWING OUT', 'OUT TO CF', etc. as OUT
    if any(k in d for k in ("OUT", "BLOWING OUT")):
        return "OUT"
    if any(k in d for k in ("IN", "BLOWING IN")):
        return "IN"
    if any(k in d for k in ("CROSS", "LEFT", "RIGHT")):
        return "CROSS"
    if any(k in d for k in ("CALM", "STILL", "0 MPH")):
        return "CALM"
    return direction  # pass through if unrecognised


def print_wind_debug_for_game(game: dict, fdg) -> None:
    """
    Print DB wind fields vs dressed ``FullyDressedGame.environment`` to stdout.
    Used with ``--debug-wind`` to trace OUT/IN/CROSS, suppression, and env_ceiling.
    """
    try:
        env = fdg.environment
        ids = fdg.identifiers
    except Exception as e:
        print(f"\n  [debug-wind] (could not read dressed game: {e})\n")
        return
    gpk = game.get("game_pk")
    ab = f"{ids.away_team_abbr}@{ids.home_team_abbr}"
    raw_dir = game.get("wind_direction")
    brief_lbl = wind_direction_label(str(raw_dir or ""))
    mph = game.get("wind_mph")
    try:
        mph_f = float(mph) if mph is not None else None
    except (TypeError, ValueError):
        mph_f = None
    mph_gate = 10.0
    mph_ok = mph_f is not None and mph_f >= mph_gate
    print("\n  [debug-wind] ───────────────────────────────────────────────────────────")
    print(f"  [debug-wind] game_pk={gpk}  {ab}  {ids.venue_name or ''}")
    print(
        f"  [debug-wind] DB row: wind_mph={mph!r}  temp_f={game.get('temp_f')!r}  "
        f"sky={game.get('sky_condition')!r}"
    )
    print(
        f"  [debug-wind] DB row: wind_direction(raw)={raw_dir!r}  "
        f"wind_effect={game.get('wind_effect')!r}  "
        f"wind_source={game.get('wind_source')!r}"
    )
    print(
        f"  [debug-wind] DB row: roof_type={game.get('roof_type')!r}  "
        f"orientation_hp={game.get('orientation_hp')!r}  "
        f"park_factor_runs={game.get('park_factor_runs')!r}"
    )
    print(
        f"  [debug-wind] brief.wind_direction_label(raw) -> {brief_lbl!r}  "
        f"(mph>={mph_gate} for in/out gates: {mph_ok})"
    )
    print(
        f"  [debug-wind] dressed: wind_dir_label={env.wind_dir_label!r}  "
        f"wind_in={env.wind_in}  wind_out={env.wind_out}"
    )
    print(
        f"  [debug-wind] dressed: is_wind_suppressed={env.is_wind_suppressed}  "
        f"env_ceiling={env.env_ceiling!r}  h3b_eligible={env.h3b_eligible}"
    )
    print(
        f"  [debug-wind] dressed: is_retractable={env.is_retractable}  "
        f"roof_status_known={env.roof_status_known}  wind_source={env.wind_source!r}"
    )
    if brief_lbl != env.wind_dir_label:
        print(
            f"  [debug-wind] NOTE: brief label {brief_lbl!r} != dressed {env.wind_dir_label!r} "
            f"(fully_dressed_game uses its own normalizer on the same row)."
        )
    print("  [debug-wind] ───────────────────────────────────────────────────────────\n")


def banner(text: str, width: int = 72) -> str:
    bar = "═" * width
    return f"\n{bar}\n  {text}\n{bar}"


def section(text: str, width: int = 72) -> str:
    return f"\n{'─' * width}\n  {text}\n{'─' * width}"


def fmt_odds(val) -> str:
    if val is None:
        return "N/A"
    v = int(val)
    return f"+{v}" if v > 0 else str(v)


def fmt_total(val) -> str:
    if val is None:
        return "N/A"
    return f"O/U {float(val):.1f}"


def missing(label: str, warn_only: bool) -> None:
    msg = f"[DATA MISSING] {label}"
    if warn_only:
        print(f"  ⚠  {msg}")
    else:
        print(f"\n  ✗  {msg}")
        print("     Re-run the required odds pull, or use --warn-missing to continue anyway.")
        sys.exit(1)


# ═══════════════════════════════════════════════════════════════════════════
# Database helpers
# ═══════════════════════════════════════════════════════════════════════════

def open_db(db_path: Path) -> sqlite3.Connection:
    if not db_path.exists():
        print(f"\n✗  Database not found: {db_path}")
        print("   Set MLB_DB_PATH or add MLB_DB_PATH=... to config/.env")
        sys.exit(1)
    conn = db_connect(str(db_path), timeout=30)  # wait up to 30s if Scout holds a lock
    conn.row_factory = sqlite3.Row
    return conn


def check_prereqs(conn: sqlite3.Connection, game_date: str, session: str) -> None:
    """
    Verify that the required odds pull for this session was completed today.
    Exits with a helpful message if the pull appears missing.
    """
    pull_map = {
        "morning": "pregame",
        "early": "pregame",
        "afternoon": "pregame",
        "primary": "pregame",
        "closing": "pregame",
        "late": "pregame",
    }
    pull_type = pull_map[session]

    # Look for an ingest log entry today for the right pull type
    cur = conn.execute(
        """
        SELECT pulled_at_utc, api_requests_used, status
        FROM   odds_ingest_log
        WHERE  date(pulled_at_utc) = ?
          AND  pull_type = ?
        ORDER  BY pulled_at_utc DESC
        LIMIT  1
        """,
        (game_date, pull_type),
    )
    row = cur.fetchone()
    if not row:
        print(f"\n✗  Prereq check failed for --session {session}:")
        print(f"   No '{pull_type}' pull found in odds_ingest_log for {game_date}.")
        print(f"   Run the required odds pull first:")
        print(f"   python load_odds.py --pregame --markets game")
        sys.exit(1)
    if row["status"] != "success":
        print(f"\n⚠  Last odds pull on {game_date} has status='{row['status']}' (not 'success').")
        print(f"   Consider re-running: python load_odds.py --pregame --markets game")
        sys.exit(1)
    print(f"  ✓ Prereq: odds pull found ({row['pulled_at_utc']}, {row['api_requests_used']} req used)")


def already_ran(
    conn: sqlite3.Connection,
    game_date: str,
    session: str,
    *,
    game_group_id: int | None = None,
) -> bool:
    """
    Return True if a brief_log row would block a non-``--force`` run.

    - Without ``--game-group-id`` (or id≤0): a row for (game_date, session) blocks only if
      ``output_file`` is set and that path still exists on disk (stale/ missing files allow re-run).
    - With ``--game-group-id N`` (N>0): only a row for the same (game_date, session, N) blocks,
      and only when the stored ``output_file`` exists.
    """
    try:
        if game_group_id is not None and int(game_group_id) > 0:
            g = int(game_group_id)
            cur = conn.execute(
                """
                SELECT output_file FROM brief_log
                WHERE game_date=? AND session=? AND IFNULL(game_group_id,0)=?
                LIMIT 1
                """,
                (game_date, session, g),
            )
        else:
            cur = conn.execute(
                "SELECT output_file FROM brief_log WHERE game_date=? AND session=? LIMIT 1",
                (game_date, session),
            )
        row = cur.fetchone()
        if not row:
            return False
        of = (row["output_file"] if hasattr(row, "keys") else row[0]) or ""
        of = str(of).strip()
        if not of:
            print(
                "  [note] brief_log has a row for this date/session but no output_file; "
                "allowing a fresh run (not a duplicate)."
            )
            return False
        out_path = Path(of)
        if not out_path.is_file():
            print(
                f"  [note] brief_log lists {of} but that file is missing; "
                "allowing a fresh run. Use --force to overwrite a still-existing file."
            )
            return False
        return True
    except sqlite3.OperationalError:
        return False


def _ensure_brief_log_game_group_id_column(conn: sqlite3.Connection) -> None:
    try:
        cols = {r[1] for r in conn.execute("PRAGMA table_info(brief_log)").fetchall()}
    except sqlite3.OperationalError:
        return
    if "game_group_id" in cols:
        return
    try:
        conn.execute("ALTER TABLE brief_log ADD COLUMN game_group_id INTEGER")
        conn.commit()
    except sqlite3.OperationalError:
        pass


def ensure_brief_log(conn: sqlite3.Connection) -> None:
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS brief_log (
            id          INTEGER PRIMARY KEY AUTOINCREMENT,
            game_date   TEXT    NOT NULL,
            session     TEXT    NOT NULL,
            generated_at TEXT   NOT NULL,
            games_covered INTEGER,
            picks_count   INTEGER,
            output_file   TEXT
        )
        """
    )
    conn.commit()
    _ensure_brief_log_game_group_id_column(conn)
    ensure_daily_pnl(conn)   # create paper-trading ledger alongside
    ensure_brief_picks(conn)  # create confirmed-picks ledger alongside


def ensure_daily_pnl(conn: sqlite3.Connection) -> None:
    """Create the paper-trading ledger table if it does not exist."""
    conn.execute("""
        CREATE TABLE IF NOT EXISTS daily_pnl (
            id              INTEGER PRIMARY KEY AUTOINCREMENT,
            game_date       TEXT    NOT NULL,
            game_pk         INTEGER NOT NULL,
            signal          TEXT    NOT NULL,   -- e.g. S1+H2, MV-F, H3b
            pick_tier       TEXT    NOT NULL,   -- top / additional / rest
            bet             TEXT    NOT NULL,   -- e.g. "BOS ML" or "OVER 8.5"
            market          TEXT    NOT NULL,   -- ML or TOTAL
            odds            INTEGER,            -- American odds
            stake_dollars   REAL    NOT NULL,
            late_season     INTEGER NOT NULL DEFAULT 0, -- 1 = half-stake applied
            result          TEXT    NOT NULL,   -- WIN / LOSS / PUSH / NO RESULT
            pnl_units       REAL    NOT NULL,
            pnl_dollars     REAL    NOT NULL,
            recorded_at     TEXT    NOT NULL    -- ET timestamp
        )
    """)
    conn.commit()


def record_daily_pnl(conn: sqlite3.Connection, rows: list,
                     now: datetime.datetime | None = None) -> None:
    """Insert paper-trading results for a game date. Idempotent on game_date+game_pk+signal."""
    if not rows:
        return
    if now is None:
        now = _now_et()
    for row in rows:
        conn.execute("""
            INSERT OR IGNORE INTO daily_pnl
                (game_date, game_pk, signal, pick_tier, bet, market, odds,
                 stake_dollars, late_season, result, pnl_units, pnl_dollars, recorded_at)
            VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)
        """, (
            row["game_date"], row["game_pk"], row["signal"], row["pick_tier"],
            row["bet"], row["market"], row["odds"],
            row["stake_dollars"], row["late_season"],
            row["result"], row["pnl_units"], row["pnl_dollars"],
            now.strftime("%Y-%m-%d %H:%M ET"),
        ))
    conn.commit()


def load_season_pnl(conn: sqlite3.Connection, game_date: str) -> dict:
    """Load season-to-date paper P&L for the season containing game_date.
    Uses the seasons table (postseason_start as season end).
    Returns dict with keys: rows, total_dollars, wins, losses, pushes, bank.
    """
    try:
        season_year = int(game_date[:4])
        row = conn.execute(
            "SELECT season_start, postseason_start FROM seasons WHERE season=?",
            (season_year,)
        ).fetchone()
        if row:
            season_start = row["season_start"]
            season_end   = row["postseason_start"] or f"{season_year}-10-01"
        else:
            season_start = f"{season_year}-04-01"
            season_end   = f"{season_year}-10-01"
    except (ValueError, TypeError):
        return {"rows": [], "total_dollars": 0.0, "wins": 0,
                "losses": 0, "pushes": 0, "bank": PAPER_BANK_START}

    try:
        rows = conn.execute("""
            SELECT game_date, signal, pick_tier, bet, odds,
                   stake_dollars, late_season, result, pnl_units, pnl_dollars
            FROM   daily_pnl
            WHERE  game_date >= ? AND game_date < ?
            ORDER  BY game_date, id
        """, (season_start, season_end)).fetchall()
    except Exception:
        return {"rows": [], "total_dollars": 0.0, "wins": 0,
                "losses": 0, "pushes": 0, "bank": PAPER_BANK_START}

    total = sum(r["pnl_dollars"] for r in rows)
    wins  = sum(1 for r in rows if r["result"] == "WIN")
    losses= sum(1 for r in rows if r["result"] == "LOSS")
    pushes= sum(1 for r in rows if r["result"] == "PUSH")
    return {
        "rows":          rows,
        "total_dollars": round(total, 2),
        "wins":          wins,
        "losses":        losses,
        "pushes":        pushes,
        "bank":          round(PAPER_BANK_START + total, 2),
    }


def ensure_brief_picks(conn: sqlite3.Connection) -> None:
    """Create table that stores confirmed picks shown in each brief.
    Migrates existing tables: total_line, model_version."""
    conn.execute("""
        CREATE TABLE IF NOT EXISTS brief_picks (
            id          INTEGER PRIMARY KEY AUTOINCREMENT,
            game_date   TEXT    NOT NULL,
            session     TEXT    NOT NULL,   -- primary / early / afternoon
            game_pk     INTEGER NOT NULL,
            pick_rank   INTEGER NOT NULL,   -- 1=top, 2-6=additional
            signal      TEXT    NOT NULL,
            bet         TEXT    NOT NULL,
            market      TEXT    NOT NULL,
            odds        INTEGER,
            total_line  REAL,              -- stored when market=TOTAL; enables movement alert
            recorded_at TEXT    NOT NULL,
            model_version TEXT DEFAULT 'legacy',
            UNIQUE (game_date, session, game_pk, signal)
        )
    """)
    conn.commit()
    try:
        cols = [r[1] for r in conn.execute("PRAGMA table_info(brief_picks)").fetchall()]
        if cols:
            if "total_line" not in cols:
                conn.execute("ALTER TABLE brief_picks ADD COLUMN total_line REAL")
            if "model_version" not in cols:
                conn.execute(
                    "ALTER TABLE brief_picks ADD COLUMN model_version TEXT DEFAULT 'legacy'"
                )
            conn.commit()
    except Exception:
        pass


def ensure_signal_state(conn: sqlite3.Connection) -> None:
    """Create signal_state table if it does not exist (append-only signal ledger)."""
    conn.execute("""
        CREATE TABLE IF NOT EXISTS signal_state (
            id          INTEGER PRIMARY KEY,
            game_date   TEXT,
            game_pk     INTEGER,
            market_type TEXT,     -- 'moneyline','spread','total'
            signal_type TEXT,     -- 'top','next','avoid'
            bet         TEXT,
            odds        INTEGER,
            session     TEXT,     -- morning, early, primary, etc.
            recorded_at TEXT
        )
    """)
    conn.commit()


def ensure_bet_ledger(conn: sqlite3.Connection) -> None:
    """Create bet_ledger table if it does not exist and enforce idempotency."""
    if not PERSIST_WRITES:
        return
    conn.execute("""
        CREATE TABLE IF NOT EXISTS bet_ledger (
            id            INTEGER PRIMARY KEY,
            game_date     TEXT,
            game_pk       INTEGER,
            market_type   TEXT,
            bet           TEXT,
            odds_taken    INTEGER,
            stake_units   REAL,
            signal_at_time TEXT,  -- 'top','next','avoid'
            session       TEXT,
            placed_at     TEXT,
            result        TEXT,   -- 'win','loss','push'
            pnl_units     REAL
        )
    """)
    # Allow top / next / avoid on the same (game_pk, market_type) as separate rows.
    try:
        conn.execute("DROP INDEX IF EXISTS idx_bet_ledger_game_market")
    except Exception:
        pass
    conn.execute("""
        CREATE UNIQUE INDEX IF NOT EXISTS idx_bet_ledger_game_market_signal
        ON bet_ledger (game_pk, market_type, IFNULL(signal_at_time, ''))
    """)
    conn.commit()


def ensure_bet_snapshots(conn: sqlite3.Connection) -> None:
    """
    Persistent snapshot of bet decisions at placement time (stake>0).
    Source of truth for PRIOR reports (no recomputation).
    """
    conn.execute(
        """
        CREATE TABLE IF NOT EXISTS bet_snapshots (
            game_date      TEXT    NOT NULL,
            game_pk        INTEGER NOT NULL,
            market_type    TEXT    NOT NULL,   -- 'ML' | 'TOTAL' | 'RL'
            bet_side       TEXT    NOT NULL,
            bet            TEXT    NOT NULL,
            odds_taken     INTEGER,
            score          INTEGER NOT NULL,
            model_p        REAL,
            implied_p      REAL,
            edge           REAL,
            eval_status    TEXT,
            signals_used   TEXT,               -- JSON list[str]
            placed_at      TEXT    NOT NULL    -- ET timestamp string
        )
        """
    )
    conn.execute(
        """
        CREATE UNIQUE INDEX IF NOT EXISTS ux_bet_snapshots_key
        ON bet_snapshots (game_date, game_pk, market_type)
        """
    )
    conn.commit()
    # Add columns for existing DBs (safe best-effort)
    try:
        cols = [r[1] for r in conn.execute("PRAGMA table_info(bet_snapshots)").fetchall()]
        if cols and "eval_status" not in cols:
            conn.execute("ALTER TABLE bet_snapshots ADD COLUMN eval_status TEXT")
            conn.commit()
    except Exception:
        pass


def save_bet_snapshot(
    conn: sqlite3.Connection,
    game_date: str,
    game_pk: int,
    market_type: str,
    bet: str,
    odds: int | None,
    scored_game: object,
    signals: list,
    *,
    placed_at: str | None = None,
) -> None:
    """
    Persist a snapshot of the bet decision at placement time.
    Called ONLY when stake > 0.
    """
    import json

    try:
        ensure_bet_snapshots(conn)
    except Exception:
        return

    now_et_string = _now_et().strftime("%Y-%m-%d %H:%M ET")
    placed_at_txt = str(placed_at or "").strip() or now_et_string

    # ``scored_game`` is expected to be the per-market eval dict (mm) with keys:
    # score, model_p, implied_p, edge, best_side.
    try:
        score = int(getattr(scored_game, "score", None) or scored_game.get("score") or 0)
    except Exception:
        score = 0
    try:
        model_p = getattr(scored_game, "model_p", None)
        if model_p is None and isinstance(scored_game, dict):
            model_p = scored_game.get("model_p")
    except Exception:
        model_p = None
    try:
        implied_p = getattr(scored_game, "implied_p", None)
        if implied_p is None and isinstance(scored_game, dict):
            implied_p = scored_game.get("implied_p")
    except Exception:
        implied_p = None
    try:
        edge = getattr(scored_game, "edge", None)
        if edge is None and isinstance(scored_game, dict):
            edge = scored_game.get("edge")
    except Exception:
        edge = None
    try:
        bet_side = getattr(scored_game, "best_side", None)
        if bet_side is None and isinstance(scored_game, dict):
            bet_side = scored_game.get("best_side")
        bet_side = str(bet_side or "")
    except Exception:
        bet_side = ""

    try:
        eval_status = None
        try:
            if isinstance(scored_game, dict):
                ev = scored_game.get("eval_status")
                eval_status = (str(ev).strip() if ev is not None else None) or None
        except Exception:
            eval_status = None

        conn.execute(
            """
            INSERT OR REPLACE INTO bet_snapshots
                (game_date, game_pk, market_type, bet_side, bet,
                 odds_taken, score, model_p, implied_p, edge, eval_status,
                 signals_used, placed_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                str(game_date),
                int(game_pk),
                str(market_type or "").strip().upper(),
                bet_side,
                str(bet or ""),
                int(odds) if odds is not None else None,
                int(score),
                float(model_p) if model_p is not None else None,
                float(implied_p) if implied_p is not None else None,
                float(edge) if edge is not None else None,
                eval_status,
                json.dumps(list(signals or [])),
                placed_at_txt,
            ),
        )
        conn.commit()
    except Exception:
        return


def backfill_missing_bet_snapshots_from_ledger(conn: sqlite3.Connection, game_date: str) -> int:
    """
    If bet_ledger already has staked rows (older runs before snapshots existed),
    reconstruct missing snapshots using the bet_ledger placed_at timestamp.
    This does not change PRIOR display logic (PRIOR still reads snapshots only).
    """
    if conn is None or not PERSIST_WRITES:
        return 0
    try:
        ensure_bet_snapshots(conn)
    except Exception:
        return 0

    try:
        rows = conn.execute(
            """
            SELECT game_pk, market_type, bet, odds_taken, placed_at
            FROM bet_ledger
            WHERE game_date = ?
              AND stake_units > 0
            """,
            (game_date,),
        ).fetchall()
        rows = [dict(r) for r in rows]
    except Exception:
        return 0

    if not rows:
        return 0

    created = 0
    for r in rows:
        gpk = int(r["game_pk"])
        mt_u = (r.get("market_type") or "").strip().lower()
        market_type = "ML" if mt_u == "moneyline" else ("TOTAL" if mt_u == "total" else ("RL" if mt_u in ("spread", "runline") else None))
        if market_type is None:
            continue

        # Skip only if snapshot exists with real content; otherwise overwrite.
        try:
            ex = conn.execute(
                """
                SELECT bet, odds_taken, signals_used
                FROM bet_snapshots
                WHERE game_date=? AND game_pk=? AND market_type=?
                LIMIT 1
                """,
                (game_date, gpk, market_type),
            ).fetchone()
            if ex:
                bet_existing = (ex[0] or "").strip()
                odds_existing = ex[1]
                sig_existing = (ex[2] or "").strip()
                if bet_existing and odds_existing is not None and sig_existing and sig_existing != "[]":
                    continue
        except Exception:
            pass

        placed_at = str(r.get("placed_at") or "").strip()
        # Pull *actual* signal labels from brief_picks (source of truth for what the brief showed),
        # rather than recomputing. Stored as a comma-separated string in brief_picks.signal.
        signals_used: list[str] = []
        try:
            mt = market_type
            row_sig = conn.execute(
                """
                SELECT signal
                FROM brief_picks
                WHERE game_date = ?
                  AND game_pk = ?
                  AND UPPER(market) = ?
                  AND pick_rank > 0
                ORDER BY recorded_at DESC, pick_rank ASC
                LIMIT 1
                """,
                (game_date, gpk, mt),
            ).fetchone()
            if row_sig and row_sig[0]:
                raw = str(row_sig[0])
                signals_used = [s.strip() for s in raw.split(",") if s.strip()]
        except Exception:
            signals_used = []

        # Save snapshot using ledger bet/odds + brief_picks signals.
        mm_stub = {"score": 0, "model_p": None, "implied_p": None, "edge": None, "best_side": "", "eval_status": "BET"}
        save_bet_snapshot(
            conn,
            game_date,
            gpk,
            market_type,
            str(r.get("bet") or ""),
            int(r["odds_taken"]) if r.get("odds_taken") is not None else None,
            mm_stub,
            signals_used,
            placed_at=placed_at,
        )
        created += 1

    return created

    now_et_string = _now_et().strftime("%Y-%m-%d %H:%M ET")
    placed_at_txt = str(placed_at or "").strip() or now_et_string

    # Pull the required fields off the market-eval dict we store on ScoredGame
    try:
        score = int(scored_game.get("score") or 0)
    except Exception:
        score = 0
    try:
        model_p = scored_game.get("model_p")
        implied_p = scored_game.get("implied_p")
        edge = scored_game.get("edge")
        bet_side = str(scored_game.get("best_side") or "")
    except Exception:
        model_p = None
        implied_p = None
        edge = None
        bet_side = ""

    try:
        conn.execute(
            """
            INSERT OR REPLACE INTO bet_snapshots
                (game_date, game_pk, market_type, bet_side, bet,
                 odds_taken, score, model_p, implied_p, edge,
                 signals_used, placed_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                game_date,
                int(game_pk),
                str(market_type or "").strip().upper(),
                bet_side,
                str(bet or ""),
                int(odds) if odds is not None else None,
                int(score),
                float(model_p) if model_p is not None else None,
                float(implied_p) if implied_p is not None else None,
                float(edge) if edge is not None else None,
                json.dumps(list(signals or [])),
                placed_at_txt,
            ),
        )
        conn.commit()
    except Exception:
        return


def save_signal_state(conn: sqlite3.Connection, game_date: str, session: str,
                      top_entry: dict | None, next_entries: list,
                      avoid_entries: list, now: datetime.datetime | None = None) -> None:
    """Persist TOP / NEXT (up to 5) / AVOID signals into signal_state (append-only)."""
    if not PERSIST_WRITES:
        return
    if conn is None or not session:
        return
    if now is None:
        now = _now_et()
    recorded_at = now.strftime("%Y-%m-%d %H:%M ET")

    try:
        ensure_signal_state(conn)
    except Exception:
        return

    def _market_type_from_market(m: str | None) -> str | None:
        m = (m or "").upper().strip()
        if m in ("ML", "MONEYLINE"):
            return "moneyline"
        if m in ("TOTAL", "OU", "O/U"):
            return "total"
        if m in ("RL", "RUNLINE", "SPREAD"):
            return "spread"
        return None

    def _best_pick_fields(entry: dict) -> tuple[str | None, str | None, int | None]:
        try:
            p = sorted(entry["sigs"]["picks"], key=lambda x: x["priority"])[0]
            bet = p.get("bet")
            market_type = _market_type_from_market(p.get("market"))
            odds = _parse_odds(p.get("odds", ""))
            return market_type, bet, odds
        except Exception:
            return None, None, None

    def _avoid_fields(entry: dict) -> tuple[str | None, str | None]:
        """
        Return (market_type, bet_text) for avoid rows.
        Keep it descriptive and non-structured (no forced bet schema).
        """
        g = (entry.get("game") or {})
        sigs = (entry.get("sigs") or {})
        reason = (sigs.get("avoid_reason") or "").strip()
        reason_u = reason.upper()

        home = (g.get("home_abbr") or "").strip()
        away = (g.get("away_abbr") or "").strip()
        total_line = g.get("total_line")

        market_type = None
        bet_text = None

        # Heuristics based on existing avoid_reason text patterns.
        if (
            "WIND SIGNALS SUPPRESSED" in reason_u
            or ("SUPPRESSED" in reason_u and "WEATHER" in reason_u)
            or "NO WEATHER EDGE" in reason_u
        ):
            market_type = "environment"
            bet_text = "Weather-signal class (no single ticket) — see brief reason"
        elif "AVOID HOME ML" in reason_u and home:
            market_type = "moneyline"
            bet_text = f"Avoid {home} ML"
        elif "AVOID AWAY ML" in reason_u and away:
            market_type = "moneyline"
            bet_text = f"Avoid {away} ML"
        elif "DO NOT BET OVER" in reason_u and total_line is not None:
            market_type = "total"
            bet_text = f"Avoid OVER {total_line}"
        elif "DO NOT BET UNDER" in reason_u and total_line is not None:
            market_type = "total"
            bet_text = f"Avoid UNDER {total_line}"
        elif "OVER" in reason_u and "DO NOT BET" in reason_u and total_line is not None:
            market_type = "total"
            bet_text = f"Avoid OVER {total_line}"
        elif "UNDER" in reason_u and "DO NOT BET" in reason_u and total_line is not None:
            market_type = "total"
            bet_text = f"Avoid UNDER {total_line}"
        elif bet_text is None and "ML" in reason_u:
            market_type = "moneyline"
            bet_text = f"Avoid {home} ML" if home else "Avoid ML"
        elif bet_text is None and (
            "TOTAL" in reason_u or "O/U" in reason_u or "OVER" in reason_u or "UNDER" in reason_u
        ):
            market_type = "total"
            bet_text = f"Avoid total" if total_line is None else f"Avoid total ({total_line})"

        # Fallback: descriptive text to make it traceable.
        if not bet_text:
            bet_text = ("Avoid: " + reason) if reason else "Avoid"

        return market_type, bet_text

    rows = []
    if top_entry is not None:
        g = (top_entry.get("game") or {})
        market_type, bet, odds = _best_pick_fields(top_entry)
        if market_type is not None:
            rows.append((game_date, g.get("game_pk"), market_type, "top", bet, odds, session, recorded_at))

    for ne in (next_entries or []):
        g = (ne.get("game") or {})
        market_type, bet, odds = _best_pick_fields(ne)
        if market_type is not None:
            rows.append((game_date, g.get("game_pk"), market_type, "next", bet, odds, session, recorded_at))

    for e in (avoid_entries or []):
        g = (e.get("game") or {})
        mt, bt = _avoid_fields(e)
        rows.append((game_date, g.get("game_pk"), mt, "avoid", bt, None, session, recorded_at))

    if not rows:
        return

    try:
        conn.executemany("""
            INSERT INTO signal_state
                (game_date, game_pk, market_type, signal_type, bet, odds, session, recorded_at)
            VALUES (?,?,?,?,?,?,?,?)
        """, rows)
        conn.commit()
    except Exception:
        pass


def _parse_recorded_at_et_ledger(s: str | None) -> datetime.datetime | None:
    if not s:
        return None
    raw = str(s).strip()
    for fmt in ("%Y-%m-%d %H:%M ET", "%Y-%m-%d %I:%M %p ET"):
        try:
            parsed = datetime.datetime.strptime(raw, fmt)
            return parsed.replace(tzinfo=_ET)
        except Exception:
            continue
    return None


def _ledger_latest_from_signal_rows(
    sig_rows: list,
    start_by_pk: dict[int, datetime.datetime],
    now_et: datetime.datetime,
    *,
    pregame_window_only: bool,
) -> dict[tuple[int, str, str], tuple[datetime.datetime, sqlite3.Row]]:
    """
    Latest signal_state row per (game_pk, market_type, signal_type).
    When pregame_window_only=True, only rows inside [start−30m, start) apply.
    """
    latest: dict[tuple[int, str, str], tuple[datetime.datetime, sqlite3.Row]] = {}
    for r in sig_rows:
        gpk = r["game_pk"]
        mt = r["market_type"]
        st = (r["signal_type"] or "").strip()
        if gpk is None or not mt or st not in ("top", "next", "avoid"):
            continue
        start_et = start_by_pk.get(int(gpk))
        if start_et is None:
            continue
        if pregame_window_only:
            window_start = start_et - datetime.timedelta(minutes=30)
            if not (window_start <= now_et < start_et):
                continue

        rec_dt = _parse_recorded_at_et_ledger(r["recorded_at"])
        if rec_dt is None:
            continue
        if rec_dt > now_et:
            continue

        key = (int(gpk), str(mt), st)
        prev = latest.get(key)
        if prev is None or rec_dt > prev[0]:
            latest[key] = (rec_dt, r)
    return latest


def _insert_bet_ledger_from_latest(
    conn: sqlite3.Connection,
    game_date: str,
    latest: dict[tuple[int, str, str], tuple[datetime.datetime, sqlite3.Row]],
) -> int:
    inserted = 0
    import json

    def _store_snapshot_if_needed(*, game_pk: int, market_type_raw: str, placed_at: str) -> None:
        """
        Store snapshot for staked bets only. Snapshot is computed once and then reused by PRIOR.
        """
        try:
            ensure_bet_snapshots(conn)
        except Exception:
            return

        mt_u = (market_type_raw or "").strip().lower()
        market_type = "ML" if mt_u == "moneyline" else ("TOTAL" if mt_u == "total" else ("RL" if mt_u in ("spread", "runline") else mt_u.upper()))
        if market_type not in ("ML", "TOTAL", "RL"):
            return

        try:
            ex = conn.execute(
                """
                SELECT bet, odds_taken, signals_used
                FROM bet_snapshots
                WHERE game_date=? AND game_pk=? AND market_type=?
                LIMIT 1
                """,
                (game_date, int(game_pk), market_type),
            ).fetchone()
            if ex:
                bet_existing = (ex[0] or "").strip()
                odds_existing = ex[1]
                sig_existing = (ex[2] or "").strip()
                if bet_existing and (odds_existing is not None) and sig_existing and sig_existing != "[]":
                    return
        except Exception:
            pass

        try:
            # Compute snapshot from data available at placement time.
            as_of = _now_et()
            games = load_games(conn, game_date, verbose=False, as_of_dt=as_of)
            gmap = {int(g["game_pk"]): g for g in games if g.get("game_pk") is not None}
            game = gmap.get(int(game_pk))
            if not game:
                return
            team_ids = list({game["home_team_id"], game["away_team_id"]})
            streaks = load_streaks(conn, game_date, team_ids, verbose=False)
            starters = load_starters(conn, game_date, verbose=False)
            sigs = evaluate_signals(conn, game, streaks, "primary", starters)
            sg = sigs.get("_scored_game")
            if sg is None:
                return
            me = getattr(sg, "market_evals", {}) or {}
            mm = me.get(market_type) or {}
            if not mm.get("evaluated"):
                return

            bet_side = str(mm.get("best_side") or "")
            bet = str(mm.get("bet") or "")
            odds_taken = mm.get("odds")
            score = int(mm.get("score") or 0)
            model_p = mm.get("model_p")
            implied_p = mm.get("implied_p")
            edge = mm.get("edge")
            # Use the *actual* evaluated signals used in scoring for this market.
            # Do NOT apply any post-threshold filtering.
            signals_used = list(mm.get("signals") or [])

            save_bet_snapshot(
                conn,
                game_date,
                int(game_pk),
                market_type,
                bet,
                int(odds_taken) if odds_taken is not None else None,
                mm,
                signals_used,
                placed_at=str(placed_at or ""),
            )
        except Exception:
            return

    # One opinion = one bet: if (game_pk, market_type) already has any staked bet, skip.
    # Also: totals are disabled temporarily (skip market_type == 'total').
    already_staked: set[tuple[int, str]] = set()
    try:
        rows = conn.execute(
            "SELECT game_pk, market_type FROM bet_ledger WHERE game_date=? AND stake_units>0",
            (game_date,),
        ).fetchall()
        for rr in rows:
            try:
                already_staked.add((int(rr[0]), str(rr[1] or "")))
            except Exception:
                continue
    except Exception:
        # Best-effort; fall back to unique index behavior only.
        pass
    for (_gpk, _mt, _st), (_dt, r) in latest.items():
        sig_type = (r["signal_type"] or "").strip()
        if sig_type not in ("top", "next", "avoid"):
            continue
        mt = str(r["market_type"] or "")
        if mt == "total":
            continue
        stake = 0.0 if sig_type == "avoid" else 1.0
        odds_val = None if sig_type == "avoid" else r["odds"]

        # If the staked ledger row already exists, still backfill the snapshot on reruns.
        if stake > 0 and (int(r["game_pk"]), mt) in already_staked:
            _store_snapshot_if_needed(
                game_pk=int(r["game_pk"]),
                market_type_raw=mt,
                placed_at=str(r["recorded_at"] or ""),
            )
            continue
        try:
            cur = conn.execute(
                """
                INSERT OR IGNORE INTO bet_ledger
                    (game_date, game_pk, market_type, bet, odds_taken, stake_units,
                     signal_at_time, session, placed_at, result, pnl_units)
                VALUES (?,?,?,?,?,?,?,?,?,?,?)
                """,
                (
                    game_date,
                    int(r["game_pk"]),
                    mt,
                    r["bet"],
                    odds_val,
                    stake,
                    sig_type,
                    r["session"],
                    r["recorded_at"],
                    None,
                    None,
                ),
            )
            if getattr(cur, "rowcount", 0) == 1:
                inserted += 1
            if stake > 0:
                # Snapshot is keyed by (date, game_pk, market). Even if the ledger row already
                # existed (rowcount==0), backfill the missing snapshot on re-runs.
                already_staked.add((int(r["game_pk"]), mt))
                _store_snapshot_if_needed(
                    game_pk=int(r["game_pk"]),
                    market_type_raw=mt,
                    placed_at=str(r["recorded_at"] or ""),
                )
        except sqlite3.OperationalError:
            continue
    return inserted


def generate_bets_from_signal_state(conn: sqlite3.Connection, game_date: str,
                                    now: datetime.datetime | None = None) -> int:
    """
    Create bet_ledger rows from signal_state in a rolling pregame window.

    A row is created when:
        game_start_utc - 30 minutes <= current_time < game_start_utc

    Selection:
      - Latest signal_state row per (game_pk, market_type, signal_type).
      - signal_type in ('top','next','avoid'). Avoid rows use stake_units=0.

    Idempotency:
      - UNIQUE (game_pk, market_type, signal_at_time) + INSERT OR IGNORE.

    Returns number of rows inserted.
    """
    if conn is None:
        return 0
    if not PERSIST_WRITES:
        return 0
    if now is None:
        now = _now_et()
    try:
        now_et = now if getattr(now, "tzinfo", None) else now.replace(tzinfo=_ET)
        now_et = now_et.astimezone(_ET)
    except Exception:
        now_et = _now_et()

    try:
        ensure_bet_ledger(conn)
    except Exception:
        return 0

    try:
        sig_rows = conn.execute(
            """
            SELECT
                game_pk, market_type, signal_type, bet, odds, session, recorded_at
            FROM signal_state
            WHERE game_date = ?
              AND game_pk IS NOT NULL
              AND market_type IS NOT NULL
              AND recorded_at IS NOT NULL
            """,
            (game_date,),
        ).fetchall()
    except sqlite3.OperationalError:
        return 0

    if not sig_rows:
        return 0

    game_pks = sorted({r["game_pk"] for r in sig_rows if r["game_pk"] is not None})
    if not game_pks:
        return 0
    placeholders = ",".join("?" * len(game_pks))
    try:
        g_rows = conn.execute(
            f"SELECT game_pk, game_start_utc FROM games WHERE game_pk IN ({placeholders})",
            tuple(game_pks),
        ).fetchall()
    except sqlite3.OperationalError:
        return 0

    start_by_pk: dict[int, datetime.datetime] = {}
    for r in g_rows:
        raw = r["game_start_utc"] or ""
        if "T" not in raw:
            continue
        try:
            utc_dt = datetime.datetime.fromisoformat(raw.rstrip("Z")).replace(
                tzinfo=datetime.timezone.utc
            )
            start_by_pk[int(r["game_pk"])] = utc_dt.astimezone(_ET)
        except Exception:
            continue

    latest = _ledger_latest_from_signal_rows(
        sig_rows, start_by_pk, now_et, pregame_window_only=True,
    )
    if not latest:
        return 0

    inserted = _insert_bet_ledger_from_latest(conn, game_date, latest)
    if inserted:
        try:
            conn.commit()
        except Exception:
            pass
    return inserted


def backfill_bet_ledger_from_signal_state(conn: sqlite3.Connection, game_date: str,
                                          now: datetime.datetime | None = None) -> int:
    """
    Materialize all top/next/avoid rows from signal_state for game_date into bet_ledger
    (no pregame window). Idempotent. Used by prior-day report so ledger matches
    archived signals including avoids.
    """
    if conn is None:
        return 0
    if now is None:
        now = _now_et()
    try:
        now_et = now if getattr(now, "tzinfo", None) else now.replace(tzinfo=_ET)
        now_et = now_et.astimezone(_ET)
    except Exception:
        now_et = _now_et()

    try:
        ensure_bet_ledger(conn)
    except Exception:
        return 0

    try:
        sig_rows = conn.execute(
            """
            SELECT
                game_pk, market_type, signal_type, bet, odds, session, recorded_at
            FROM signal_state
            WHERE game_date = ?
              AND game_pk IS NOT NULL
              AND market_type IS NOT NULL
              AND recorded_at IS NOT NULL
            """,
            (game_date,),
        ).fetchall()
    except sqlite3.OperationalError:
        return 0

    if not sig_rows:
        return 0

    game_pks = sorted({r["game_pk"] for r in sig_rows if r["game_pk"] is not None})
    placeholders = ",".join("?" * len(game_pks))
    try:
        g_rows = conn.execute(
            f"SELECT game_pk, game_start_utc FROM games WHERE game_pk IN ({placeholders})",
            tuple(game_pks),
        ).fetchall()
    except sqlite3.OperationalError:
        return 0

    start_by_pk: dict[int, datetime.datetime] = {}
    for r in g_rows:
        raw = r["game_start_utc"] or ""
        if "T" not in raw:
            continue
        try:
            utc_dt = datetime.datetime.fromisoformat(raw.rstrip("Z")).replace(
                tzinfo=datetime.timezone.utc
            )
            start_by_pk[int(r["game_pk"])] = utc_dt.astimezone(_ET)
        except Exception:
            continue

    latest = _ledger_latest_from_signal_rows(
        sig_rows, start_by_pk, now_et, pregame_window_only=False,
    )
    if not latest:
        return 0

    inserted = _insert_bet_ledger_from_latest(conn, game_date, latest)
    if inserted:
        try:
            conn.commit()
        except Exception:
            pass
    return inserted


def strip_avoid_bet_label(bet_text: str) -> str:
    """Strip leading 'Avoid:' / 'Avoid ' so ML/total parsers can grade counterfactuals."""
    s = (bet_text or "").strip()
    su = s.upper()
    if su.startswith("AVOID "):
        return s[6:].strip()
    if su.startswith("AVOID:"):
        return s[6:].strip()
    return s


def avoid_entry_bet_fields(entry: dict) -> tuple[str, str, float | None]:
    """Return (market, bet_label, total_line) for an AVOID entry (brief / prior display)."""
    g = (entry.get("game") or {})
    sigs = (entry.get("sigs") or {})
    reason = (sigs.get("avoid_reason") or "").strip()
    reason_u = reason.upper()
    home = (g.get("home_abbr") or "").strip()
    away = (g.get("away_abbr") or "").strip()
    total_line = g.get("total_line")
    venue = (g.get("venue_name") or "").strip()

    market = "OTHER"
    bet_text = None

    # Venue / environment — not a single ML or O/U ticket (wind-edge class warning).
    if (
        "WIND SIGNALS SUPPRESSED" in reason_u
        or ("SUPPRESSED" in reason_u and "WEATHER" in reason_u)
        or "NO WEATHER EDGE" in reason_u
    ):
        market = "ENV"
        where = f" ({venue})" if venue else ""
        bet_text = (
            f"Weather-driven model plays only (wind/MV-B/H3b-style edges) — "
            f"no single numbered bet{where}; venue treated as non-actionable for weather signals"
        )
    if bet_text is None:
        if "AVOID HOME ML" in reason_u and home:
            market = "ML"
            bet_text = f"Avoid {home} ML"
        elif "AVOID AWAY ML" in reason_u and away:
            market = "ML"
            bet_text = f"Avoid {away} ML"
        elif ("DO NOT BET OVER" in reason_u or ("OVER" in reason_u and "DO NOT BET" in reason_u)) and total_line is not None:
            market = "TOTAL"
            bet_text = f"Avoid OVER {total_line}"
        elif ("DO NOT BET UNDER" in reason_u or ("UNDER" in reason_u and "DO NOT BET" in reason_u)) and total_line is not None:
            market = "TOTAL"
            bet_text = f"Avoid UNDER {total_line}"
        elif "ML" in reason_u:
            market = "ML"
            bet_text = f"Avoid {home} ML" if home else "Avoid ML"
        elif "TOTAL" in reason_u or "O/U" in reason_u or "OVER" in reason_u or "UNDER" in reason_u:
            market = "TOTAL"
            bet_text = f"Avoid total" if total_line is None else f"Avoid total ({total_line})"

    if not bet_text:
        bet_text = ("Avoid: " + reason) if reason else "Avoid"

    return market, bet_text, (float(total_line) if total_line is not None else None)


def prior_avoid_outcome_lines(entry: dict) -> tuple[str, str, str] | None:
    """
    Human lines for prior report: (bet_to_skip, counterfactual, verdict).
    None if game not final or not parseable.
    """
    g = entry.get("game") or {}
    hs, as_ = g.get("home_score"), g.get("away_score")
    if hs is None or as_ is None:
        return None
    market, bet_label, _tl = avoid_entry_bet_fields(entry)
    if market == "ENV":
        return (
            f"  BET TO SKIP: {bet_label}",
            "  Counterfactual: N/A — this flags a class of weather-driven model plays, "
            "not one specific ML or O/U ticket at the closing line.",
            "  Verdict: N/A — informational; see model detail for venue / weather context.",
        )
    equiv = strip_avoid_bet_label(bet_label)
    home_abbr = (g.get("home_abbr") or "").strip().upper()
    away_abbr = (g.get("away_abbr") or "").strip().upper()
    runs = hs + as_

    def _parse_ml_team_local(bt: str) -> str | None:
        s = (bt or "").strip().upper()
        if not s:
            return None
        return s.split()[0]

    def _parse_total_local(bt: str) -> tuple[str | None, float | None]:
        s = (bt or "").strip().upper()
        if not s:
            return None, None
        side = "over" if s.startswith("OVER") else ("under" if s.startswith("UNDER") else None)
        if side is None:
            return None, None
        try:
            return side, float(s.split()[1])
        except Exception:
            return side, None

    hypo: str | None = None
    if market == "ML":
        team = _parse_ml_team_local(equiv)
        if not team:
            return None
        if team == home_abbr:
            hypo = "win" if hs > as_ else "loss"
        elif team == away_abbr:
            hypo = "win" if as_ > hs else "loss"
        else:
            return None
    elif market == "TOTAL":
        side, line = _parse_total_local(equiv)
        if side is None or line is None:
            return None
        if runs == line:
            hypo = "push"
        else:
            won = (side == "over" and runs > line) or (side == "under" and runs < line)
            hypo = "win" if won else "loss"
    else:
        bet_line = f"  BET TO SKIP: {bet_label}"
        tail = (
            "  (Counterfactual not parsed for this flag type — see model note below.)"
        )
        return (bet_line, tail, "  Verdict: —")

    bet_line = f"  BET TO SKIP: {bet_label}"
    if hypo == "push":
        cf = "  If you had taken it: PUSH vs closing total (no win/loss edge)."
        ver = "  Verdict on the call: — (push — neither good nor bad in 1u terms)."
    elif hypo == "win":
        cf = "  If you had taken it: would have WON (vs closing line / side above)."
        ver = "  Verdict on the call: ✗ Poor avoid — you skipped a winning side."
    else:
        cf = "  If you had taken it: would have LOST."
        ver = "  Verdict on the call: ✓ Good avoid — you dodged a losing bet."

    return (bet_line, cf, ver)


def grade_bet_ledger(conn: sqlite3.Connection, game_date: str | None = None) -> int:
    """
    Grade bets in bet_ledger for games that are Final.

    - Joins bet_ledger -> games on game_pk
    - Only processes games where games.status = 'Final'
    - Top/next rows: result in ('win','loss','push'), pnl_units at stake 1
    - Avoid rows (signal_at_time='avoid'): result good_avoid / bad_avoid / push_avoid —
      counterfactual for the skipped bet; pnl_units 0 (informational only)

    Market rules:
      moneyline:
        bet team wins -> win else loss
      total:
        compare total line vs actual runs
      run line:
        compare spread vs final score
    """
    if conn is None:
        return 0

    def _pnl_units_from_odds(odds: int | None, won: bool, push: bool) -> float:
        if push:
            return 0.0
        if not won:
            return -1.0
        if odds is None:
            # Missing odds -> assume even money
            return 1.0
        return (odds / 100.0) if odds > 0 else (100.0 / abs(odds))

    def _parse_total_bet(bet_text: str) -> tuple[str | None, float | None]:
        s = (bet_text or "").strip().upper()
        if not s:
            return None, None
        side = "over" if s.startswith("OVER") else ("under" if s.startswith("UNDER") else None)
        if side is None:
            return None, None
        # e.g. "OVER 8.5"
        try:
            num = float(s.split()[1])
            return side, num
        except Exception:
            return side, None

    def _parse_runline_bet(bet_text: str) -> tuple[str | None, float | None]:
        # e.g. "NYY -1.5" or "STL +1.5"
        s = (bet_text or "").strip().upper()
        if not s:
            return None, None
        parts = s.split()
        if len(parts) < 2:
            return None, None
        team = parts[0]
        try:
            line = float(parts[1])
        except Exception:
            return team, None
        return team, line

    def _parse_ml_team(bet_text: str) -> str | None:
        # e.g. "CHC ML" or "BAL ML"
        s = (bet_text or "").strip().upper()
        if not s:
            return None
        parts = s.split()
        if not parts:
            return None
        return parts[0]

    where_date = ""
    params: tuple = ()
    if game_date:
        where_date = "AND bl.game_date = ?"
        params = (game_date,)

    # Only grade ungraded bets (result is NULL or empty) for Final games.
    rows = conn.execute(
        f"""
        SELECT
            bl.id,
            bl.game_pk,
            g.game_date_et AS game_date_et,
            bl.market_type,
            bl.bet,
            bl.odds_taken,
            bl.signal_at_time,
            g.home_team_id,
            g.away_team_id,
            g.home_score,
            g.away_score,
            th.abbreviation AS home_abbr,
            ta.abbreviation AS away_abbr
        FROM bet_ledger bl
        JOIN games g ON g.game_pk = bl.game_pk
        JOIN teams th ON th.team_id = g.home_team_id
        JOIN teams ta ON ta.team_id = g.away_team_id
        WHERE g.status = 'Final'
          AND (bl.result IS NULL OR TRIM(bl.result) = '')
          {where_date}
        """,
        params,
    ).fetchall()

    if not rows:
        return 0

    # ── Calibration log (CSV append-only) ────────────────────────────────
    # For each graded staked bet, append: date, game_pk, bet_type, score, model_p, implied_p, edge, result.
    # Uses score_game via evaluate_signals() on the same slate date (Final games included).
    import csv
    from pathlib import Path

    from batch.pipeline.edge_utils import american_to_implied_prob, score_to_model_prob

    cal_path = Path(_REPO_ROOT) / "data" / "calibration_log.csv"
    cal_path.parent.mkdir(parents=True, exist_ok=True)
    _wrote_header = cal_path.exists() is False

    # Cache per game_pk so we only score each game once during grading.
    _scored_cache: dict[int, object] = {}

    def _get_scored_for_game_pk(gpk: int, gd: str) -> object | None:
        if gpk in _scored_cache:
            return _scored_cache[gpk]
        try:
            as_of = _now_et()  # non-None → load_games includes Final rows
            games = load_games(conn, gd, verbose=False, as_of_dt=as_of)
            gmap = {int(g["game_pk"]): g for g in games if g.get("game_pk") is not None}
            game = gmap.get(int(gpk))
            if not game:
                _scored_cache[gpk] = None
                return None
            team_ids = list({game["home_team_id"], game["away_team_id"]})
            streaks = load_streaks(conn, gd, team_ids, verbose=False)
            starters = load_starters(conn, gd, verbose=False)
            sigs = evaluate_signals(conn, game, streaks, "primary", starters)
            sg = sigs.get("_scored_game")
            _scored_cache[gpk] = sg
            return sg
        except Exception:
            _scored_cache[gpk] = None
            return None

    def _append_cal_row(*, gd: str, gpk: int, bet_type: str, market_type: str, odds_taken: int | None, result_str: str) -> None:
        nonlocal _wrote_header
        sg = _get_scored_for_game_pk(gpk, gd)
        if sg is None:
            return
        # Market-specific calibration: do NOT mix ML/TOTAL/RL in one dataset.
        try:
            me = getattr(sg, "market_evals", {}) or {}
            mt = str(market_type or "").strip().upper()
            mm = me.get(mt) or {}
            score = int(mm.get("score") or 0)
            model_p = float(mm.get("model_p") or score_to_model_prob(score))
            implied_p = float(mm.get("implied_p")) if mm.get("implied_p") is not None else american_to_implied_prob(int(odds_taken) if odds_taken is not None else None)
            edge = float(mm.get("edge")) if mm.get("edge") is not None else ((model_p - implied_p) if implied_p is not None else None)
        except Exception:
            score = 0
            model_p = float(score_to_model_prob(score))
            implied_p = american_to_implied_prob(int(odds_taken) if odds_taken is not None else None)
            edge = (model_p - implied_p) if implied_p is not None else None
        res_u = (result_str or "").strip().lower()
        if res_u == "win":
            res_val = 1
        elif res_u == "loss":
            res_val = 0
        else:
            # ignore pushes/avoids/no-result for calibration
            return
        with cal_path.open("a", newline="", encoding="utf-8") as fh:
            w = csv.writer(fh)
            if _wrote_header:
                w.writerow(["date", "game_pk", "market_type", "bet_type", "score", "model_p", "implied_p", "edge", "result"])
                _wrote_header = False
            w.writerow(
                [
                    gd,
                    int(gpk),
                    str(market_type or "").strip().upper(),
                    bet_type,
                    int(score),
                    f"{model_p:.3f}",
                    (f"{implied_p:.3f}" if implied_p is not None else "NA"),
                    (f"{edge:.3f}" if edge is not None else "NA"),
                    int(res_val),
                ]
            )

    updated = 0
    for r in rows:
        market = (r["market_type"] or "").strip().lower()
        bet_text = r["bet"] or ""
        hs = r["home_score"]
        as_ = r["away_score"]
        if hs is None or as_ is None:
            continue

        res = None
        pnl = None
        sig_at = (r["signal_at_time"] or "").strip().lower()

        if sig_at == "avoid":
            equiv = strip_avoid_bet_label(bet_text)
            home_abbr = (r["home_abbr"] or "").upper()
            away_abbr = (r["away_abbr"] or "").upper()
            runs = hs + as_
            hypo: str | None = None

            if market == "moneyline":
                team = _parse_ml_team(equiv)
                if not team:
                    continue
                if team == home_abbr:
                    won = hs > as_
                elif team == away_abbr:
                    won = as_ > hs
                else:
                    continue
                hypo = "win" if won else "loss"
            elif market == "total":
                side, line = _parse_total_bet(equiv)
                if side is None or line is None:
                    continue
                if runs == line:
                    hypo = "push"
                else:
                    won = (side == "over" and runs > line) or (side == "under" and runs < line)
                    hypo = "win" if won else "loss"
            elif market in ("spread", "runline"):
                team, line = _parse_runline_bet(equiv)
                if team is None or line is None:
                    continue
                if team == home_abbr:
                    adj = hs + line
                    opp = as_
                elif team == away_abbr:
                    adj = as_ + line
                    opp = hs
                else:
                    continue
                if adj == opp:
                    hypo = "push"
                else:
                    won = adj > opp
                    hypo = "win" if won else "loss"
            else:
                continue

            if hypo == "push":
                res = "push_avoid"
            elif hypo == "win":
                res = "bad_avoid"
            else:
                res = "good_avoid"
            pnl = 0.0

            conn.execute(
                "UPDATE bet_ledger SET result = ?, pnl_units = ? WHERE id = ?",
                (res, float(round(pnl, 4)), r["id"]),
            )
            updated += 1
            continue

        if market == "moneyline":
            team = _parse_ml_team(bet_text)
            if not team:
                continue
            home_abbr = (r["home_abbr"] or "").upper()
            away_abbr = (r["away_abbr"] or "").upper()
            if team == home_abbr:
                won = hs > as_
            elif team == away_abbr:
                won = as_ > hs
            else:
                # Unknown team token
                continue
            res = "win" if won else "loss"
            pnl = _pnl_units_from_odds(r["odds_taken"], won=won, push=False)

        elif market == "total":
            side, line = _parse_total_bet(bet_text)
            if side is None or line is None:
                continue
            runs = hs + as_
            if runs == line:
                res = "push"
                pnl = _pnl_units_from_odds(r["odds_taken"], won=False, push=True)
            else:
                won = (side == "over" and runs > line) or (side == "under" and runs < line)
                res = "win" if won else "loss"
                pnl = _pnl_units_from_odds(r["odds_taken"], won=won, push=False)

        elif market in ("spread", "runline"):
            team, line = _parse_runline_bet(bet_text)
            if team is None or line is None:
                continue
            home_abbr = (r["home_abbr"] or "").upper()
            away_abbr = (r["away_abbr"] or "").upper()
            if team == home_abbr:
                adj = hs + line
                opp = as_
            elif team == away_abbr:
                adj = as_ + line
                opp = hs
            else:
                continue
            if adj == opp:
                res = "push"
                pnl = _pnl_units_from_odds(r["odds_taken"], won=False, push=True)
            else:
                won = adj > opp
                res = "win" if won else "loss"
                pnl = _pnl_units_from_odds(r["odds_taken"], won=won, push=False)
        else:
            continue

        if res is None or pnl is None:
            continue

        conn.execute(
            "UPDATE bet_ledger SET result = ?, pnl_units = ? WHERE id = ?",
            (res, float(round(pnl, 4)), r["id"]),
        )
        updated += 1
        try:
            # Append calibration row for staked bets only.
            if (
                float(r.get("stake_units") or 0.0) > 0
                and (r["odds_taken"] is not None)
                and (market in ("moneyline", "spread", "runline", "total"))
            ):
                mt_u = (r.get("market_type") or "").strip().lower()
                mlabel = "ML" if mt_u == "moneyline" else ("TOTAL" if mt_u == "total" else ("RL" if mt_u in ("spread", "runline") else str(r.get("market_type") or "").upper()))
                _append_cal_row(
                    gd=str(r.get("game_date_et") or game_date or ""),
                    gpk=int(r["game_pk"]),
                    bet_type=str(r["market_type"] or ""),
                    market_type=mlabel,
                    odds_taken=int(r["odds_taken"]) if r["odds_taken"] is not None else None,
                    result_str=str(res),
                )
        except Exception:
            pass

    if updated:
        conn.commit()
    return updated


def save_brief_picks(
    conn: sqlite3.Connection,
    game_date: str,
    session: str,
    pick_entries: list | None,
    *,
    avoid_entries: list | None = None,
    now: datetime.datetime | None = None,
    model_version: str = "legacy",
) -> None:
    """Record picks shown in this brief for prior-report grading.
    pick_entries: sorted list of entry dicts from the brief builder.
    Records top pick (rank 1) and additional picks (ranks 2-6).
    Also records AVOID flags for research as pick_rank=0 with signal='AVOID'.
    AVOID rows sync into bet_ledger via signal_state + backfill (stake 0).
    Idempotent — INSERT OR IGNORE on (game_date, session, game_pk, signal).

    Skips insert when the game has already started (vs ``now``) or when the same
    game_date + game_pk + signal + market was stored in an earlier session.
    """
    if now is None:
        now = _now_et()
    now_et = now.strftime("%Y-%m-%d %H:%M ET")

    ensure_brief_picks(conn)

    wrote_any = False

    if pick_entries:
        for rank, entry in enumerate(pick_entries[:6], start=1):
            g = entry["game"]
            # Use highest-priority pick for this game
            p = sorted(entry["sigs"]["picks"], key=lambda x: x["priority"])[0]
            signal = ", ".join(entry["sigs"]["signals"])
            if _game_already_started_for_brief_picks(g, now):
                continue
            if _brief_pick_exists_cross_session(
                conn, game_date, int(g["game_pk"]), signal, p["market"],
            ):
                continue
            odds_raw = _parse_odds(p.get("odds", ""))
            total_line_val = None
            if p.get("market") == "TOTAL":
                total_line_val = entry["game"].get("total_line")
            try:
                conn.execute(
                    """
                    INSERT OR IGNORE INTO brief_picks
                        (game_date, session, game_pk, pick_rank, signal,
                         bet, market, odds, total_line, recorded_at, model_version)
                    VALUES (?,?,?,?,?,?,?,?,?,?,?)
                    """,
                    (
                        game_date,
                        session,
                        g["game_pk"],
                        rank,
                        signal,
                        p["bet"],
                        p["market"],
                        odds_raw,
                        total_line_val,
                        now_et,
                        model_version,
                    ),
                )
                wrote_any = True
            except Exception:
                pass

    if avoid_entries:
        for entry in avoid_entries:
            g = (entry.get("game") or {})
            gpk = g.get("game_pk")
            if gpk is None:
                continue
            if _game_already_started_for_brief_picks(g, now):
                continue
            market, bet_text, total_line_val = avoid_entry_bet_fields(entry)
            if _brief_pick_exists_cross_session(
                conn, game_date, int(gpk), "AVOID", market,
            ):
                continue
            try:
                conn.execute(
                    """
                    INSERT OR IGNORE INTO brief_picks
                        (game_date, session, game_pk, pick_rank, signal,
                         bet, market, odds, total_line, recorded_at, model_version)
                    VALUES (?,?,?,?,?,?,?,?,?,?,?)
                    """,
                    (
                        game_date,
                        session,
                        int(gpk),
                        0,
                        "AVOID",
                        bet_text,
                        market,
                        None,
                        total_line_val,
                        now_et,
                        model_version,
                    ),
                )
                wrote_any = True
            except Exception:
                pass

    if wrote_any:
        conn.commit()


def load_brief_picks(
    conn: sqlite3.Connection,
    game_date: str,
    session: str = "primary",
    *,
    include_avoids: bool = False,
) -> list:
    """
    Load confirmed picks that were shown in yesterday's primary brief.
    Returns list of dicts ordered by pick_rank.

    By default excludes AVOID research rows (pick_rank=0, signal='AVOID').
    """
    try:
        avoid_clause = "" if include_avoids else " AND pick_rank > 0 "
        rows = conn.execute("""
            SELECT game_pk, pick_rank, signal, bet, market, odds
            FROM   brief_picks
            WHERE  game_date = ? AND session = ?
              {avoid_clause}
            ORDER  BY pick_rank
        """.format(avoid_clause=avoid_clause), (game_date, session)).fetchall()
        return [dict(r) for r in rows]
    except Exception:
        return []


def load_todays_prior_sessions(conn: sqlite3.Connection, game_date: str,
                               current_session: str) -> list:
    """
    Load all picks recorded in earlier sessions TODAY, ordered by session then rank.
    Used to build the intra-day Signal Tracker.

    Session order: early < afternoon < primary  (closing is never a source for tracker)
    Returns list of dicts, each with:
        session, pick_rank, game_pk, signal, bet, market, odds
    """
    SESSION_ORDER = ["early", "afternoon", "primary", "closing", "late"]
    try:
        current_idx = SESSION_ORDER.index(current_session)
    except ValueError:
        current_idx = len(SESSION_ORDER)

    prior_sessions = SESSION_ORDER[:current_idx]
    if not prior_sessions:
        return []

    placeholders = ",".join("?" * len(prior_sessions))
    try:
        rows = conn.execute(f"""
            SELECT session, game_pk, pick_rank, signal, bet, market, odds,
                   total_line, recorded_at
            FROM   brief_picks
            WHERE  game_date  = ?
              AND  session    IN ({placeholders})
              AND  pick_rank  > 0
            ORDER  BY
                CASE session
                    WHEN 'early'     THEN 1
                    WHEN 'afternoon' THEN 2
                    WHEN 'primary'   THEN 3
                    ELSE 4
                END,
                pick_rank
        """, (game_date, *prior_sessions)).fetchall()
        return [dict(r) for r in rows]
    except Exception:
        return []


# Session display labels for the tracker
_SESSION_LABELS = {
    "early":     "EARLY",
    "afternoon": "AFTERNOON",
    "primary":   "PRIMARY",
    "closing":   "CLOSING",
    "late":      "LATE GAMES",
}

# Signal tracker rules
SIGNAL_MIN_SCORE = 15



def movement_alert(conn: sqlite3.Connection, game_date: str,
                   current_session: str, game_pk: int,
                   current_total: float, current_home_ml: int) -> str:
    """
    Compare current total/ML against the earliest prior session pick for this
    game today. Returns a formatted alert string if movement exceeds thresholds,
    or empty string if no prior pick or no significant movement.

    Thresholds (matching ops guide):
      Total  : ≥ 0.5 points in either direction
      ML     : ≥ 10 cents implied probability (not raw points — catches -130→-160 etc.)

    Arrow convention:
      ⬇ Total dropped   = counter-signal for OVER picks (sharps see lower scoring)
      ⬆ Total rose      = confirms OVER pick
      ML move shown as implied probability shift
    """
    SESSION_ORDER = ["early", "afternoon", "primary", "closing", "late"]
    try:
        current_idx = SESSION_ORDER.index(current_session)
    except ValueError:
        return ""
    if current_idx == 0:
        return ""  # no prior session today

    prior_sessions = SESSION_ORDER[:current_idx]
    placeholders   = ",".join("?" * len(prior_sessions))

    try:
        row = conn.execute(f"""
            SELECT market, total_line, odds
            FROM   brief_picks
            WHERE  game_date = ? AND game_pk = ?
              AND  session IN ({placeholders})
              AND  pick_rank > 0
            ORDER  BY
                CASE session
                    WHEN 'early'     THEN 1
                    WHEN 'afternoon' THEN 2
                    WHEN 'primary'   THEN 3
                    ELSE 4
                END
            LIMIT 1
        """, (game_date, game_pk, *prior_sessions)).fetchone()
    except Exception:
        return ""

    if not row:
        return ""

    alerts = []

    # ── Total movement ────────────────────────────────────────────────────
    prior_total = row["total_line"] if isinstance(row, dict) else row[1]
    if prior_total is not None and current_total is not None:
        try:
            delta = round(float(current_total) - float(prior_total), 2)
            if abs(delta) >= 0.5:
                direction = "⬇" if delta < 0 else "⬆"
                counter   = "  ← COUNTER-SIGNAL for OVER" if delta < 0 else "  ← confirms OVER"
                alerts.append(
                    f"⚡ MOVEMENT ALERT: Total {prior_total} → {current_total} "
                    f"({delta:+.1f}) {direction}{counter}"
                )
        except (TypeError, ValueError):
            pass

    # ── ML movement (implied probability shift) ───────────────────────────
    prior_ml_raw = row["odds"] if isinstance(row, dict) else row[2]
    prior_market  = row["market"] if isinstance(row, dict) else row[0]
    if prior_market == "ML" and prior_ml_raw is not None and current_home_ml is not None:
        try:
            def to_imp(ml):
                ml = int(ml)
                return abs(ml)/(abs(ml)+100) if ml < 0 else 100/(ml+100)
            prior_imp   = to_imp(prior_ml_raw)
            current_imp = to_imp(current_home_ml)
            shift_cents = round((current_imp - prior_imp) * 100, 1)
            if abs(shift_cents) >= 10:
                direction = "HOME more favoured" if shift_cents > 0 else "HOME less favoured"
                alerts.append(
                    f"⚡ MOVEMENT ALERT: ML shifted {shift_cents:+.0f}¢ implied "
                    f"({direction})"
                )
        except (TypeError, ValueError):
            pass

    return "\n".join(f"  {a}" for a in alerts)


def build_signal_tracker_block(prior_picks: list, current_games: list,
                                streaks: dict, current_session: str) -> str:
    """
    Build the intra-day Signal Tracker section for action briefs.

    Compares picks from earlier today's sessions against the current slate:
      · Still firing with same bet → ACTIVE (unchanged)
      · Still in picks but odds moved → ACTIVE (line moved N pts)
      · Dropped from picks entirely → DROPPED (with reason)
      · Was Top Pick, now lower rank → DEMOTED
      · New pick not in earlier sessions → shown in main picks section, not here

    Returns formatted text block (empty string if no prior sessions had picks).
    """
    if not prior_picks:
        return ""

    lines = []
    lines.append(section("📡  SIGNAL TRACKER  —  Today's Pick Status vs Earlier Sessions"))
    lines.append(
        "\n  Tracks how today's signals have evolved across brief sessions.\n"
        "  Early price = odds when signal first fired. Current = now.\n"
        "  If you got in early: your position is noted. Decision is yours.\n"
    )

    # Build a lookup: game_pk → current recomputed status (no carry-forward of prior signals).
    # We recompute for *every* evaluated game each brief run; tracker only *displays* games
    # that meet SIGNAL_MIN_SCORE or are eligible bets by edge.
    from batch.pipeline.edge_utils import EDGE_MIN

    wind_ids = {"MV-F", "MV-B", "H3b"}
    streak_ids = {"S1+H2", "S1H2", "S1"}
    matchup_ids = {"LHP_FADE", "LHP_FADE_RL", "NF4"}

    def _grouped_categories_from_scored(scored: object) -> list[str]:
        cats: set[str] = set()
        sigs = list(getattr(scored, "signals_fired", []) or [])
        for s in sigs:
            if not bool(getattr(s, "fires", False)):
                continue
            sid = str(getattr(s, "signal_id", "") or "")
            if sid in wind_ids:
                cats.add("wind")
            elif sid in streak_ids:
                cats.add("streak")
            elif sid in matchup_ids:
                cats.add("matchup")
        return sorted(cats)

    current_lookup: dict[int, dict] = {}
    for entry in (current_games or []):
        g = entry.get("game") or {}
        pk = g.get("game_pk")
        if pk is None:
            continue
        sg = (entry.get("sigs") or {}).get("_scored_game")
        if sg is None:
            continue
        try:
            score = int(getattr(sg, "best_aggregate_score", 0) or 0)
        except Exception:
            score = 0
        try:
            edge = float(getattr(sg, "edge", 0.0) or 0.0)
        except Exception:
            edge = 0.0

        if edge >= EDGE_MIN:
            status = "ACTIVE BET SIGNAL"
        elif score >= SIGNAL_MIN_SCORE:
            status = "WEAK SIGNAL (no edge)"
        else:
            status = "NO SIGNAL"

        current_lookup[int(pk)] = {
            "score": score,
            "edge": edge,
            "status": status,
            "cats": _grouped_categories_from_scored(sg),
        }

    # Sanity check: if too many games are showing signals, flag as over-triggering.
    eligible = [
        v for v in current_lookup.values()
        if (v["edge"] >= EDGE_MIN) or (int(v["score"]) >= SIGNAL_MIN_SCORE)
    ]
    if current_games:
        try:
            frac = len(eligible) / float(len(current_games))
        except Exception:
            frac = 0.0
        if frac > 0.50:
            lines.append(color_text("\n  ⚠  Sanity check: signals on >50% of games — over-triggering\n", "yellow"))

    # Group prior picks by game_pk, keep the EARLIEST session's record per game
    seen_pks = {}
    for row in prior_picks:
        pk = row["game_pk"]
        if pk not in seen_pks:
            seen_pks[pk] = row   # first occurrence = earliest session

    if not seen_pks:
        lines.append("  No picks were recorded in prior sessions today.\n")
        return "\n".join(lines)

    for pk, prior in seen_pks.items():
        sess_label   = _SESSION_LABELS.get(prior["session"], prior["session"].upper())
        prior_rank   = prior["pick_rank"]
        prior_odds   = prior["odds"]
        prior_bet    = prior["bet"]
        rank_label   = "TOP PICK" if prior_rank == 1 else f"#{prior_rank}"
        prior_odds_str = fmt_odds(prior_odds) if prior_odds else "N/A"

        cur = current_lookup.get(int(pk))
        if cur is None:
            continue

        score = int(cur.get("score") or 0)
        edge = float(cur.get("edge") or 0.0)
        status = str(cur.get("status") or "NO SIGNAL")
        cats = list(cur.get("cats") or [])
        cats_txt = (" · ".join(cats)) if cats else "none"

        # Only list games that meet the threshold (or are bet-eligible).
        from batch.pipeline.edge_utils import EDGE_MIN as _EDGE_MIN
        if not (edge >= _EDGE_MIN or score >= SIGNAL_MIN_SCORE):
            continue

        lines.append(
            f"\n  {prior_bet}  —  was {rank_label} @ {sess_label} ({prior_odds_str})\n"
            f"      NOW: {status}   |   score={score} edge={edge:.3f}\n"
            f"      GROUPED SIGNALS: {cats_txt}\n"
        )

    lines.append("")
    return "\n".join(lines)


def compute_paper_picks(evaluated: list, game_date: str) -> list:
    """Convert graded pick entries into paper-trading rows ready for daily_pnl.
    evaluated: list of entry dicts from build_prior_day_report.
    Returns list of row dicts.
    """
    pick_entries = sorted(
        [e for e in evaluated if e["graded"]],
        key=lambda e: min(p["priority"] for p in e["graded"])
    )
    top_pick   = pick_entries[:1]
    next_picks = pick_entries[1:6]
    # rest_picks (7th+) get $0 — not recorded

    pnl_rows = []
    for tier_label, tier_entries, base_stake in [
        ("top",        top_pick,   PAPER_STAKE_TOP),
        ("additional", next_picks, PAPER_STAKE_ADD),
    ]:
        for entry in tier_entries:
            g = entry["game"]
            # Use the highest-priority pick for this game
            p = sorted(entry["graded"], key=lambda x: x["priority"])[0]
            signal = ", ".join(entry["sigs"]["signals"])

            # Late-season H3b half-stake detection
            try:
                month = int(game_date[5:7])
            except (ValueError, TypeError):
                month = 0
            is_h3b    = "H3b" in (entry["sigs"].get("signal_ids") or [])
            late_flag = 1 if (is_h3b and month in H3B_LATE_SEASON_MONTHS) else 0
            stake     = round(base_stake * (PAPER_LATE_FACTOR if late_flag else 1.0), 2)

            # Dollar P&L from units
            pnl_dollars = round(p["pnl"] * stake, 2)

            # Clean result string
            raw_result = p["result"]  # e.g. "✓ WIN", "✗ LOSS", "— PUSH"
            if "WIN"  in raw_result: result_clean = "WIN"
            elif "LOSS" in raw_result: result_clean = "LOSS"
            elif "PUSH" in raw_result: result_clean = "PUSH"
            else: result_clean = "NO RESULT"

            pnl_rows.append({
                "game_date":     game_date,
                "game_pk":       g["game_pk"],
                "signal":        signal,
                "pick_tier":     tier_label,
                "bet":           p["bet"],
                "market":        p["market"],
                "odds":          p.get("odds_raw") or _parse_odds(p.get("odds", "")),
                "stake_dollars": stake,
                "late_season":   late_flag,
                "result":        result_clean,
                "pnl_units":     round(p["pnl"], 4),
                "pnl_dollars":   pnl_dollars,
            })
    return pnl_rows


def _parse_odds(odds_str: str) -> int | None:
    """Parse a formatted odds string like '-130' or '+115' to int.
    Returns None if unparseable.
    """
    try:
        s = str(odds_str).replace("\u2212", "-").replace("\u202f", "").strip()
        return int(s)
    except (ValueError, TypeError):
        return None


def format_dollar_pnl_block(today_rows: list, season_pnl: dict,
                            game_date: str) -> str:
    """Format the dollar P&L section for the prior-day report."""
    lines = []
    lines.append(section("💵  PAPER ACCOUNT  —  Dollar Performance"))

    if not today_rows:
        lines.append("\n  No bets placed (no signals fired yesterday).\n")
    else:
        top_rows  = [r for r in today_rows if r["pick_tier"] == "top"]
        add_rows  = [r for r in today_rows if r["pick_tier"] == "additional"]
        today_pnl = sum(r["pnl_dollars"] for r in today_rows)
        today_staked = sum(r["stake_dollars"] for r in today_rows)

        # Stake summary line
        stake_parts = []
        for r in top_rows:
            flag = " ½-stake H3b late-season" if r["late_season"] else ""
            stake_parts.append(f"Top ${r['stake_dollars']:.0f}{flag}")
        if add_rows:
            stake_parts.append(f"{len(add_rows)} additional × ${add_rows[0]['stake_dollars']:.0f}")
        lines.append(f"\n  Bets placed:   {'  ·  '.join(stake_parts)}")
        lines.append(f"  Total staked:  ${today_staked:.2f}")

        # Per-bet detail
        lines.append("")
        for r in today_rows:
            tier  = "TOP" if r["pick_tier"] == "top" else "ADD"
            sign  = "+" if r["pnl_dollars"] >= 0 else ""
            flag  = " [½-stake]" if r["late_season"] else ""
            lines.append(
                f"  [{tier}] {r['bet']:<22} {r['result']:<8} "
                f"${r['stake_dollars']:.0f} → {sign}${r['pnl_dollars']:.2f}{flag}"
            )

        # Day total
        sign = "+" if today_pnl >= 0 else ""
        lines.append(f"\n  TODAY:   {sign}${today_pnl:.2f}")

    # Season to date
    stv = season_pnl
    sign = "+" if stv["total_dollars"] >= 0 else ""
    rec  = f"W:{stv['wins']}  L:{stv['losses']}"
    if stv["pushes"]:
        rec += f"  P:{stv['pushes']}"
    lines.append(
        f"  SEASON:  {sign}${stv['total_dollars']:.2f}  "
        f"({rec})"
    )
    lines.append(
        f"  BANK:    ${stv['bank']:.2f}  "
        f"(started ${PAPER_BANK_START:.0f})"
    )
    lines.append("")
    return "\n".join(lines)


def log_brief(
    conn,
    game_date,
    session,
    games_covered,
    picks_count,
    output_file,
    pick_entries=None,
    avoid_entries=None,
    now: datetime.datetime | None = None,
    game_group_id: int | None = None,
):
    if now is None:
        now = _now_et()
    generated_at_et = now.strftime("%Y-%m-%d %H:%M ET")
    _gid: int | None = None
    if game_group_id is not None and int(game_group_id) > 0:
        _gid = int(game_group_id)
    try:
        conn.execute(
            """
            INSERT INTO brief_log
                (game_date, session, generated_at, games_covered, picks_count, output_file, game_group_id)
            VALUES (?, ?, ?, ?, ?, ?, ?)
            """,
            (game_date, session, generated_at_et, games_covered, picks_count, output_file, _gid),
        )
        conn.commit()
        # Save confirmed picks for prior-report grading (action sessions only)
        if pick_entries is not None and session in ("primary", "early", "afternoon", "late"):
            save_brief_picks(conn, game_date, session, pick_entries, avoid_entries=avoid_entries, now=now)
    except sqlite3.OperationalError:
        # Best-effort logging: don't fail the brief if another process holds a DB lock.
        try:
            conn.rollback()
        except Exception:
            pass
        return


# ═══════════════════════════════════════════════════════════════════════════
# Data loaders
# ═══════════════════════════════════════════════════════════════════════════

def load_games(conn: sqlite3.Connection, game_date: str, verbose: bool,
               as_of_dt: datetime.datetime | None = None) -> list:
    """
    Load today's games with odds, weather, and team info.
    Returns list of dicts with all fields needed for signal evaluation.

    When as_of_dt is set (historical simulation), loads all regular-season games
    for the date, including Final rows, so main() can filter by simulated clock.
    Live runs use as_of_dt=None and exclude Final games in SQL.
    """
    status_line = (
        ""
        if as_of_dt is not None
        else "          AND  g.status    != 'Final'          -- skip already-completed games\n"
    )
    cur = conn.execute(
        """
        SELECT
            g.game_pk,
            g.game_date_et AS game_date,
            g.season,
            g.venue_id,
            g.game_start_utc,
            v.name          AS venue_name,
            g.temp_f,
            g.wind_mph,
            g.wind_direction,
            g.sky_condition,
            COALESCE(g.wind_source, 'actual') AS wind_source,

            -- Venue intelligence (populated by add_stadium_data.py)
            v.wind_effect,
            v.wind_note,
            v.roof_type,
            v.elevation_ft,
            v.park_factor_runs,
            v.park_factor_hr,
            v.park_factor_hr,
            v.orientation_hp,

            th.team_id      AS home_team_id,
            th.abbreviation AS home_abbr,
            th.name         AS home_name,

            ta.team_id      AS away_team_id,
            ta.abbreviation AS away_abbr,
            ta.name         AS away_name,

            -- Moneyline fields
            ml.home_ml,
            ml.away_ml,
            -- Total fields
            tot.total_line,
            tot.over_odds,
            tot.under_odds,
            -- Run line fields
            rl.home_rl_line  AS home_rl,
            rl.away_rl_line  AS away_rl,
            rl.home_rl_odds,
            rl.away_rl_odds,
            COALESCE(ml.captured_at_utc, tot.captured_at_utc) AS odds_captured_at,
            ml.bookmaker AS odds_bookmaker

        FROM   games g
        JOIN   teams  th ON th.team_id  = g.home_team_id
        JOIN   teams  ta ON ta.team_id  = g.away_team_id
        LEFT JOIN venues v  ON v.venue_id   = g.venue_id
        -- Join each market type separately — one row per game
        LEFT JOIN v_closing_game_odds ml  ON ml.game_pk  = g.game_pk
                                         AND ml.market_type = 'moneyline'
        LEFT JOIN v_closing_game_odds tot ON tot.game_pk = g.game_pk
                                         AND tot.market_type = 'total'
        LEFT JOIN v_closing_game_odds rl  ON rl.game_pk  = g.game_pk
                                         AND rl.market_type = 'runline'
        WHERE  g.game_date_et = ?
          AND  g.game_type = 'R'          -- regular season only; Spring Training / Exhibition excluded
"""
        + status_line
        + """        ORDER  BY g.game_start_utc
        """,
        (game_date,),
    )
    rows = [dict(r) for r in cur.fetchall()]

    if verbose:
        print(f"\n  [verbose] {len(rows)} upcoming games found for {game_date}")
        for r in rows:
            print(f"           {r['away_abbr']} @ {r['home_abbr']}  "
                  f"ML:{fmt_odds(r['home_ml'])}/{fmt_odds(r['away_ml'])}  "
                  f"Tot:{r['total_line']}  "
                  f"Wind:{r['wind_mph']} mph {r['wind_direction']}  "
                  f"WindEffect:{r['wind_effect']}")
    return rows


def load_streaks(conn: sqlite3.Connection, game_date: str, team_ids: list, verbose: bool) -> dict:
    """
    Compute current win/loss streak for each team entering today's game.
    Returns dict of {team_id: streak_int}  (+N = N-game win streak, -N = N-game loss streak).
    """
    streaks = {}
    for tid in team_ids:
        cur = conn.execute(
            """
            SELECT
                CASE WHEN home_team_id = ? THEN
                    CASE WHEN home_score > away_score THEN 'W' ELSE 'L' END
                ELSE
                    CASE WHEN away_score > home_score THEN 'W' ELSE 'L' END
                END AS result
            FROM   games
            WHERE  (home_team_id = ? OR away_team_id = ?)
              AND  status  = 'Final'
              AND  game_date_et < ?
            ORDER  BY game_date_et DESC, game_start_utc DESC
            LIMIT  15
            """,
            (tid, tid, tid, game_date),
        )
        results = [r["result"] for r in cur.fetchall()]
        if not results:
            streaks[tid] = 0
            continue
        current = results[0]
        count   = 0
        for r in results:
            if r == current:
                count += 1
            else:
                break
        streaks[tid] = count if current == "W" else -count

    if verbose:
        print(f"\n  [verbose] Team streaks entering {game_date}:")
        for tid, s in streaks.items():
            label = f"W{s}" if s > 0 else f"L{abs(s)}" if s < 0 else "—"
            print(f"           team_id={tid}  streak={label}")

    return streaks


def load_starters(conn: sqlite3.Connection, game_date: str, verbose: bool) -> dict:
    """
    Load probable starting pitchers for today's games.
    Returns dict of {game_pk: {home_starter, away_starter}} where available.
    Falls back gracefully if the table/column doesn't exist yet.
    """
    starters = {}
    try:
        cur = conn.execute(
            """
            SELECT
                gp.game_pk,
                gp.team_id,
                p.full_name,
                p.era_season,
                p.throws,
                trs.rolling_ops        AS team_rolling_ops,
                trs.games_in_window    AS team_games_in_window
            FROM   game_probable_pitchers gp
            JOIN   players p   ON p.player_id  = gp.player_id
            JOIN   games   g   ON g.game_pk    = gp.game_pk
            LEFT JOIN team_rolling_stats trs
                   ON  trs.game_pk = gp.game_pk
                   AND trs.team_id = gp.team_id
            WHERE  g.game_date_et = ?
            """,
            (game_date,),
        )
        rows = cur.fetchall()
        for r in rows:
            if r["game_pk"] not in starters:
                starters[r["game_pk"]] = {}
            starters[r["game_pk"]][r["team_id"]] = {
                "name":        r["full_name"],
                "era":         r["era_season"],
                "throws":     r["throws"],            # 'L', 'R', or None
                "rolling_ops": r["team_rolling_ops"],  # home team rolling OPS from trs
                "ops_window":  r["team_games_in_window"],
            }
        if verbose:
            print(f"\n  [verbose] Starters loaded for {len(starters)} games")
    except sqlite3.OperationalError as e:
        # Table not yet populated — non-fatal
        if verbose:
            print(f"\n  [verbose] game_probable_pitchers not available: {e}")
    return starters


def load_line_movement(conn: sqlite3.Connection, game_date: str, verbose: bool) -> dict:
    """
    Load open-to-close line movement for today's games.
    Returns dict of {game_pk: movement_row}.
    Only populated after --compute-movement has run (closing session).
    """
    movement = {}
    try:
        cur = conn.execute(
            """
            SELECT
                lm.game_pk,
                lm.ml_move_cents,
                lm.total_move,
                lm.move_direction,
                lm.steam_move,
                lm.reverse_line_move
            FROM   line_movement lm
            JOIN   games g ON g.game_pk = lm.game_pk
            WHERE  g.game_date_et = ?
            """,
            (game_date,),
        )
        for r in cur.fetchall():
            movement[r["game_pk"]] = dict(r)
        if verbose:
            print(f"\n  [verbose] Line movement data found for {len(movement)} games")
    except sqlite3.OperationalError as e:
        if verbose:
            print(f"\n  [verbose] line_movement table not available: {e}")
    return movement


# ═══════════════════════════════════════════════════════════════════════════
# Signal evaluation
# ═══════════════════════════════════════════════════════════════════════════

def enrich_game_with_starters(game: dict, starters: dict) -> None:
    """
    Inject away starter handedness, ERA, name, and home team rolling OPS
    into the game dict. Used by ``enrich_game()`` before dressing; may also be
    called alone when only the raw dict needs starter fields.
    Modifies game in place. Safe to call even if starters is empty.
    """
    gpk     = game.get("game_pk")
    home_id = game.get("home_team_id")
    away_id = game.get("away_team_id")
    game_starters = starters.get(gpk, {})

    away_s = game_starters.get(away_id, {})
    home_s = game_starters.get(home_id, {})

    game["away_starter_throw"] = away_s.get("throws")
    game["away_starter_era"]   = away_s.get("era")
    game["away_starter_name"]  = away_s.get("name")
    game["home_rolling_ops"]   = home_s.get("rolling_ops")
    game["home_ops_window"]    = home_s.get("ops_window") or 0


def enrich_game(conn: sqlite3.Connection, game: dict, starters: dict):
    """
    Starter-enrich the brief row dict and dress it for scoring.

    Returns ``FullyDressedGame`` (see ``batch.pipeline.score_game.dress_game_for_brief``).
    """
    enrich_game_with_starters(game, starters)
    from batch.pipeline.score_game import dress_game_for_brief

    return dress_game_for_brief(conn, game)


def evaluate_signals(
    conn: sqlite3.Connection,
    game: dict,
    streaks: dict,
    session: str,
    starters: dict | None = None,
    *,
    verbose: bool = False,
    debug_wind: bool = False,
) -> dict:
    """
    Evaluate all model signals for a single game.
    Returns a dict describing which signals fired and what they recommend.

    Flow: ``enrich_game`` → ``score_game`` → ``scored_game_to_eval_dict``. Extra keys:
    ``output_tier``, ``tier_basis``, ``stake_multiplier`` (from ``ScoredGame``);
    ``best_aggregate_score``; ``signal_brief`` (Core / Support / Minor by per-signal score).
    Per-pick ``aggregate_score`` = sum of 1–10 signal scores in that bet bucket
    (ML/total/…); the brief shows ``[HIGH · nn%]`` / ``[MED · nn%]`` / ``[LOW · nn%]``
    (tier label + percent from that aggregate). ``signals`` = short reader labels
    (no internal codes); ``signal_ids`` = internal ids for filters and analytics.
    ``_scored_game`` is the in-process ``ScoredGame`` for helpers like ``format_bet_block``
    (omit when serializing).
    When ``verbose`` is True, prints tier_basis per game to stdout.
    When ``debug_wind`` is True, prints a wind classification dump per game (DB vs dressed env).
    """
    from batch.pipeline.score_game import score_game, scored_game_to_eval_dict
    from batch.pipeline.dressed_game_blocks import SignalFinding

    if conn is None:
        raise TypeError("evaluate_signals requires a sqlite3.Connection (conn)")

    starters = starters if starters is not None else {}

    hid = int(game["home_team_id"])
    aid = int(game["away_team_id"])
    home_streak = int(streaks.get(hid, 0))
    away_streak = int(streaks.get(aid, 0))

    # --- Dress (may fail); never gate score_game on signals/picks ---
    fdg = None
    dress_exc: BaseException | None = None
    try:
        fdg = enrich_game(conn, game, starters)
        from dataclasses import replace as _dc_replace

        fdg = _dc_replace(
            fdg,
            brief_session=session,
            home_streak=home_streak,
            away_streak=away_streak,
        )
        if debug_wind:
            print_wind_debug_for_game(game, fdg)
    except Exception as e:
        dress_exc = e
        fdg = None

    # --- score_game ALWAYS runs once we have a dressed game (no if signals / if picks) ---
    scored = None
    score_exc: BaseException | None = None
    if fdg is not None:
        try:
            gd = fdg.identifiers.game_date_et
            game_month = int(gd[5:7]) if len(gd) >= 7 else 0
            # Guarantee a non-empty "signals" list on the dressed game for debugging.
            # This does NOT create a real model edge because it does not fire and has no bet_side impact.
            if not list(getattr(fdg, "signals", None) or []):
                fdg = _dc_replace(
                    fdg,
                    signals=[
                        SignalFinding(
                            signal_id="baseline",
                            signal_strength="weak",
                            bet_side="",
                            odds="N/A",
                            edge_basis="Baseline placeholder (debug-only).",
                            fires=False,
                        )
                    ],
                )
            if os.getenv("DEBUG_SCORE_GAME") == "1":
                try:
                    game_pk = int(game.get("game_pk"))
                except Exception:
                    game_pk = -1
                pre = list(getattr(fdg, "signals", None) or [])
                signals = [getattr(s, "signal_id", str(s)) for s in pre]
                print(f"[DEBUG BEFORE SCORE] game={game_pk} signals={signals}")
            scored = score_game(fdg, home_streak, game_month)
        except Exception as e:
            score_exc = e
            scored = None

    out: dict | None = None

    if dress_exc is not None:
        out = {
            "signals": [],
            "signal_ids": [],
            "picks": [],
            "signal_brief": "",
            "best_aggregate_score": 0,
            "avoid": False,
            "avoid_reason": None,
            "watch": False,
            "watch_reason": None,
            "data_flags": [
                f"Dress pipeline failed (score_game not run): {dress_exc}",
                "Score penalty: no dressed game — treat as no edge (aggregate held at 0).",
            ],
            "output_tier": None,
            "tier_basis": "",
            "stake_multiplier": 0.0,
            "_scored_game": None,
            "market_evals": {},
        }
    elif score_exc is not None:
        out = {
            "signals": [],
            "signal_ids": [],
            "picks": [],
            "signal_brief": "",
            "best_aggregate_score": 0,
            "avoid": False,
            "avoid_reason": None,
            "watch": False,
            "watch_reason": None,
            "data_flags": [
                f"score_game() failed after successful dress: {score_exc}",
                "Score penalty: scoring exception — treat as no edge (aggregate held at 0).",
            ],
            "output_tier": None,
            "tier_basis": "",
            "stake_multiplier": 0.0,
            "_scored_game": None,
            "market_evals": {},
        }
    elif scored is not None:
        try:
            out = scored_game_to_eval_dict(scored, session)
            out["output_tier"] = scored.output_tier
            out["tier_basis"] = scored.tier_basis
            out["stake_multiplier"] = scored.stake_multiplier
            if verbose and scored.tier_basis:
                ab = f"{fdg.identifiers.away_team_abbr}@{fdg.identifiers.home_team_abbr}"
                print(f"  [verbose] score {ab}: tier={scored.output_tier!r} | {scored.tier_basis}")
            out["_scored_game"] = scored
            out["market_evals"] = getattr(scored, "market_evals", {}) or {}

            if PERSIST_WRITES:
                try:
                    gpk = int(game.get("game_pk"))
                    gd = str(fdg.identifiers.game_date_et)
                    placed_at = _now_et().strftime("%Y-%m-%d %H:%M ET")
                    for mt, mm in (out["market_evals"] or {}).items():
                        if not isinstance(mm, dict) or not mm.get("evaluated"):
                            continue
                        save_bet_snapshot(
                            conn,
                            gd,
                            gpk,
                            str(mt),
                            str(mm.get("bet") or ""),
                            int(mm["odds"]) if mm.get("odds") is not None else None,
                            mm,
                            list(mm.get("signals") or []),
                            placed_at=placed_at,
                        )
                except Exception:
                    pass
        except Exception as e:
            flags = list(getattr(scored, "data_flags", None) or [])
            flags.append(f"Brief mapping failed after score_game (ScoredGame retained): {e}")
            out = {
                "signals": [],
                "signal_ids": [],
                "picks": [],
                "signal_brief": "",
                "best_aggregate_score": int(scored.best_aggregate_score or 0),
                "avoid": False,
                "avoid_reason": None,
                "watch": False,
                "watch_reason": None,
                "data_flags": flags,
                "output_tier": scored.output_tier,
                "tier_basis": scored.tier_basis or "",
                "stake_multiplier": float(scored.stake_multiplier or 0.0),
                "_scored_game": scored,
                "market_evals": getattr(scored, "market_evals", {}) or {},
            }

    if out is None:
        raise RuntimeError("evaluate_signals: internal error (no result set)")

    if os.getenv("DEBUG_SCORE_GAME") == "1":
        try:
            game_pk = int(game.get("game_pk"))
        except Exception:
            game_pk = -1
        raw_signals = list(out.get("signal_ids") or [])
        if not raw_signals:
            raw_signals = ["baseline"]
        print(f"[DEBUG SIGNAL GEN FINAL] game={game_pk} signals={raw_signals}")

    return out


# ═══════════════════════════════════════════════════════════════════════════
# Brief formatting helpers
# ═══════════════════════════════════════════════════════════════════════════

# -----------------------------------------------------------------------------
# Confidence conversion
# -----------------------------------------------------------------------------
def score_to_confidence(score: int) -> int:
    """Normalize roughly 5–30 range → 50–95%."""
    return min(95, max(50, int(50 + (score - 5) * 2)))


def tier_label(tier: str | None) -> str:
    return {
        "Tier1": "HIGH",
        "Tier2": "MED",
        "Tier3": "LOW",
    }.get(tier, "NO EDGE")


# -----------------------------------------------------------------------------
# Signal name mapping (display only; internal ids unchanged)
# -----------------------------------------------------------------------------
SIGNAL_LABELS = {
    "S1+H2": "Streak Fade",
    "S1H2": "Streak Fade",
    "S1": "Streak Pressure",
    "MV-F": "Wind Fade (ML)",
    "MV-B": "Wind Boost (Over)",
    "LHP_FADE": "LHP Mismatch",
    "LHP_FADE_RL": "LHP RL Edge",
    "NF4": "Pitching Edge",
    "H3b": "Wind → Over",
}


def humanize_signal(sig_id: str) -> str:
    if sig_id in SIGNAL_LABELS:
        return SIGNAL_LABELS[sig_id]
    from batch.pipeline.score_game import signal_display_name

    return signal_display_name(sig_id)


# -----------------------------------------------------------------------------
# Signal grouping (per-signal confidence_score on model findings)
# -----------------------------------------------------------------------------
def group_signals(signals: list) -> dict[str, list[str]]:
    core: list[str] = []
    support: list[str] = []
    minor: list[str] = []

    def _sc(s: object) -> int:
        v = getattr(s, "confidence_score", None)
        try:
            return int(v) if v is not None else 0
        except (TypeError, ValueError):
            return 0

    for s in sorted(signals, key=_sc, reverse=True):
        sid = str(getattr(s, "signal_id", "") or "")
        name = humanize_signal(sid)
        sc = _sc(s)
        if sc >= 7:
            core.append(name)
        elif sc >= 5:
            support.append(name)
        else:
            minor.append(name)

    return {
        "Core": core,
        "Support": support,
        "Minor": minor,
    }


def generate_why_line(signals: list) -> str:
    if not signals:
        return "No clear edge."

    sigs = sorted(signals, key=lambda x: int(getattr(x, "confidence_score", 0) or 0), reverse=True)

    label_map = {
        "S1+H2": "streak fade",
        "S1H2": "streak fade",
        "S1": "streak pressure",
        "MV-F": "wind fade conditions",
        "MV-B": "wind boost for scoring",
        "H3b": "over conditions from wind",
        "LHP_FADE": "lefty/righty mismatch",
        "LHP_FADE_RL": "run line matchup edge",
        "NF4": "pitching mismatch",
    }

    top: list[str] = []
    for s in sigs:
        sid = str(getattr(s, "signal_id", "") or "")
        label = label_map.get(sid, sid.lower())
        if label and label not in top:
            top.append(label)
        if len(top) == 3:
            break

    if not top:
        return "No clear edge."

    # Match example tone: if the top driver is very strong, prefix "strong".
    try:
        top_score = int(getattr(sigs[0], "confidence_score", 0) or 0)
    except Exception:
        top_score = 0
    if top_score >= 8 and not top[0].startswith("strong "):
        top[0] = "strong " + top[0]

    if len(top) == 1:
        return f"WHY: {top[0].capitalize()} driving the edge."
    if len(top) == 2:
        return f"WHY: {top[0].capitalize()} with support from {top[1]}."
    return f"WHY: {top[0].capitalize()} with support from {top[1]} and {top[2]}."


def color_text(text: str, color: str) -> str:
    colors = {
        "green": "\033[92m",
        "yellow": "\033[93m",
        "red": "\033[91m",
        "dim": "\033[2m",
        "bold": "\033[1m",
        "underline": "\033[4m",
        "reset": "\033[0m",
    }
    c = colors.get(color, "")
    r = colors["reset"] if c else ""
    return f"{c}{text}{r}"


def format_bet_block(scored_game: object) -> str:
    """
    ASCII pick card from a ``ScoredGame`` (primary / additional brief rows).

    ``[HIGH · nn%]`` uses **game** ``best_aggregate_score`` + ``output_tier``.
    BET shows the aggregate bucket id (e.g. ``AWAY_ML``). Signals: all
    ``signals_fired``, sorted by per-signal ``confidence_score``, grouped
    Core (≥7) / Support (5–6) / Minor (≤4) with ``humanize_signal`` labels.
    """
    tier = getattr(scored_game, "output_tier", None)
    stake = float(getattr(scored_game, "stake_multiplier", 0.0) or 0.0)
    best_score = int(getattr(scored_game, "best_aggregate_score", 0) or 0)
    confidence = score_to_confidence(best_score)
    tier_txt = tier_label(tier)

    grouped = group_signals(list(getattr(scored_game, "signals_fired", []) or []))

    def fmt_group(label: str, items: list[str]) -> str:
        return f"{label}: " + " · ".join(items) if items else ""

    signal_lines = "\n".join(
        filter(
            None,
            [
                fmt_group("Core", grouped["Core"]),
                fmt_group("Support", grouped["Support"]),
                fmt_group("Minor", grouped["Minor"]),
            ],
        )
    )
    if not signal_lines.strip():
        signal_lines = "(none)"

    why_line = generate_why_line(list(getattr(scored_game, "active_bets", []) or []))

    pick = getattr(scored_game, "top_pick", None)
    best_side = getattr(scored_game, "best_side", None)
    if pick is not None and getattr(pick, "bet_side", None):
        bet_txt = str(pick.bet_side).upper()
    else:
        bet_txt = str(best_side or "unknown").upper()

    if stake > 0:
        head = color_text("🔥 BET:", "green")
        head_line = f"│  {head}  {bet_txt:<18}  [{tier_txt} · {confidence}%]"
        stake_line = color_text(f"│  STAKE: {stake:.2f}u  ← PLAY THIS", "green")
    else:
        head = color_text("❌ NO BET:", "dim")
        head_line = f"│  {head}  {bet_txt:<18}  [{confidence}%]"
        stake_line = color_text("│  STAKE: 0.00u  ← SKIP", "dim")

    return f"""
┌─────────────────────────────────────────────────────────┐
{head_line}
│
│  {why_line}
│
│  SIGNAL:
{signal_lines}
│
│  MODEL SCORE: {best_score}
{stake_line}
└─────────────────────────────────────────────────────────┘
""".strip()


def signal_summary_for_doc(sigs: dict) -> str:
    """Word table SIGNAL column: same Core / Support / Minor grouping as ``format_bet_block``."""
    sg = sigs.get("_scored_game")
    if sg is not None:
        grouped = group_signals(list(getattr(sg, "signals_fired", []) or []))
        parts: list[str] = []
        if grouped["Core"]:
            parts.append("Core: " + " · ".join(grouped["Core"]))
        if grouped["Support"]:
            parts.append("Support: " + " · ".join(grouped["Support"]))
        if grouped["Minor"]:
            parts.append("Minor: " + " · ".join(grouped["Minor"]))
        if parts:
            return "\n".join(parts)
    fb = (sigs.get("signal_brief") or "").strip()
    if fb:
        return fb
    return ", ".join(sigs.get("signals") or []) or ""


def matchup_line(game: dict) -> str:
    """Return formatted matchup string with (h) home indicator and ET start time."""
    away  = game.get("away_abbr", "AWAY")
    home  = game.get("home_abbr", "HOME")
    venue = game.get("venue_name") or ""

    # Convert game_start_utc → Eastern Time for display
    start_str = ""
    raw = game.get("game_start_utc") or ""
    if raw and "T" in raw:
        try:
            utc_dt = datetime.datetime.fromisoformat(raw.rstrip("Z")).replace(
                tzinfo=datetime.timezone.utc)
            et_dt  = utc_dt.astimezone(_ET)
            # %-I is POSIX-only and fails on Windows.
            # lstrip('0') on %I:%M %p achieves the same result cross-platform.
            # Guard: if stripping leaves ':MM %p' (midnight edge), keep '12'.
            t = et_dt.strftime('%I:%M %p').lstrip('0') or et_dt.strftime('%I:%M %p')
            start_str = f"  {t} ET"
        except (ValueError, AttributeError):
            pass

    return f"{away}  vs  {home} (h)    [{venue}]{start_str}"


def weather_line(game: dict, *, wind_signal_hints: bool = True) -> str:
    """Return formatted weather / conditions line.

    When ``wind_signal_hints`` is False (morning sneak peek), omit in-line
    wind “signal” callouts — only factual temp/wind/venue suppression text.
    """
    wind_effect = (game.get("wind_effect") or "HIGH").upper()
    roof_type   = game.get("roof_type") or "Open"

    parts = []

    # Roof / dome status prefix
    if wind_effect == "SUPPRESSED":
        # Use specific note when available, fall back to roof type
        note = game.get("wind_note") or ""
        first_sentence = note.split(".")[0] if note else f"{roof_type} — wind signals suppressed"
        parts.append(f"⚠ {first_sentence}")
    # LOW wind_effect: no extra venue headline (temp/wind only).

    # Temperature
    if game.get("temp_f") is not None:
        parts.append(f"{int(game['temp_f'])}°F")

    # Wind — only show signal flag for HIGH-effect venues
    if game.get("wind_mph") is not None:
        mph       = game["wind_mph"]
        direction = wind_direction_label(game.get("wind_direction") or "")
        parts.append(f"{mph} mph wind {direction}")
        if wind_signal_hints:
            if wind_effect == "HIGH" and mph >= WIND_OUT_MIN_MPH and direction in ("OUT", "IN"):
                parts.append("⚑ WIND SIGNAL")
            elif wind_effect == "MODERATE" and mph >= WIND_OUT_MIN_MPH and direction in ("OUT", "IN"):
                parts.append("~ wind (moderate venue)")

    return "  ".join(parts) if parts else "Conditions not available"


def odds_summary_line(game: dict) -> str:
    """Return compact odds summary with bookmaker source."""
    home = game.get("home_abbr", "HOME")
    away = game.get("away_abbr", "AWAY")
    hml  = fmt_odds(game.get("home_ml"))
    aml  = fmt_odds(game.get("away_ml"))
    tot  = fmt_total(game.get("total_line"))
    hrl  = game.get("home_rl")
    hrl_o = game.get("home_rl_odds")
    arl = game.get("away_rl")
    arl_o = game.get("away_rl_odds")
    book = game.get("odds_bookmaker") or ""
    # Clean up bookmaker name for display
    book_display = {
        "draftkings": "DK", "fanduel": "FD", "betmgm": "MGM",
        "betonlineag": "BOL", "fanatics": "FAN", "betrivers": "BR",
        "williamhill_us": "WH", "bovada": "BOV", "betus": "BetUS",
        "lowvig": "LV", "mybookieag": "MYB",
    }.get(book, book.upper() if book else "")
    src = f"  [{book_display}]" if book_display else ""
    rl_parts: list[str] = []
    if hrl is not None:
        rl_parts.append(f"{home} {'+' if float(hrl) >= 0 else ''}{hrl:g} ({fmt_odds(hrl_o)})")
    if arl is not None:
        rl_parts.append(f"{away} {'+' if float(arl) >= 0 else ''}{arl:g} ({fmt_odds(arl_o)})")
    rl_str = f"  RL: {' / '.join(rl_parts)}" if rl_parts else ""
    return f"ML: {home} {hml} / {away} {aml}  |  {tot}{rl_str}{src}"


def _avoid_scope_line(game: dict, sigs: dict) -> str | None:
    """
    Return a concise, explicit description of WHAT is being avoided.
    Avoid flags should never read like "avoid the whole game" unless it truly is.
    """
    scope = (sigs.get("avoid_scope") or "").strip()
    if scope:
        return scope

    # Fallback for older avoids: infer from reason text.
    reason = (sigs.get("avoid_reason") or "").strip()
    ru = reason.upper()
    if "WIND" in ru and ("SUPPRESSED" in ru or "NO WEATHER" in ru):
        return "WIND SIGNALS ONLY (totals/wind edges) — other analysis OK"
    if "AVOID HOME ML" in ru or "AVOID AWAY ML" in ru or " ML" in ru:
        return "MONEYLINE ONLY (ML) — other markets OK"
    if "OVER" in ru or "UNDER" in ru or "TOTAL" in ru or "O/U" in ru:
        return "TOTALS ONLY (O/U) — other markets OK"
    return None


def starter_line(game: dict, starters: dict) -> str:
    """Return starting pitcher line if available."""
    gpk = game.get("game_pk")
    if not gpk or gpk not in starters:
        return "Starters: not yet confirmed"
    s       = starters[gpk]
    home_id = game.get("home_team_id")
    away_id = game.get("away_team_id")
    home_s  = s.get(home_id, {})
    away_s  = s.get(away_id, {})

    def fmt_starter(d):
        if not d:
            return "TBD"
        era = f" (ERA {d['era']:.2f})" if d.get("era") is not None else ""
        return f"{d.get('name', 'TBD')}{era}"

    home_name = game.get("home_abbr", "HOME")
    away_name = game.get("away_abbr", "AWAY")
    return f"Starters: {away_name} {fmt_starter(away_s)}  vs  {home_name} (h) {fmt_starter(home_s)}"


def streak_line(game: dict, streaks: dict) -> str:
    """Return home/away streak line."""
    home_id = game.get("home_team_id")
    away_id = game.get("away_team_id")
    home_s  = streaks.get(home_id, 0)
    away_s  = streaks.get(away_id, 0)

    def label(s):
        if s > 0:
            return f"W{s}"
        if s < 0:
            return f"L{abs(s)}"
        return "—"

    h = game.get("home_abbr", "HOME")
    a = game.get("away_abbr", "AWAY")
    flags = []
    if abs(home_s) >= STREAK_THRESHOLD:
        flags.append(f"⚑ {h} (h) streak signal")
    return (f"Streak: {a} {label(away_s)}  |  {h} (h) {label(home_s)}"
            + (f"  {'  '.join(flags)}" if flags else ""))


def movement_line(game: dict, movement: dict) -> str:
    """Return line movement summary for closing session."""
    gpk = game.get("game_pk")
    if not gpk or gpk not in movement:
        return ""
    m = movement[gpk]
    parts = []
    if m.get("ml_move_cents") is not None:
        cents = float(m["ml_move_cents"])
        arrow = "▲" if cents > 0 else "▼" if cents < 0 else "→"
        parts.append(f"ML move: {arrow} {abs(cents):.1f}¢ ({m.get('move_direction', '?')})")
    if m.get("total_move") is not None:
        tm = float(m["total_move"])
        arrow = "▲" if tm > 0 else "▼" if tm < 0 else "→"
        parts.append(f"Total move: {arrow} {abs(tm):.1f}")
    flags = []
    if m.get("steam_move"):
        flags.append("🔥 STEAM")
    if m.get("reverse_line_move"):
        flags.append("↔ REVERSE LINE")
    if flags:
        parts.append("  ".join(flags))
    return "  |  ".join(parts) if parts else "No significant movement"


# ═══════════════════════════════════════════════════════════════════════════
# Brief body builders per session
# ═══════════════════════════════════════════════════════════════════════════


def build_prior_day_report(conn: sqlite3.Connection, game_date: str,
                            verbose: bool, now: datetime.datetime | None = None,
                            *, debug_wind: bool = False) -> str:
    """
    Prior day performance report — full slate with signal grading.

    Shows every game played yesterday with:
      · Final score, runs, winner, conditions
      · Closing odds and O/U outcome
      · Which model signals fired and whether they won
      · Full slate per-game bet outcomes (or NO SIGNAL)
      · Day-level P&L summary
    """
    lines = []
    if now is None:
        now = _now_et()
    generated_ts = now.strftime("%Y-%m-%d %I:%M %p ET").lstrip("0")
    lines.append(banner(f"MLB Scout · PRIOR DAY REPORT  ·  Results for {game_date}"))
    lines.append(f"  Generated: {generated_ts}\n")
    lines.append(
        "\n  Full slate results with model signal grading.\n"
        "  Every game shown. Signal picks highlighted with ✓ WIN / ✗ LOSS / — PUSH.\n"
    )

    # ── Load all completed games ──────────────────────────────────────────
    cur = conn.execute("""
        SELECT
            g.game_pk,
            g.game_date_et AS game_date,
            g.season,
            g.venue_id,
            g.home_score, g.away_score,
            g.wind_mph, g.wind_direction, g.temp_f, g.sky_condition,
            COALESCE(g.wind_source, 'actual') AS wind_source,
            g.game_start_utc,
            th.team_id   AS home_team_id,
            ta.team_id   AS away_team_id,
            th.abbreviation AS home_abbr,
            ta.abbreviation AS away_abbr,
            th.name         AS home_name,
            ta.name         AS away_name,
            v.name          AS venue_name,
            v.wind_effect, v.wind_note, v.roof_type,
            v.park_factor_runs, v.park_factor_hr, v.orientation_hp,
            go_ml.home_ml,   go_ml.away_ml,
            go_tot.total_line, go_tot.over_odds, go_tot.under_odds,
            rl.home_rl_line  AS home_rl,
            rl.away_rl_line  AS away_rl,
            rl.home_rl_odds,
            rl.away_rl_odds
        FROM   games g
        JOIN   teams  th     ON th.team_id  = g.home_team_id
        JOIN   teams  ta     ON ta.team_id  = g.away_team_id
        LEFT JOIN venues v   ON v.venue_id  = g.venue_id
        LEFT JOIN v_closing_game_odds go_ml  ON go_ml.game_pk  = g.game_pk
                                             AND go_ml.market_type  = 'moneyline'
        LEFT JOIN v_closing_game_odds go_tot ON go_tot.game_pk = g.game_pk
                                             AND go_tot.market_type = 'total'
        LEFT JOIN v_closing_game_odds rl     ON rl.game_pk = g.game_pk
                                             AND rl.market_type = 'runline'
        WHERE  g.game_date_et = ?
          AND  g.game_type = 'R'
          AND  g.status    = 'Final'
        ORDER  BY g.game_start_utc, g.game_pk
    """, (game_date,))
    games = [dict(r) for r in cur.fetchall()]

    if not games:
        lines.append(f"  No completed regular-season games found for {game_date}.")
        lines.append(f"  Ensure the 6 AM stats pull ran: python load_mlb_stats.py\n")
        lines.append(CAVEAT)
        return "\n".join(lines)

    # Backfill snapshots from existing ledger rows if needed (older days before snapshots existed).
    try:
        backfill_missing_bet_snapshots_from_ledger(conn, game_date)
    except Exception:
        pass

    # Load snapshots (source of truth; PRIOR report does not recompute signals/scores)
    try:
        ensure_bet_snapshots(conn)
    except Exception:
        pass
    try:
        snap_rows = conn.execute(
            """
            SELECT game_pk, market_type, bet_side, bet, odds_taken, score, model_p, implied_p, edge,
                   eval_status, signals_used, placed_at
            FROM bet_snapshots
            WHERE game_date = ?
            """,
            (game_date,),
        ).fetchall()
        snapshots = {(int(r["game_pk"]), str(r["market_type"] or "").upper()): dict(r) for r in snap_rows}
    except Exception:
        snapshots = {}

    # ── Skipped edge summary ─────────────────────────────────────────────
    def _snap_status(s: dict) -> str:
        st = str(s.get("eval_status") or "").strip()
        if not st:
            # Legacy snapshots (staked bets) before eval_status existed.
            return "BET" if s.get("odds_taken") is not None else "NO_EDGE"
        return st

    skipped = [s for s in snapshots.values() if _snap_status(s) == "SKIPPED_EDGE"]
    if skipped:
        edges = [float(s["edge"]) for s in skipped if s.get("edge") is not None]
        avg_edge = (sum(edges) / len(edges)) if edges else 0.0
        best = max(skipped, key=lambda x: float(x.get("edge") or 0.0))
        best_edge = float(best.get("edge") or 0.0)
        best_bet = str(best.get("bet") or "")
    else:
        avg_edge = 0.0
        best_edge = 0.0
        best_bet = ""

    lines.append(section("🟡  SKIPPED EDGE SUMMARY"))
    lines.append(f"\n  Skipped Edges: {len(skipped)}")
    lines.append(f"  Avg Edge: +{avg_edge*100:.1f}%")
    if best_bet:
        lines.append(f"  Max Skipped: +{best_edge*100:.1f}% ({best_bet})")
    lines.append("")

    # ── Grading helpers ───────────────────────────────────────────────────
    def grade_ml(bet_side, hs, as_, odds):
        if hs is None or as_ is None:
            return "— NO RESULT", 0.0
        if odds in (None, 0):
            return "— NO RESULT", 0.0
        won = (bet_side == "home" and hs > as_) or (bet_side == "away" and as_ > hs)
        if won:
            pnl = odds / 100.0 if odds > 0 else 100.0 / abs(odds)
            return "✓ WIN", round(pnl, 2)
        return "✗ LOSS", -1.0

    def grade_total(bet, hs, as_, total, odds):
        if hs is None or as_ is None or total is None:
            return "— NO RESULT", 0.0
        runs = hs + as_
        if runs == total:
            return "— PUSH", 0.0
        hit = (bet == "over" and runs > total) or (bet == "under" and runs < total)
        if hit:
            pnl = odds / 100.0 if (odds and odds > 0) else 100.0 / abs(odds or 110)
            return "✓ WIN", round(pnl, 2)
        return "✗ LOSS", -1.0

    def grade_runline(team: str, line: float, hs: int, as_: int,
                      *, home_abbr: str, away_abbr: str,
                      odds: int | None):
        """Grade a runline/spread bet for one side at one line."""
        if hs is None or as_ is None or line is None:
            return "— NO RESULT", 0.0
        side = (team or "").strip().upper()
        home_abbr_u = (home_abbr or "").strip().upper()
        away_abbr_u = (away_abbr or "").strip().upper()
        if side == home_abbr_u:
            adj = float(hs) + float(line)
            opp = float(as_)
        elif side == away_abbr_u:
            adj = float(as_) + float(line)
            opp = float(hs)
        else:
            return "— NO RESULT", 0.0
        if adj == opp:
            return "— PUSH", 0.0
        won = adj > opp
        if won:
            if odds is None:
                return "✓ WIN", 0.0
            pnl = odds / 100.0 if odds > 0 else 100.0 / abs(odds)
            return "✓ WIN", round(pnl, 2)
        return "✗ LOSS", -1.0

    def _parse_runline_bet(bet_text: str) -> tuple[str | None, float | None]:
        try:
            parts = (bet_text or "").strip().split()
            if len(parts) < 2:
                return None, None
            team = parts[0].strip().upper()
            line = float(parts[1])
            return team, line
        except Exception:
            return None, None

    # ── Snapshot-based grading (no recomputation) ──────────────────────────
    evaluated = []
    for g in games:
        hs = g["home_score"]; as_ = g["away_score"]
        tot = g["total_line"]
        runs = (hs + as_) if (hs is not None and as_ is not None) else None
        ou_label = ""
        if tot is not None and runs is not None:
            ou_label = (f"OVER {tot} ({runs} runs)"  if runs > tot else
                        f"UNDER {tot} ({runs} runs)" if runs < tot else
                        f"PUSH {tot} ({runs} runs)")
        evaluated.append({
            "game": g,
            "runs": runs,
            "ou_label": ou_label,
            "winner": g["home_abbr"] if (hs or 0) > (as_ or 0) else g["away_abbr"],
            "snapshots": {
                "ML": snapshots.get((int(g["game_pk"]), "ML")),
                "TOTAL": snapshots.get((int(g["game_pk"]), "TOTAL")),
                "RL": snapshots.get((int(g["game_pk"]), "RL")),
            },
        })

    def pnl_str(pnl):
        return f"+{pnl:.2f}u" if pnl > 0 else (f"{pnl:.2f}u" if pnl < 0 else "push")

    def game_score_line(e, indent="  "):
        g = e["game"]
        hs = g["home_score"]; as_ = g["away_score"]
        s  = f"{g['away_abbr']} {as_}  –  {g['home_abbr']} {hs}  ({e['winner']} wins)"
        if e["ou_label"]:
            s += f"  |  {e['ou_label']}"
        return f"{indent}Final: {s}"

    def game_odds_line(e, indent="  "):
        g   = e["game"]
        hml = g["home_ml"]; aml = g["away_ml"]; tot = g["total_line"]
        if not hml and not aml:
            return ""
        return (f"{indent}ML: {g['home_abbr']} {fmt_odds(hml)} / "
                f"{g['away_abbr']} {fmt_odds(aml)}"
                + (f"  |  O/U {tot}" if tot else ""))

    # ════════════════════════════════════════════════════════════════════
    # BET LEDGER (source of truth) — backfill from signal_state, then grade
    # ════════════════════════════════════════════════════════════════════
    try:
        if PERSIST_WRITES:
            backfill_bet_ledger_from_signal_state(conn, game_date, now=now)
    except Exception:
        pass
    try:
        grade_bet_ledger(conn, game_date=game_date)
    except Exception:
        pass

    try:
        bet_rows = conn.execute("""
            SELECT
                id, game_pk, market_type, bet, odds_taken, stake_units,
                signal_at_time, session, placed_at, result, pnl_units
            FROM bet_ledger
            WHERE game_date = ?
            ORDER BY placed_at
        """, (game_date,)).fetchall()
        bet_rows = [dict(r) for r in bet_rows]
    except Exception:
        bet_rows = []

    def _ledger_signal_label(sig: str | None) -> str:
        s = (sig or "").strip().lower()
        if s == "top":
            return "TOP PICK"
        if s == "next":
            return "NEXT (additional)"
        if s == "avoid":
            return "AVOID (do not bet)"
        return (sig or "").upper() or "—"

    def _summarise_bets(rows: list) -> dict:
        stake_rows = [
            r for r in rows
            if (r.get("signal_at_time") or "").lower() != "avoid"
        ]
        avoid_rows = [
            r for r in rows
            if (r.get("signal_at_time") or "").lower() == "avoid"
        ]
        graded = [r for r in stake_rows if r.get("result") in ("win", "loss", "push")]
        wins   = sum(1 for r in graded if r["result"] == "win")
        losses = sum(1 for r in graded if r["result"] == "loss")
        pushes = sum(1 for r in graded if r["result"] == "push")
        units  = sum(float(r.get("pnl_units") or 0.0) for r in graded)
        n      = len(graded)
        roi    = (100.0 * units / n) if n else 0.0
        av_done = [r for r in avoid_rows if r.get("result") in ("good_avoid", "bad_avoid", "push_avoid")]
        agood = sum(1 for r in av_done if r["result"] == "good_avoid")
        abad = sum(1 for r in av_done if r["result"] == "bad_avoid")
        apush = sum(1 for r in av_done if r["result"] == "push_avoid")
        return {
            "bets": n,
            "wins": wins,
            "losses": losses,
            "pushes": pushes,
            "units": units,
            "roi": roi,
            "avoid_graded": len(av_done),
            "avoid_good": agood,
            "avoid_bad": abad,
            "avoid_push": apush,
        }

    bet_summary = _summarise_bets(bet_rows)
    by_market: dict[str, list] = {}
    for r in bet_rows:
        m = (r.get("market_type") or "unknown").strip() or "unknown"
        by_market.setdefault(m, []).append(r)

    # ════════════════════════════════════════════════════════════════════
    # FULL SLATE — per-game, per-bet grading
    # ════════════════════════════════════════════════════════════════════
    lines.append(section(f"📋  FULL SLATE  ({len(evaluated)} games)"))
    lines.append(
        "\n  For each game: prints a moneyline (ML) line and a totals (O/U) line.\n"
        "  If a model signal fired for that market, it is shown with P&L.\n"
        "  If not, SIGNAL shows 'No Signal' and P&L is N/A.\n"
    )

    def _clean_result(res: str) -> str:
        r = (res or "").upper()
        if "WIN" in r:
            return "WIN"
        if "LOSS" in r:
            return "LOSS"
        if "PUSH" in r:
            return "PUSH"
        return (res or "—").strip() or "—"

    def _best_pick_for_market(graded: list, market: str) -> dict | None:
        picks = [p for p in graded if (p.get("market") or "").upper() == market.upper()]
        if not picks:
            return None
        return sorted(picks, key=lambda x: x.get("priority", 99))[0]

    for e in evaluated:
        g = e["game"]
        lines.append(f"\n  {matchup_line(g)}")
        lines.append(f"  {weather_line(g)}")
        if g.get("home_score") is not None:
            lines.append(game_score_line(e))
        ol = game_odds_line(e)
        if ol:
            lines.append(ol)

        def _signal_display_for_snapshot(snap: dict | None) -> str:
            if not snap:
                return "No Signal"
            status = _snap_status(snap)
            edge = snap.get("edge")
            if status == "BET":
                try:
                    import json
                    raw = snap.get("signals_used") or "[]"
                    sigs = json.loads(raw) if isinstance(raw, str) else list(raw)
                    return ", ".join(str(x) for x in sigs) if sigs else "Signal Missing (BUG)"
                except Exception:
                    return "Signal Missing (BUG)"
            if status == "SKIPPED_EDGE":
                try:
                    return f"Weak Edge (+{float(edge)*100:.1f}%)" if edge is not None else "Weak Edge (+?.?%)"
                except Exception:
                    return "Weak Edge (+?.?%)"
            if status == "NO_MODEL":
                return "No Model"
            return "No Edge"

        # ML line
        ml_snap = (e.get("snapshots") or {}).get("ML")
        ml_sig = _signal_display_for_snapshot(ml_snap)
        if ml_snap and _snap_status(ml_snap) == "BET":
            bet_txt = str(ml_snap.get("bet") or "")
            odds = ml_snap.get("odds_taken")
            side = "away" if bet_txt.startswith(g["away_abbr"]) else "home"
            res, pnl_v = grade_ml(side, hs, as_, int(odds) if odds is not None else 0)
            pnl_disp = pnl_str(float(pnl_v))
            lines.append("  ────────────────────────────")
            lines.append(color_text("  🔥 ACTIONABLE BET", "bold"))
            if "WIN" in str(res):
                prefix = color_text("✔ BET", "green")
            elif "LOSS" in str(res):
                prefix = color_text("✘ BET", "red")
            else:
                prefix = color_text("• BET", "bold")
            lines.append(
                f"  {prefix}  {bet_txt:<12} | SIGNAL: {ml_sig:<22} | "
                f"RESULT: {_clean_result(res):<10} | P&L: {pnl_disp}"
            )
        else:
            winner = e.get("winner") or ""
            # Non-bet outcomes use eval_status for coloring/label if snapshot exists.
            if ml_snap:
                st = _snap_status(ml_snap)
                bet_txt = str(ml_snap.get("bet") or f"{winner} ML")
                edge = ml_snap.get("edge")
                edge_txt = f" [Edge +{float(edge)*100:.1f}%]" if edge is not None else ""
                if st == "SKIPPED_EDGE":
                    prefix = color_text("⚠ SKIPPED", "yellow")
                    line = f"  {prefix}  {bet_txt:<12}{edge_txt} | SIGNAL: {ml_sig}"
                    lines.append(line)
                elif st == "NO_MODEL":
                    line = f"  ⚠ NO MODEL  {bet_txt:<12} | SIGNAL: {ml_sig}"
                    lines.append(color_text(line, "dim"))
                else:
                    line = f"  … NO EDGE   {bet_txt:<12} | SIGNAL: {ml_sig}"
                    lines.append(color_text(line, "dim"))
            else:
                line = (
                    f"  NO BET  {winner} ML{' ' * max(0, 9-len(winner))} | SIGNAL: {ml_sig:<22} | "
                    f"RESULT: WIN        | P&L: N/A"
                )
                lines.append(color_text(line, "dim"))

        # TOTAL line (always show the closing-line outcome side)
        tot = g.get("total_line")
        runs = e.get("runs")
        if tot is not None and runs is not None:
            if runs > tot:
                outcome_bet = f"OVER {tot}"
                outcome_res = "WIN"
            elif runs < tot:
                outcome_bet = f"UNDER {tot}"
                outcome_res = "WIN"
            else:
                outcome_bet = f"PUSH {tot}"
                outcome_res = "PUSH"

            tot_snap = (e.get("snapshots") or {}).get("TOTAL")
            tot_sig = _signal_display_for_snapshot(tot_snap)
            if tot_snap and _snap_status(tot_snap) == "BET":
                bet_txt = str(tot_snap.get("bet") or "")
                odds = tot_snap.get("odds_taken")
                bet = "over" if "OVER" in bet_txt.upper() else "under"
                res, pnl_v = grade_total(bet, hs, as_, tot, int(odds) if odds is not None else None)
                pnl_disp = pnl_str(float(pnl_v))
                lines.append("  ────────────────────────────")
                lines.append(color_text("  🔥 ACTIONABLE BET", "bold"))
                if "WIN" in str(res):
                    prefix = color_text("✔ BET", "green")
                elif "LOSS" in str(res):
                    prefix = color_text("✘ BET", "red")
                else:
                    prefix = color_text("• BET", "bold")
                lines.append(
                    f"  {prefix}  {bet_txt:<12} | SIGNAL: {tot_sig:<22} | "
                    f"RESULT: {_clean_result(res):<10} | P&L: {pnl_disp}"
                )
            else:
                if tot_snap:
                    st = _snap_status(tot_snap)
                    bet_txt = str(tot_snap.get("bet") or outcome_bet)
                    edge = tot_snap.get("edge")
                    edge_txt = f" [Edge +{float(edge)*100:.1f}%]" if edge is not None else ""
                    if st == "SKIPPED_EDGE":
                        prefix = color_text("⚠ SKIPPED", "yellow")
                        lines.append(f"  {prefix}  {bet_txt:<12}{edge_txt} | SIGNAL: {tot_sig}")
                    elif st == "NO_MODEL":
                        lines.append(color_text(f"  ⚠ NO MODEL  {bet_txt:<12} | SIGNAL: {tot_sig}", "dim"))
                    else:
                        lines.append(color_text(f"  … NO EDGE   {bet_txt:<12} | SIGNAL: {tot_sig}", "dim"))
                else:
                    line = (
                        f"  NO BET  {outcome_bet:<12} | SIGNAL: {tot_sig:<22} | "
                        f"RESULT: {outcome_res:<10} | P&L: N/A"
                    )
                    lines.append(color_text(line, "dim"))
        else:
            tot_sig = "No Signal"
            line = f"  NO BET  TOTAL N/A     | SIGNAL: {tot_sig:<22} | RESULT: —         | P&L: N/A"
            lines.append(color_text(line, "dim"))

        # RUNLINE line (spread) — show when data available
        hrl = g.get("home_rl")
        arl = g.get("away_rl")
        hrl_o = g.get("home_rl_odds")
        arl_o = g.get("away_rl_odds")
        hs = g.get("home_score")
        as_ = g.get("away_score")

        rl_snap = (e.get("snapshots") or {}).get("RL")
        rl_sig = _signal_display_for_snapshot(rl_snap)
        if rl_snap and _snap_status(rl_snap) == "BET":
            bet_txt = str(rl_snap.get("bet") or "")
            odds = rl_snap.get("odds_taken")
            team, line = _parse_runline_bet(bet_txt)
            res, pnl_v = grade_runline(
                team or "", float(line) if line is not None else 0.0, hs, as_,
                home_abbr=g.get("home_abbr") or "", away_abbr=g.get("away_abbr") or "",
                odds=int(odds) if odds is not None else None,
            )
            lines.append("  ────────────────────────────")
            lines.append(color_text("  🔥 ACTIONABLE BET", "bold"))
            if "WIN" in str(res):
                prefix = color_text("✔ BET", "green")
            elif "LOSS" in str(res):
                prefix = color_text("✘ BET", "red")
            else:
                prefix = color_text("• BET", "bold")
            lines.append(
                f"  {prefix}  {bet_txt:<12} | SIGNAL: {rl_sig:<22} | "
                f"RESULT: {_clean_result(res):<10} | P&L: {pnl_str(float(pnl_v))}"
            )
        elif hrl is not None and arl is not None and hs is not None and as_ is not None:
            # No signal: print the side that covered at the closing line.
            home_adj = float(hs) + float(hrl)
            away_adj = float(as_) + float(arl)
            if home_adj == float(as_) or away_adj == float(hs):
                outcome_bet = f"PUSH {g.get('home_abbr','')} {hrl:+g}"
                outcome_res = "PUSH"
            elif home_adj > float(as_):
                outcome_bet = f"{g.get('home_abbr','')} {hrl:+g}"
                outcome_res = "WIN"
            else:
                outcome_bet = f"{g.get('away_abbr','')} {arl:+g}"
                outcome_res = "WIN"
            if rl_snap:
                st = _snap_status(rl_snap)
                bet_txt = str(rl_snap.get("bet") or outcome_bet)
                edge = rl_snap.get("edge")
                edge_txt = f" [Edge +{float(edge)*100:.1f}%]" if edge is not None else ""
                if st == "SKIPPED_EDGE":
                    prefix = color_text("⚠ SKIPPED", "yellow")
                    lines.append(f"  {prefix}  {bet_txt:<12}{edge_txt} | SIGNAL: {rl_sig}")
                elif st == "NO_MODEL":
                    lines.append(color_text(f"  ⚠ NO MODEL  {bet_txt:<12} | SIGNAL: {rl_sig}", "dim"))
                else:
                    lines.append(color_text(f"  … NO EDGE   {bet_txt:<12} | SIGNAL: {rl_sig}", "dim"))
            else:
                line = (
                    f"  NO BET  {outcome_bet:<12} | SIGNAL: {rl_sig:<22} | "
                    f"RESULT: {outcome_res:<10} | P&L: N/A"
                )
                lines.append(color_text(line, "dim"))
        else:
            line = f"  NO BET  RUNLINE N/A   | SIGNAL: {rl_sig:<22} | RESULT: —         | P&L: N/A"
            lines.append(color_text(line, "dim"))

        # Snapshot-driven: no recomputed data_flags in PRIOR.
        lines.append("")

    # ════════════════════════════════════════════════════════════════════
    # BET LEDGER SUMMARY (P&L source of truth)
    # ════════════════════════════════════════════════════════════════════
    lines.append(color_text("\n📈  BET LEDGER SUMMARY", "bold") + color_text("\n" + ("─" * 72), "underline"))
    lines.append(
        color_text(
            f"\n  Staked plays: {bet_summary['bets']}   "
            f"{bet_summary['wins']}W {bet_summary['losses']}L {bet_summary['pushes']}P   "
            f"Units: {bet_summary['units']:+.2f}u   ROI: {bet_summary['roi']:.1f}%",
            "bold",
        )
    )

    # Optional grouping by market_type
    for m, rows_m in sorted(by_market.items(), key=lambda x: x[0]):
        s = _summarise_bets(rows_m)
        if s["bets"]:
            lines.append(
                f"\n  {m:<10} {s['bets']} bet(s)  "
                f"{s['wins']}W {s['losses']}L {s['pushes']}P   "
                f"Units: {s['units']:+.2f}u   ROI: {s['roi']:.1f}%"
            )

    # Season-to-date (same season as this game_date, up through game_date_et)
    try:
        season_row = conn.execute(
            "SELECT season FROM games WHERE game_date_et = ? AND game_type = 'R' LIMIT 1",
            (game_date,),
        ).fetchone()
        season_int = int(season_row[0]) if season_row and season_row[0] is not None else int(str(game_date)[:4])
    except Exception:
        season_int = int(str(game_date)[:4])

    try:
        srows = conn.execute(
            """
            SELECT
                bl.game_date, bl.game_pk, bl.market_type, bl.bet, bl.odds_taken,
                bl.stake_units, bl.signal_at_time, bl.session, bl.placed_at,
                bl.result, bl.pnl_units
            FROM bet_ledger bl
            JOIN games g ON g.game_pk = bl.game_pk
            WHERE g.season = ?
              AND g.game_type = 'R'
              AND g.game_date_et <= ?
            ORDER BY g.game_date_et, bl.placed_at
            """,
            (season_int, game_date),
        ).fetchall()
        srows = [dict(r) for r in srows]
        ssummary = _summarise_bets(srows)
        lines.append(
            f"\n  Season-to-date ({season_int} through {game_date}): {ssummary['bets']} bet(s)   "
            f"{ssummary['wins']}W {ssummary['losses']}L {ssummary['pushes']}P   "
            f"Units: {ssummary['units']:+.2f}u   ROI: {ssummary['roi']:.1f}%"
        )
    except Exception:
        pass

    ou_games = [e for e in evaluated if e["game"]["total_line"] and e["runs"] is not None]
    if ou_games:
        overs  = sum(1 for e in ou_games if e["runs"] > e["game"]["total_line"])
        unders = sum(1 for e in ou_games if e["runs"] < e["game"]["total_line"])
        pushes = len(ou_games) - overs - unders
        lines.append(
            f"\n  Slate O/U:  {overs} Over / {unders} Under"
            + (f" / {pushes} Push" if pushes else "")
            + f"  ({overs/len(ou_games):.0%} over rate)"
        )
    lines.append("")

    lines.append(CAVEAT)
    return "\n".join(lines)

def build_morning_brief(games, streaks, starters, game_date,
                        conn=None, session=None,
                        now: datetime.datetime | None = None,
                        verbose: bool = False,
                        debug_wind: bool = False):
    """
    Morning session = pipeline ``early_peek`` sneak peek: slate only.

    One section (Today's Slate): matchup, venue, start time, weather, starters
    when known, and Vegas lines (ML, total, runline). No signal evaluation,
    no picks, no streak monitor, no avoids, no signal_state persistence.
    """
    _ = (streaks, conn, session, verbose, debug_wind)  # API compatibility; unused (slate-only).
    lines = []
    if now is None:
        now = _now_et()
    generated_ts = now.strftime("%Y-%m-%d %I:%M %p ET").lstrip("0")
    lines.append(banner(f"MLB BETTING BRIEF  ·  MORNING SNEAK PEEK  ·  {game_date}"))
    lines.append(f"  Generated: {generated_ts}\n")
    lines.append(
        "\n  Today's schedule and lines only — no model signals at this run time.\n"
        "  Weather and odds reflect the latest load; re-check before you bet.\n"
    )

    lines.append(section(f"📋  TODAY'S SLATE  ({len(games)} games)"))
    if not games:
        lines.append("\n  No games on the slate for this date.\n")
    for game in games:
        lines.append(f"\n  {matchup_line(game)}")
        lines.append(f"  {weather_line(game, wind_signal_hints=False)}")
        lines.append(f"  {starter_line(game, starters)}")
        lines.append(f"  {odds_summary_line(game)}")
    lines.append("")
    lines.append(CAVEAT)
    return "\n".join(lines)


def build_primary_brief(games, streaks, starters, game_date,
                        session_label="PRIMARY", s6_fires=None,
                        conn=None, session=None, now: datetime.datetime | None = None,
                        verbose: bool = False,
                        debug_wind: bool = False):
    if s6_fires is None:
        s6_fires = {}
    _action = {
        "EARLY GAMES": "✅  EARLY GAMES ACTION WINDOW — All unplayed games. Decide before first pitch.",
        "AFTERNOON":   "✅  AFTERNOON ACTION WINDOW — All unplayed games. Decide before first pitch.",
        "PRIMARY":     "✅  PRIMARY ACTION WINDOW — All unplayed games. Make your betting decisions NOW.",
        "LATE GAMES":  "✅  LATE GAMES ACTION WINDOW — West Coast games still unstarted. Check odds and act now.",
    }.get(session_label, "✅  ACTION WINDOW — Make your betting decisions NOW.")
    lines = []
    if now is None:
        now = _now_et()
    generated_ts = now.strftime("%Y-%m-%d %I:%M %p ET").lstrip("0")
    lines.append(banner(f"MLB BETTING BRIEF  ·  {session_label} SESSION  ·  {game_date}"))
    lines.append(f"  Generated: {generated_ts}\n")
    lines.append(
        f"\n  {_action}\n"
        "  Injury news absorbed. Lineups posting. Lines near closing.\n"
        "  Confirm each line before placing. No bet above 2% of bankroll.\n"
    )
    from batch.pipeline.score_game import format_aggregate_for_brief

    evaluated_entries: list[dict] = []
    all_picks: list[dict] = []
    no_signal: list[dict] = []

    for game in games:
        sigs = evaluate_signals(
            conn, game, streaks, "primary", starters,
            verbose=verbose, debug_wind=debug_wind,
        )
        entry = {
            "game":    game,
            "sigs":    sigs,
            "starter": starter_line(game, starters),
            "streak":  streak_line(game, streaks),
        }
        evaluated_entries.append(entry)
        best_score = int(sigs.get("best_aggregate_score") or 0)
        if best_score >= 5:
            all_picks.append(entry)
        else:
            # No published "bets to avoid" — treat no-pick games (incl. internal avoid flags) as no-signal slate
            no_signal.append(entry)

    # ── Daily bet cap: keep top N by edge (stake>0) ──────────────────────
    max_bets_per_day = 5
    bet_entries: list[dict] = []
    lean_entries: list[dict] = []
    for e in all_picks:
        sg = (e.get("sigs") or {}).get("_scored_game")
        if sg is None:
            lean_entries.append(e)
            continue
        try:
            st = float(getattr(sg, "stake_multiplier", 0.0) or 0.0)
        except Exception:
            st = 0.0
        if st > 0:
            bet_entries.append(e)
        else:
            lean_entries.append(e)

    def _edge_val(entry: dict) -> float:
        sg = (entry.get("sigs") or {}).get("_scored_game")
        try:
            return float(getattr(sg, "edge", 0.0) or 0.0) if sg is not None else 0.0
        except Exception:
            return 0.0

    bet_entries.sort(key=_edge_val, reverse=True)
    kept_bets = bet_entries[:max_bets_per_day]
    skipped_bets = bet_entries[max_bets_per_day:]

    # We'll still show lean/no-bet games, but only keep the top N actual bets.
    all_picks = kept_bets + lean_entries

    # ── ACTION SUMMARY (stake > 0 only) ──────────────────────────────────
    bets_list: list[str] = []
    for e in all_picks:
        sg = (e.get("sigs") or {}).get("_scored_game")
        if sg is None:
            continue
        try:
            st = float(getattr(sg, "stake_multiplier", 0.0) or 0.0)
        except Exception:
            st = 0.0
        if st <= 0:
            continue
        # Prefer readable bet label from the scored object
        bs = getattr(getattr(sg, "top_pick", None), "bet_side", None) or getattr(sg, "best_side", None)
        bet_side = str(bs or "").strip()
        if bet_side == "away_ml":
            bet_lbl = f"{(sg.game.identifiers.away_team_abbr or 'AWAY')} ML"
        elif bet_side == "home_ml":
            bet_lbl = f"{(sg.game.identifiers.home_team_abbr or 'HOME')} ML"
        else:
            bet_lbl = bet_side.upper() if bet_side else "BET"
        bets_list.append(f"- {bet_lbl} ({st:.2f}u)")

    lines.append(section("🔥  ACTION SUMMARY"))
    lines.append(f"\n  Bets Today: {len(bets_list)}")
    lines.append(f"  Games Evaluated: {len(games)}\n")
    if bets_list:
        lines.extend(["  " + b for b in bets_list])
        lines.append("")

    if skipped_bets:
        lines.append(color_text("  (Note: additional edges existed but were skipped due to the 5-bets/day cap.)\n", "dim"))

    # Sort picks by priority (lower = higher priority)
    all_picks.sort(key=lambda e: min(p["priority"] for p in e["sigs"]["picks"]))

    # ── Persist signal state (TOP / NEXT only for customer briefs) ───────
    # Do not affect computation or report output; best-effort insert only.
    if conn is not None and session is not None:
        top_entry = all_picks[0] if len(all_picks) >= 1 else None
        next_entries = all_picks[1:6] if len(all_picks) >= 2 else []
        save_signal_state(conn, game_date, session, top_entry, next_entries, [], now=now)

    # ── Signal Tracker — intra-day pick status vs earlier sessions ────────
    # Only runs when conn is available and there were prior sessions today.
    if conn is not None and session is not None:
        prior_picks = load_todays_prior_sessions(conn, game_date, session)
        tracker_block = build_signal_tracker_block(
            prior_picks, evaluated_entries, streaks, session
        )
        if tracker_block:
            lines.append(tracker_block)

    # ── Top Pick ─────────────────────────────────────────────────────────
    lines.append(section("🔺  TOP PICK  —  Highest Probability Signal"))
    if not all_picks:
        lines.append("\n  No confirmed signals fire on today's slate.\n")
        lines.append("  Wait for tomorrow. Discipline > action.\n")
    else:
        top = all_picks[0]
        g   = top["game"]
        p   = sorted(top["sigs"]["picks"], key=lambda x: x["priority"])[0]

        lines.append(f"\n  {matchup_line(g)}")
        lines.append(f"  {weather_line(g)}")
        lines.append(f"  {top['starter']}")
        lines.append(f"  {top['streak']}")
        # Movement alert — compare vs earliest prior session pick today
        alert = movement_alert(conn, game_date, session,
                               g["game_pk"], g.get("total_line"), g.get("home_ml"))
        if alert:
            lines.append("")
            lines.append(alert)
        sg = top["sigs"].get("_scored_game")
        if sg is not None:
            lines.append("")
            for bl in format_bet_block(sg).splitlines():
                lines.append("  " + bl)
        else:
            score_txt = format_aggregate_for_brief(
                p.get("aggregate_score"), top["sigs"].get("output_tier")
            )
            lines.append(f"\n  BET: {p['bet']}  ODDS: {p['odds']}  {score_txt}")
            stack_fb = (top["sigs"].get("signal_brief") or "").strip()
            if stack_fb:
                lines.append("")
                for stack_line in stack_fb.splitlines():
                    lines.append(f"  {stack_line}")
        if p.get("alt"):
            alt = p["alt"]
            alt_score = format_aggregate_for_brief(
                alt.get("aggregate_score"), top["sigs"].get("output_tier")
            )
            lines.append(
                f"  ALT: {alt['bet']:<20}  ODDS: {alt['odds']:<8}  {alt_score}"
            )
            note = str(alt.get("note") or "")
            if note:
                lines.append(
                    f"      {textwrap.fill(note, width=66, subsequent_indent='      ')}"
                )
        if p.get("info"):
            info = p["info"]
            lines.append(f"  INFO: {info['bet']:<20}  ODDS: {info['odds']:<8}")
            note = str(info.get("note") or "")
            if note:
                lines.append(
                    f"      {textwrap.fill(note, width=66, subsequent_indent='      ')}"
                )
        lines.append(f"\n  {odds_summary_line(g)}")
        lines.append(f"\n  REASON: {textwrap.fill(p['reason'], width=66, subsequent_indent='          ')}")
        if top["sigs"]["data_flags"]:
            for f in top["sigs"]["data_flags"]:
                lines.append(f"  ⚠ DATA: {f}")
        lines.append("")

    # ── Additional Picks ─────────────────────────────────────────────────
    rest = all_picks[1:]
    lines.append(section(f"📋  ADDITIONAL MODEL SELECTIONS  ({len(rest)})"))
    if not rest:
        lines.append("\n  No additional confirmed signals today.\n")
    for i, entry in enumerate(rest, start=1):
        g    = entry["game"]
        sigs = entry["sigs"]
        best = sorted(sigs["picks"], key=lambda x: x["priority"])[0]
        lines.append(f"\n  #{i}  {matchup_line(g)}")
        lines.append(f"       {weather_line(g)}")
        lines.append(f"       {entry['starter']}")
        lines.append(f"       {entry['streak']}")
        alert = movement_alert(conn, game_date, session,
                               g["game_pk"], g.get("total_line"), g.get("home_ml"))
        if alert:
            lines.append("")
            lines.append(alert)
        sg2 = sigs.get("_scored_game")
        ind = "       "
        if sg2 is not None:
            lines.append("")
            for bl in format_bet_block(sg2).splitlines():
                lines.append(ind + bl)
        else:
            score_txt = format_aggregate_for_brief(
                best.get("aggregate_score"), sigs.get("output_tier")
            )
            lines.append(
                f"{ind}BET: {best['bet']:<20} ODDS: {best['odds']:<8} {score_txt:<7}"
            )
            stack2 = (sigs.get("signal_brief") or "").strip()
            if stack2:
                for sl in stack2.splitlines():
                    lines.append(ind + sl)
        if best.get("alt"):
            alt = best["alt"]
            alt_score = format_aggregate_for_brief(
                alt.get("aggregate_score"), sigs.get("output_tier")
            )
            lines.append(
                f"{ind}ALT: {alt['bet']:<20} ODDS: {alt['odds']:<8}  {alt_score}"
            )
            note = str(alt.get("note") or "")
            if note:
                lines.append(
                    f"            {textwrap.fill(note, width=66, subsequent_indent='            ')}"
                )
        if best.get("info"):
            info = best["info"]
            lines.append(f"{ind}INFO: {info['bet']:<19} ODDS: {info['odds']:<8}")
            note = str(info.get("note") or "")
            if note:
                lines.append(
                    f"             {textwrap.fill(note, width=66, subsequent_indent='             ')}"
                )
        lines.append(f"       {textwrap.fill(best['reason'], width=66, subsequent_indent='       ')}")
        if sigs["data_flags"]:
            for f in sigs["data_flags"]:
                lines.append(f"       ⚠ DATA: {f}")
        lines.append("")

    # ── Hot pitcher streak (separate monitor; half-stake until N≥50) ──────
    # Double-confirmation when the home team is also on a long win streak.
    lines.append(section(f"🔬  HOT PITCHER STREAK MONITOR  ({len(s6_fires)} fire(s))"))
    if not s6_fires:
        lines.append(
            "\n  No long starter win streaks (7+ starts) flagged today.\n"
        )
    else:
        lines.append(
            "\n  STATUS: Monitoring signal — +25.0% ROI on 27 fires (2018–2025, all bookmakers).\n"
            "  74% of fires are independent of long team win-streak patterns (additive, not redundant).\n"
            "  Stake: 0.5 unit until cumulative N ≥ 50. Full unit when the home team is also on a long win streak.\n"
        )
        # Build game lookup for matchup display
        game_lookup = {g["game_pk"]: g for g in games}
        for gk, fire in s6_fires.items():
            g = game_lookup.get(gk)
            if g:
                lines.append(f"  {matchup_line(g)}")
                lines.append(f"  {odds_summary_line(g)}")
            lines.append("  SIGNAL:  Hot pitcher fade (monitoring)")
            lines.append(f"  BET: {fire['bet_side'].upper()} ML  "
                         f"  STREAK: W{fire['win_streak']} ({fire['start_count']} starts this season)")
            stake_note = (
                "  ★ FULL STAKE — long home team win streak also active (double confirmation)"
                if fire["s1_also_active"]
                else f"  ◆ 0.5 UNIT STAKE — monitoring (N={fire['cumulative_n']}/{50})"
            )
            lines.append(stake_note)
            lines.append(f"\n  REASON: {textwrap.fill(fire['reason_text'], width=66, subsequent_indent='          ')}")
            lines.append("")

    # ── No-signal slate ──────────────────────────────────────────────────
    if no_signal:
        lines.append(section(f"—  NO SIGNAL  ({len(no_signal)} games — market efficient or insufficient data)"))
        for entry in no_signal:
            g = entry["game"]
            lines.append(f"  {matchup_line(g)}  |  {weather_line(g)}")
            lines.append(f"    {odds_summary_line(g)}")
            if entry["sigs"]["data_flags"]:
                for f in entry["sigs"]["data_flags"]:
                    lines.append(f"    ⚠ {f}")
        lines.append("")

    lines.append(
        "\n  ─────────────────────────────────────────────────────────────────\n"
        "  All bets assume flat-unit sizing. No bet above 2% of bankroll.\n"
        "  Confirm wind direction and speed at game time via Weather.com\n"
        "  or Windy.com. Lines shown are from last odds pull.\n"
        "  ─────────────────────────────────────────────────────────────────\n"
    )
    lines.append(CAVEAT)
    return "\n".join(lines)


def build_closing_brief(games, streaks, starters, movement, game_date,
                        conn=None, session=None,
                        now: datetime.datetime | None = None,
                        verbose: bool = False,
                        debug_wind: bool = False):
    lines = []
    if now is None:
        now = _now_et()
    generated_ts = now.strftime("%Y-%m-%d %I:%M %p ET").lstrip("0")
    lines.append(banner(f"MLB BETTING BRIEF  ·  CLOSING SESSION  ·  {game_date}"))
    lines.append(f"  Generated: {generated_ts}\n")
    lines.append(
        "\n  CLOSING CONFIRMATION — Compare against Primary Brief picks.\n"
        "  No new bets unless closing price is BETTER than Primary Brief price.\n"
        "  Flag any line that moved 3+ cents since the Primary Brief.\n"
        "  Per game, “no pick” means the model has no published bet at this price snapshot —\n"
        "  not that some unstated earlier signal expired (check Primary for that matchup).\n"
    )

    # For persistence (does not affect output)
    all_picks_entries = []

    for game in games:
        if _game_started_or_in_progress_for_closing(game, now):
            continue
        sigs = evaluate_signals(
            conn, game, streaks, "closing", starters,
            verbose=verbose, debug_wind=debug_wind,
        )
        g    = game
        lines.append(f"\n  {matchup_line(g)}")
        lines.append(f"  {weather_line(g)}")
        lines.append(f"  {starter_line(g, starters)}")
        lines.append(f"  {streak_line(g, streaks)}")
        lines.append(f"  {odds_summary_line(g)}")

        # Line movement
        mov = movement_line(g, movement)
        if mov:
            lines.append(f"  Movement: {mov}")

        # Signal hold check (NO SIGNAL is aggregated-score driven; do not gate on empty pick lists)
        best_score = int(sigs.get("best_aggregate_score") or 0)
        if best_score >= 5:
            lines.append(f"  ✅ Model pick tier at this snapshot: {sigs.get('output_tier') or 'Tier?'} (score {best_score})")
            if sigs.get("signals"):
                lines.append(f"     Signals: {', '.join(sigs['signals'])}")
            for p in (sigs.get("picks") or []):
                lines.append(f"     → {p['bet']}  {p['odds']}  ({p['reason'][:60]}…)")
        else:
            lines.append(
                "  — NO SIGNAL at these lines (best aggregated score < 5; no published bet card)."
            )

        # Persistable classifications (no change to signal logic)
        entry = {
            "game":    g,
            "sigs":    sigs,
            "starter": starter_line(g, starters),
            "streak":  streak_line(g, streaks),
        }
        if int(sigs.get("best_aggregate_score") or 0) >= 5:
            all_picks_entries.append(entry)

        # Closing-specific flag: steam or reverse line move
        gpk = game.get("game_pk")
        if gpk and gpk in movement:
            m = movement[gpk]
            if m.get("steam_move") or m.get("reverse_line_move"):
                lines.append(f"  🔥 SHARP ACTION DETECTED — market may have information.")

        if sigs["data_flags"]:
            for f in sigs["data_flags"]:
                lines.append(f"  ⚠ DATA: {f}")
        lines.append("")

    # ── Persist signal state (TOP / NEXT / AVOID) ─────────────────────────
    # Best-effort insert only; does not affect computation or report output.
    if conn is not None and session is not None:
        try:
            all_picks_entries.sort(key=lambda e: min(p["priority"] for p in e["sigs"]["picks"]))
        except Exception:
            pass
        top_entry    = all_picks_entries[0] if len(all_picks_entries) >= 1 else None
        next_entries = all_picks_entries[1:6] if len(all_picks_entries) >= 2 else []
        save_signal_state(conn, game_date, session, top_entry, next_entries, [], now=now)

    lines.append(
        "  ─────────────────────────────────────────────────────────────────\n"
        "  CLV Note: Your edge is measured against these closing prices.\n"
        "  If you bet at Primary Brief prices, compare to these lines now.\n"
        "  Run --compute-movement at 11:30 PM for full movement analysis.\n"
        "  ─────────────────────────────────────────────────────────────────\n"
    )
    lines.append(CAVEAT)
    return "\n".join(lines)



# ═══════════════════════════════════════════════════════════════════════════
# WORD (DOCX) OUTPUT
# ═══════════════════════════════════════════════════════════════════════════
# Requires: pip install python-docx
# Produces outputs/briefs/YYYY-MM-DD_SESSION.docx alongside the .txt file (default).
#
# Design:
#   · Dark accent colour (#1F3864 navy) for headers — readable on print
#   · Signal picks in a shaded table — stands out immediately
#   · Full slate as a compact 4-column table (away, home, ML, O/U)
#   · Watch flags as bold callout paragraphs (no avoid callouts in customer docs)
#   · All sessions supported: prior, morning, early, afternoon, primary, closing, late
# ═══════════════════════════════════════════════════════════════════════════

def _set_cell_bg(cell, hex_color: str):
    """Set table cell background colour (shading)."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _cell_para(cell, text: str, bold=False, size_pt=10,
               color_hex: str = "000000", align=None):
    """Clear a cell and add a single styled paragraph."""
    if align is None:
        align = WD_ALIGN_PARAGRAPH.LEFT
    cell.paragraphs[0].clear()
    p   = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold      = bold
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor(
        int(color_hex[0:2], 16),
        int(color_hex[2:4], 16),
        int(color_hex[4:6], 16),
    )
    return p


def _add_heading(doc, text: str, level: int = 1):
    """
    Banner heading (level=1) or section heading (level=2).
    Uses direct paragraph formatting so it works without style conflicts.
    """
    p   = doc.add_paragraph()
    run = p.add_run(text.upper())
    if level == 1:
        run.font.size  = Pt(16)
        run.font.bold  = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Navy shaded background via paragraph border hack
        pPr  = p._p.get_or_add_pPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "1F3864")
        pPr.append(shd)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(4)
        p.paragraph_format.left_indent  = Inches(0.1)
    else:
        run.font.size  = Pt(12)
        run.font.bold  = True
        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(2)
        # Bottom border only
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    "6")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), "1F3864")
        pBdr.append(bot)
        pPr.append(pBdr)
    return p


def _add_note(doc, text: str, italic=True, color_hex="475569"):
    """Small italic note paragraph."""
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.italic     = italic
    run.font.size  = Pt(9)
    run.font.color.rgb = RGBColor(
        int(color_hex[0:2], 16),
        int(color_hex[2:4], 16),
        int(color_hex[4:6], 16),
    )
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    return p


def _add_matchup_block(doc, game: dict, streaks: dict, starters: dict,
                       sigs: dict, show_movement: bool = False,
                       movement: dict = None, show_picks: bool = True,
                       *, include_streak: bool = True,
                       wind_signal_hints: bool = True):
    """
    Render one game block: matchup line, weather, odds, streak, starters,
    then any signals/picks/avoid/watch in a shaded pick table if relevant.
    """
    # ── Matchup header ───────────────────────────────────────────────────
    home = game.get("home_abbr", "HOME")
    away = game.get("away_abbr", "AWAY")
    venue = game.get("venue_name") or ""
    t     = _game_start_et(game)          # ET start time — always included
    start_suffix = f"  {t}" if t else ""
    matchup_txt = f"{away}  vs  {home} (h)    [{venue}]{start_suffix}"

    p   = doc.add_paragraph()
    run = p.add_run(matchup_txt)
    run.bold      = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(1)

    # ── Details line (weather | odds | [streak] | starters) ─────────────
    details = [
        weather_line(game, wind_signal_hints=wind_signal_hints),
        odds_summary_line(game),
    ]
    if include_streak:
        details.append(streak_line(game, streaks))
    details.append(starter_line(game, starters))
    for detail in details:
        p   = doc.add_paragraph()
        run = p.add_run(detail)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Inches(0.2)

    # ── Line movement (closing session) ──────────────────────────────────
    if show_movement and movement:
        mov = movement_line(game, movement)
        if mov:
            p   = doc.add_paragraph()
            run = p.add_run(f"Movement: {mov}")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x8B, 0x5C, 0xF6)
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.space_after = Pt(0)

    # ── Picks table ──────────────────────────────────────────────────────
    if sigs["picks"] and show_picks:
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Table Grid"
        tbl.autofit = False
        # Column widths: Signal | Bet | Odds | Reason  (total 9360 DXA = 6.5")
        widths = [1170, 1440, 720, 6030]   # DXA units
        for i, cell in enumerate(tbl.rows[0].cells):
            cell.width = widths[i]

        # Header row
        headers = ["SIGNAL", "BET", "ODDS", "REASON"]
        header_bg = "1F3864"
        for i, (cell, hdr) in enumerate(zip(tbl.rows[0].cells, headers)):
            _set_cell_bg(cell, header_bg)
            _cell_para(cell, hdr, bold=True, size_pt=9,
                       color_hex="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)

        for pick in sorted(sigs["picks"], key=lambda x: x["priority"]):
            row_cells = tbl.add_row().cells
            pick_bg   = "E8F5E9" if pick["market"] == "TOTAL" else "E3F2FD"
            for cell in row_cells:
                _set_cell_bg(cell, pick_bg)
            _cell_para(
                row_cells[0],
                signal_summary_for_doc(sigs) or ", ".join(sigs["signals"]),
                bold=True, size_pt=9, color_hex="1F3864",
            )
            _cell_para(row_cells[1], pick["bet"],
                       bold=True, size_pt=10, color_hex="0D47A1")
            _cell_para(row_cells[2], pick["odds"],
                       bold=True, size_pt=10, color_hex="1B5E20",
                       align=WD_ALIGN_PARAGRAPH.CENTER)
            # Wrap reason text
            reason_short = (pick["reason"][:300] + "…")                            if len(pick["reason"]) > 300 else pick["reason"]
            _cell_para(row_cells[3], reason_short, size_pt=8, color_hex="333333")

        doc.add_paragraph()  # spacer

    # (No "bets to avoid" block in customer Word output.)

    # ── Watch flag ───────────────────────────────────────────────────────
    if sigs.get("watch") and sigs.get("watch_reason"):
        p   = doc.add_paragraph()
        run = p.add_run(f"▶  WATCH: {sigs['watch_reason']}")
        run.bold      = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xE6, 0x5C, 0x00)
        p.paragraph_format.left_indent  = Inches(0.2)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)

    # ── Data flags ───────────────────────────────────────────────────────
    for flag in sigs.get("data_flags", []):
        p   = doc.add_paragraph()
        run = p.add_run(f"⚠  {flag}")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(1)


def _add_full_slate_table(doc, games: list):
    """Compact 5-column table: Time | Away | Home | ML | O/U."""
    _add_heading(doc, "Full Slate", level=2)
    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = "Table Grid"
    tbl.autofit = False
    # Total usable width ~7200 twips (page minus margins).
    # TIME col is narrowest; ML col is widest to fit odds pairs.
    widths = [900, 1200, 1200, 2400, 1200]
    for i, cell in enumerate(tbl.rows[0].cells):
        cell.width = widths[i]

    headers = ["TIME (ET)", "AWAY", "HOME", "MONEYLINE", "O/U"]
    for cell, hdr in zip(tbl.rows[0].cells, headers):
        _set_cell_bg(cell, "1F3864")
        _cell_para(cell, hdr, bold=True, size_pt=9,
                   color_hex="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)

    for game in games:
        home    = game.get("home_abbr", "HOME")
        away    = game.get("away_abbr", "AWAY")
        hml     = fmt_odds(game.get("home_ml"))
        aml     = fmt_odds(game.get("away_ml"))
        ml_txt  = f"{home} {hml} / {away} {aml}"
        tot_txt = fmt_total(game.get("total_line"))
        time_et = _game_start_et(game)   # "7:10 PM ET" or ""

        row_cells = tbl.add_row().cells
        _set_cell_bg(row_cells[0], "EEF2FF")   # light blue tint for time col
        _set_cell_bg(row_cells[1], "F5F5F5")
        _set_cell_bg(row_cells[2], "FFFFFF")
        _set_cell_bg(row_cells[3], "F5F5F5")
        _set_cell_bg(row_cells[4], "FFFFFF")
        _cell_para(row_cells[0], time_et or "—", size_pt=9, color_hex="1F3864",
                   align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_para(row_cells[1], away,            size_pt=10, color_hex="333333")
        _cell_para(row_cells[2], f"{home} (h)",   bold=True, size_pt=10, color_hex="1F3864")
        _cell_para(row_cells[3], ml_txt,           size_pt=9, color_hex="333333")
        _cell_para(row_cells[4], tot_txt,          size_pt=9, color_hex="333333",
                   align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()


def _new_brief_doc(session: str, game_date: str) -> "Document":
    """Create and configure a new Document with header/footer."""
    doc = Document()

    # Page margins — 0.75" all sides for denser layout
    for section in doc.sections:
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(0.75)
        section.right_margin  = Inches(0.75)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # Header
    hdr = doc.sections[0].header
    hdr.paragraphs[0].clear()
    htxt = f"MLB Scout  ·  {session.upper()} BRIEF  ·  {game_date}"
    run  = hdr.paragraphs[0].add_run(htxt)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    run.bold = True
    hdr.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    return doc


def build_docx_brief(
    conn: sqlite3.Connection,
    session: str,
    game_date: str,
    games: list,
    streaks: dict,
    starters: dict,
    movement: dict = None,
    verbose: bool = False,
    debug_wind: bool = False,
) -> "Document":
    """
    Build a Word Document for any session type.
    Returns a python-docx Document ready to save.
    """
    doc = _new_brief_doc(session, game_date)

    # ── Banner ────────────────────────────────────────────────────────────
    session_titles = {
        "morning":   "Morning Sneak Peek — Today's Slate",
        "early":     "Early Games Brief — 1 PM Action",
        "afternoon": "Afternoon Brief — 4 PM Action",
        "primary":   "Primary Brief — Evening Action",
        "closing":   "Closing Brief — Final Confirmation",
        "late":      "Late Games Brief — West Coast Action",
        "prior":     "Prior Day Report",
    }
    _add_heading(doc, f"MLB Betting Model  ·  {session_titles.get(session, session.upper())}  ·  {game_date}")

    # ── Session note ──────────────────────────────────────────────────────
    notes = {
        "morning":   "Slate listing only — schedule, weather, starters when known, and current lines (ML / O/U / runline). No model signals.",
        "early":     "All unplayed games on today's slate. Decide now before first pitch. Confirm lineup and weather.",
        "afternoon": "All unplayed games on today's slate. Decide now before first pitch. Confirm lineup and weather.",
        "primary":   "PRIMARY ACTION WINDOW — All unplayed games on today's slate. Make your betting decisions NOW. Confirm each line before placing.",
        "closing":   "CLOSING CONFIRMATION — Compare against Primary Brief picks. No new bets unless price is BETTER than Primary.",
        "late":      "LATE GAMES ACTION WINDOW — West Coast games still unstarted. Check odds and act now. Confirm each line before placing.",
        "prior":     "Yesterday's results. Compare against picks from your Primary Brief.",
    }
    _add_note(doc, notes.get(session, ""), italic=True, color_hex="475569")
    doc.add_paragraph()

    if not games:
        _add_note(doc, "No games found for this session.", color_hex="C62828")
        return doc

    # ── MORNING = pipeline ``early_peek`` — slate only, no signal eval ──
    if session == "morning":
        _add_heading(doc, f"Today's Slate  ({len(games)} games)", level=2)
        empty_sigs = {
            "picks": [],
            "signals": [],
            "signal_ids": [],
            "avoid": False,
            "avoid_reason": None,
            "watch": False,
            "watch_reason": None,
            "data_flags": [],
        }
        for game in games:
            _add_matchup_block(
                doc, game, streaks, starters, empty_sigs,
                show_picks=False,
                include_streak=False,
                wind_signal_hints=False,
            )
        return doc

    # ── Evaluate signals for all games ───────────────────────────────────
    entries = []
    for game in games:
        sigs = evaluate_signals(
            conn, game, streaks, session, starters,
            verbose=verbose, debug_wind=debug_wind,
        )
        entries.append({"game": game, "sigs": sigs})

    # ── PRIMARY / EARLY / AFTERNOON layout ───────────────────────────────
    if session in ("primary", "early", "afternoon", "late"):
        # NO SIGNAL policy:
        # - best_aggregate_score < 5 → No Signal
        # - otherwise → tier from score_game (Tier1/Tier2/Tier3) and must NOT be labeled No Signal
        picks_entries  = [e for e in entries if int(e["sigs"].get("best_aggregate_score") or 0) >= 5]
        nosig_entries  = [e for e in entries if int(e["sigs"].get("best_aggregate_score") or 0) < 5]
        picks_entries.sort(key=lambda e: min(p["priority"] for p in e["sigs"]["picks"]))

        # Top Pick
        _add_heading(doc, "Top Pick  —  Highest Probability Signal", level=2)
        if not picks_entries:
            _add_note(doc, "No published picks (best aggregated score < 5 across the slate). Wait for tomorrow. Discipline > action.", color_hex="C62828")
        else:
            _add_matchup_block(doc, picks_entries[0]["game"], streaks, starters,
                               picks_entries[0]["sigs"])

        # Additional picks
        rest = picks_entries[1:]
        _add_heading(doc, f"Additional Model Selections  ({len(rest)})", level=2)
        if not rest:
            _add_note(doc, "No additional confirmed signals today.", color_hex="475569")
        for e in rest:
            _add_matchup_block(doc, e["game"], streaks, starters, e["sigs"])

        # No signal
        _add_heading(doc, f"No Signal  ({len(nosig_entries)} games)", level=2)
        _add_full_slate_table(doc, [e["game"] for e in nosig_entries])

    # ── PRIOR DAY layout ──────────────────────────────────────────────────
    elif session == "prior":
        # Re-evaluate signals retroactively on completed games and grade results.
        # Mirrors build_prior_day_report() logic using the docx helper functions.

        def _grade_ml(bet_side, hs, as_, odds):
            if hs is None or as_ is None:
                return "NO RESULT", 0.0
            won = (bet_side == "home" and hs > as_) or (bet_side == "away" and as_ > hs)
            if won:
                pnl = (odds / 100.0) if odds and odds > 0 else (100.0 / abs(odds)) if odds else 0.0
                return "WIN", round(pnl, 2)
            return "LOSS", -1.0

        def _grade_total(bet, hs, as_, total, odds):
            if hs is None or as_ is None or total is None:
                return "NO RESULT", 0.0
            runs = hs + as_
            if runs == total:
                return "PUSH", 0.0
            hit = (bet == "over" and runs > total) or (bet == "under" and runs < total)
            if hit:
                pnl = (odds / 100.0) if odds and odds > 0 else (100.0 / abs(odds or 110))
                return "WIN", round(pnl, 2)
            return "LOSS", -1.0

        def _ou_label(g):
            tot  = g.get("total_line"); hs = g.get("home_score"); as_ = g.get("away_score")
            if tot is None or hs is None or as_ is None:
                return ""
            runs = hs + as_
            if runs > tot:  return f"OVER {tot} ({runs} runs)"
            if runs < tot:  return f"UNDER {tot} ({runs} runs)"
            return f"PUSH {tot} ({runs} runs)"

        def _result_color(result: str) -> str:
            return "006400" if result == "WIN" else ("C00000" if result == "LOSS" else "555555")

        def _pnl_str(pnl: float) -> str:
            return f"+{pnl:.2f}u" if pnl > 0 else (f"{pnl:.2f}u" if pnl < 0 else "push")

        # Build evaluated list
        evaluated = []
        for g in games:
            sigs = evaluate_signals(
                conn, g, streaks, "primary", starters,
                verbose=verbose, debug_wind=debug_wind,
            )
            hs   = g.get("home_score"); as_ = g.get("away_score")
            tot  = g.get("total_line"); hml = g.get("home_ml"); aml = g.get("away_ml")
            runs = (hs + as_) if (hs is not None and as_ is not None) else None

            graded = []
            for p in sigs["picks"]:
                if p["market"] == "ML":
                    side = "away" if p["bet"].startswith(g.get("away_abbr", "")) else "home"
                    odds = aml if side == "away" else hml
                    res, pnl = _grade_ml(side, hs, as_, odds or 0)
                elif p["market"] == "TOTAL":
                    bet  = "over" if "OVER" in p["bet"].upper() else "under"
                    odds = g.get("over_odds") if bet == "over" else g.get("under_odds")
                    res, pnl = _grade_total(bet, hs, as_, tot, odds)
                else:
                    res, pnl = "NO RESULT", 0.0
                graded.append({**p, "result": res, "pnl": pnl})

            winner = g.get("home_abbr", "") if (hs or 0) > (as_ or 0) else g.get("away_abbr", "")
            evaluated.append({
                "game": g, "sigs": sigs, "graded": graded,
                "ou_label": _ou_label(g), "runs": runs, "winner": winner,
            })

        # FULL SLATE: one block per game; picks show graded table; otherwise "NO SIGNAL".
        # (No separate "Bets to Avoid" section in prior report.)

        def _add_graded_pick_table(doc, graded_picks: list, sigs: dict):
            """Pick table with result and P&L columns added for prior-day grading."""
            tbl = doc.add_table(rows=1, cols=5)
            tbl.style = "Table Grid"
            tbl.autofit = False
            widths = [1080, 1260, 720, 1080, 5220]   # Signal | Bet | Odds | Result | Reason
            for i, cell in enumerate(tbl.rows[0].cells):
                cell.width = widths[i]
            headers = ["SIGNAL", "BET", "ODDS", "RESULT", "REASON"]
            for cell, hdr in zip(tbl.rows[0].cells, headers):
                _set_cell_bg(cell, "1F3864")
                _cell_para(cell, hdr, bold=True, size_pt=9,
                           color_hex="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)
            for pick in sorted(graded_picks, key=lambda x: x.get("priority", 99)):
                row_cells = tbl.add_row().cells
                result    = pick.get("result", "")
                pick_bg   = ("E2EFDA" if result == "WIN"
                             else "FCE4D6" if result == "LOSS"
                             else "FFF2CC")
                for cell in row_cells:
                    _set_cell_bg(cell, pick_bg)
                _cell_para(
                    row_cells[0],
                    signal_summary_for_doc(sigs) or ", ".join(sigs["signals"]),
                    bold=True, size_pt=9, color_hex="1F3864",
                )
                _cell_para(row_cells[1], pick["bet"],
                           bold=True, size_pt=10, color_hex="0D47A1")
                _cell_para(row_cells[2], pick.get("odds", ""),
                           bold=True, size_pt=10, color_hex="1B5E20",
                           align=WD_ALIGN_PARAGRAPH.CENTER)
                result_disp = f"{'✓' if result=='WIN' else '✗' if result=='LOSS' else '—'} {result}  {_pnl_str(pick['pnl'])}"
                _cell_para(row_cells[3], result_disp, bold=True, size_pt=9,
                           color_hex=_result_color(result),
                           align=WD_ALIGN_PARAGRAPH.CENTER)
                reason_short = (pick["reason"][:260] + "…") if len(pick["reason"]) > 260 else pick["reason"]
                _cell_para(row_cells[4], reason_short, size_pt=8, color_hex="333333")
            doc.add_paragraph()

        def _add_score_line(doc, e: dict):
            """Score and O/U outcome line for a completed game."""
            g   = e["game"]
            hs  = g.get("home_score"); as_ = g.get("away_score")
            if hs is None:
                return
            ou  = f"  |  {e['ou_label']}" if e["ou_label"] else ""
            txt = f"Final: {g.get('away_abbr','')} {as_}  –  {g.get('home_abbr','')} {hs}  ({e['winner']} wins){ou}"
            p   = doc.add_paragraph()
            run = p.add_run(txt)
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            p.paragraph_format.left_indent  = Inches(0.2)

        def _add_odds_line(doc, g: dict):
            hml = g.get("home_ml"); aml = g.get("away_ml"); tot = g.get("total_line")
            if not hml and not aml:
                return
            txt = (f"ML: {g.get('home_abbr','')} {fmt_odds(hml)} / "
                   f"{g.get('away_abbr','')} {fmt_odds(aml)}"
                   + (f"  |  O/U {tot}" if tot else ""))
            p   = doc.add_paragraph()
            run = p.add_run(txt)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(2)
            p.paragraph_format.left_indent  = Inches(0.2)

        def _clean_result(res: str) -> str:
            r = (res or "").upper()
            if "WIN" in r:
                return "WIN"
            if "LOSS" in r:
                return "LOSS"
            if "PUSH" in r:
                return "PUSH"
            return (res or "—").strip() or "—"

        def _best_pick_for_market(graded: list, market: str) -> dict | None:
            picks = [p for p in graded if (p.get("market") or "").upper() == market.upper()]
            if not picks:
                return None
            return sorted(picks, key=lambda x: x.get("priority", 99))[0]

        _add_heading(doc, f"Full Slate  ({len(evaluated)} games)", level=2)
        for e in evaluated:
            g = e["game"]
            _add_matchup_block(doc, g, streaks, starters, e["sigs"], show_picks=False)
            _add_score_line(doc, e)
            _add_odds_line(doc, g)

            best_sc = int((e.get("sigs") or {}).get("best_aggregate_score") or 0)
            sig_label = (
                ", ".join(e["sigs"].get("signals") or [])
                if best_sc >= 5
                else "No Signal"
            )

            # ML bet line
            ml_pick = _best_pick_for_market(e.get("graded") or [], "ML")
            if ml_pick:
                txt = (
                    f"BET LINE: {ml_pick.get('bet','')}  | SIGNAL: {sig_label}  | "
                    f"RESULT: {_clean_result(ml_pick.get('result'))}  | "
                    f"P&L: {('+' if float(ml_pick.get('pnl') or 0.0) > 0 else '')}{float(ml_pick.get('pnl') or 0.0):.2f}u"
                )
            else:
                txt = f"BET LINE: {e.get('winner','')} ML  | SIGNAL: No Signal  | RESULT: WIN  | P&L: N/A"
            _add_note(doc, txt, italic=False, color_hex="333333")

            # TOTAL bet line (closing outcome side)
            tot = g.get("total_line")
            runs = e.get("runs")
            if tot is not None and runs is not None:
                if runs > tot:
                    outcome_bet = f"OVER {tot}"
                    outcome_res = "WIN"
                elif runs < tot:
                    outcome_bet = f"UNDER {tot}"
                    outcome_res = "WIN"
                else:
                    outcome_bet = f"PUSH {tot}"
                    outcome_res = "PUSH"

                tot_pick = _best_pick_for_market(e.get("graded") or [], "TOTAL")
                if tot_pick:
                    txt = (
                        f"BET LINE: {tot_pick.get('bet','')}  | SIGNAL: {sig_label}  | "
                        f"RESULT: {_clean_result(tot_pick.get('result'))}  | "
                        f"P&L: {('+' if float(tot_pick.get('pnl') or 0.0) > 0 else '')}{float(tot_pick.get('pnl') or 0.0):.2f}u"
                    )
                else:
                    txt = f"BET LINE: {outcome_bet}  | SIGNAL: No Signal  | RESULT: {outcome_res}  | P&L: N/A"
            else:
                txt = "BET LINE: TOTAL N/A  | SIGNAL: No Signal  | RESULT: —  | P&L: N/A"
            _add_note(doc, txt, italic=False, color_hex="333333")

            # RUNLINE bet line
            hrl = g.get("home_rl")
            arl = g.get("away_rl")
            hs = g.get("home_score")
            as_ = g.get("away_score")
            if hrl is not None and arl is not None and hs is not None and as_ is not None:
                home_adj = float(hs) + float(hrl)
                away_adj = float(as_) + float(arl)
                if home_adj == float(as_) or away_adj == float(hs):
                    outcome_bet = f"PUSH {g.get('home_abbr','')} {hrl:+g}"
                    outcome_res = "PUSH"
                elif home_adj > float(as_):
                    outcome_bet = f"{g.get('home_abbr','')} {hrl:+g}"
                    outcome_res = "WIN"
                else:
                    outcome_bet = f"{g.get('away_abbr','')} {arl:+g}"
                    outcome_res = "WIN"
                txt = f"BET LINE: {outcome_bet}  | SIGNAL: No Signal  | RESULT: {outcome_res}  | P&L: N/A"
            else:
                txt = "BET LINE: RUNLINE N/A  | SIGNAL: No Signal  | RESULT: —  | P&L: N/A"
            _add_note(doc, txt, italic=False, color_hex="333333")

            doc.add_paragraph()

        # ── MODEL P&L SUMMARY ────────────────────────────────────────────
        _add_heading(doc, "📈  Model P&L Summary", level=2)
        all_signal_entries = [e for e in evaluated if e["graded"]]

        def _summarise(entries):
            picks = [p for e in entries for p in e["graded"]]
            if not picks: return None
            w = sum(1 for p in picks if p["pnl"] > 0)
            l = sum(1 for p in picks if p["pnl"] < 0)
            pu = sum(1 for p in picks if p["pnl"] == 0)
            return w, l, pu, sum(x["pnl"] for x in picks)

        if not _summarise(all_signal_entries):
            _add_note(doc, "No model signals fired yesterday.", color_hex="C62828")
        else:
            # Summary table
            tbl = doc.add_table(rows=1, cols=5)
            tbl.style = "Table Grid"
            tbl.autofit = False
            widths = [2400, 900, 900, 900, 2160]
            for i, cell in enumerate(tbl.rows[0].cells):
                cell.width = widths[i]
            for cell, hdr in zip(tbl.rows[0].cells, ["CATEGORY", "W", "L", "P", "NET P&L"]):
                _set_cell_bg(cell, "1F3864")
                _cell_para(cell, hdr, bold=True, size_pt=9,
                           color_hex="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)
            for label, entries in [("All signals", all_signal_entries)]:
                s = _summarise(entries)
                if not s: continue
                w, l, pu, total_pnl = s
                row_cells = tbl.add_row().cells
                bg = "E2EFDA" if total_pnl > 0 else "FCE4D6" if total_pnl < 0 else "FFFFFF"
                for cell in row_cells: _set_cell_bg(cell, bg)
                _cell_para(row_cells[0], label, bold=True, size_pt=9, color_hex="1F3864")
                _cell_para(row_cells[1], str(w), bold=True, size_pt=10, color_hex="006400", align=WD_ALIGN_PARAGRAPH.CENTER)
                _cell_para(row_cells[2], str(l), bold=True, size_pt=10, color_hex="C00000", align=WD_ALIGN_PARAGRAPH.CENTER)
                _cell_para(row_cells[3], str(pu), size_pt=10, color_hex="555555", align=WD_ALIGN_PARAGRAPH.CENTER)
                pnl_disp = f"+{total_pnl:.2f}u" if total_pnl > 0 else f"{total_pnl:.2f}u"
                _cell_para(row_cells[4], pnl_disp, bold=True, size_pt=10,
                           color_hex=_result_color("WIN" if total_pnl > 0 else "LOSS" if total_pnl < 0 else ""),
                           align=WD_ALIGN_PARAGRAPH.CENTER)
            doc.add_paragraph()

        # Slate O/U rate
        ou_games = [e for e in evaluated if e["game"].get("total_line") and e["runs"] is not None]
        if ou_games:
            overs  = sum(1 for e in ou_games if e["runs"] > e["game"]["total_line"])
            unders = sum(1 for e in ou_games if e["runs"] < e["game"]["total_line"])
            pushes = len(ou_games) - overs - unders
            ou_txt = (f"Slate O/U:  {overs} Over / {unders} Under"
                      + (f" / {pushes} Push" if pushes else "")
                      + f"  ({overs/len(ou_games):.0%} over rate)")
            _add_note(doc, ou_txt, italic=False, color_hex="333333")

        # CLV reminder
        doc.add_paragraph()
        _add_note(doc,
            "CLV Reminder: compare your bet prices vs yesterday's closing lines. "
            "Positive CLV (you beat the close) = process is working correctly. "
            "If --compute-movement was not run last night, run it now: "
            "python load_odds.py --compute-movement",
            color_hex="475569")

    # ── CLOSING layout ────────────────────────────────────────────────────
    elif session == "closing":
        _add_heading(doc, f"Closing Confirmation  —  {len(games)} Games", level=2)
        for e in entries:
            _add_matchup_block(doc, e["game"], streaks, starters, e["sigs"],
                               show_movement=True, movement=movement or {})
        _add_note(doc,
            "CLV Note: Your edge is measured against these closing prices. "
            "Run --compute-movement at 11:30 PM for full movement analysis.",
            color_hex="475569")

    # ── Footer note ───────────────────────────────────────────────────────
    doc.add_paragraph()
    _add_note(doc,
        "All bets assume flat-unit sizing. No bet above 2% of bankroll. "
        "Confirm wind and lineup at game time. Lines shown from last odds pull.",
        color_hex="64748b")

    # ── Legal caveat ─────────────────────────────────────────────────────
    doc.add_paragraph()
    caveat_tbl = doc.add_table(rows=1, cols=1)
    caveat_tbl.style = "Table Grid"
    caveat_cell = caveat_tbl.rows[0].cells[0]
    _set_cell_bg(caveat_cell, "FFF3CD")
    caveat_cell._tc.get_or_add_tcPr()

    # Heading line — bold
    heading_p = caveat_cell.paragraphs[0]
    heading_p.clear()
    heading_p.paragraph_format.space_before = Pt(4)
    heading_p.paragraph_format.space_after  = Pt(4)
    heading_p.paragraph_format.left_indent  = Inches(0.1)
    h_run = heading_p.add_run(
        "EDUCATIONAL USE ONLY \u2014 NOT FINANCIAL ADVICE"
    )
    h_run.bold            = True
    h_run.font.size       = Pt(9)
    h_run.font.color.rgb  = RGBColor(0x7B, 0x4F, 0x00)

    # Body lines
    caveat_body = (
        "This brief is produced by a personal statistical model for research "
        "and educational purposes only. Nothing in this output constitutes "
        "financial advice, investment advice, or a recommendation to place any "
        "wager of any kind. Sports betting carries substantial financial risk "
        "and is not appropriate for all individuals.\n"
        "Past signal performance does not guarantee future results. Odds, lines, "
        "and conditions can change materially between the time this brief is "
        "generated and game time. Always verify every line independently with "
        "your bookmaker before placing any bet.\n"
        "You are solely responsible for all betting decisions and any resulting "
        "financial outcomes. Never bet more than you can afford to lose. "
        "If gambling is causing financial or personal harm, contact the National "
        "Problem Gambling Helpline: 1-800-522-4700 or visit ncpgambling.org."
    )
    for line in caveat_body.split("\n"):
        body_p = caveat_cell.add_paragraph()
        body_p.paragraph_format.space_before = Pt(2)
        body_p.paragraph_format.space_after  = Pt(2)
        body_p.paragraph_format.left_indent  = Inches(0.1)
        b_run = body_p.add_run(line)
        b_run.font.size      = Pt(8)
        b_run.font.color.rgb = RGBColor(0x7B, 0x4F, 0x00)

    return doc

# ═══════════════════════════════════════════════════════════════════════════
# Main
# ═══════════════════════════════════════════════════════════════════════════

def _parse_as_of_arg(value: str) -> datetime.datetime:
    try:
        naive = datetime.datetime.strptime(value.strip(), "%Y-%m-%d %H:%M")
    except ValueError as exc:
        raise argparse.ArgumentTypeError(
            "expected 'YYYY-MM-DD HH:MM' (America/New_York wall time)"
        ) from exc
    return naive.replace(tzinfo=_ET)


def _parse_as_of_time_arg(value: str) -> datetime.datetime | datetime.time:
    """
    Parse --as-of-time as either:
    - full ET instant: YYYY-MM-DD HH:MM (same wall form as --as-of), or
    - time-only HH:MM[:SS] combined with --date (or today).
    """
    s = value.strip()
    try:
        naive = datetime.datetime.strptime(s, "%Y-%m-%d %H:%M")
        return naive.replace(tzinfo=_ET)
    except ValueError:
        pass
    for fmt in ("%H:%M", "%H:%M:%S"):
        try:
            tt = datetime.datetime.strptime(s, fmt).time()
            return datetime.time(tt.hour, tt.minute, 0, 0)
        except ValueError:
            continue
    raise argparse.ArgumentTypeError(
        "expected --as-of-time as HH:MM or full 'YYYY-MM-DD HH:MM' (America/New_York)"
    )


def _et_datetime_from_date_and_time(game_date: str, t: datetime.time) -> datetime.datetime:
    d = datetime.date.fromisoformat(game_date)
    return datetime.datetime.combine(d, t, tzinfo=_ET)


def parse_args():
    p = argparse.ArgumentParser(
        description="MLB Betting Model — Daily Brief Generator",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=textwrap.dedent("""\
            EXAMPLES
            --------
            Prior day report (after 6 AM stats pull):
              python generate_daily_brief.py --session prior

            Morning brief (after 9 AM odds pull):
              python generate_daily_brief.py --session morning

            Early games brief (after 12 PM pull — 1 PM starters):
              python generate_daily_brief.py --session early

            Afternoon brief (after 3:30 PM pull — 4 PM starters):
              python generate_daily_brief.py --session afternoon

            Primary brief (after 5 PM odds pull — evening games) — use clock for hybrid session:
              python generate_daily_brief.py --session primary --date 2026-04-17 --as-of-time 17:30

            Closing brief (after 6:30 PM odds pull):
              python generate_daily_brief.py --session closing --as-of-time 18:50

            Rerun yesterday's primary brief (recovery):
              python generate_daily_brief.py --session primary --date 2026-03-25 --as-of-time 17:30 --force

            Dry-run to preview without writing to log:
              python generate_daily_brief.py --session primary --dry-run

            Run with prereq check (recommended for scheduled runs):
              python generate_daily_brief.py --session primary --check-prereqs

            Sync bet_ledger from signal_state only (pregame T−30 window; no brief output):
              python generate_daily_brief.py --sync-bet-ledger-only --date YYYY-MM-DD

            Wind classification investigation:
              python generate_daily_brief.py --session primary --verbose --debug-wind
        """),
    )
    p.add_argument(
        "--sync-bet-ledger-only", action="store_true",
        help="Only run generate_bets_from_signal_state for --date (default today). "
             "No brief file, no brief_log duplicate check. Implies writes unless --dry-run.",
    )
    p.add_argument(
        "--session", required=False, default=None,
        choices=["prior", "morning", "early", "afternoon", "primary", "closing", "late"],
        help="Brief kind: ``prior`` and ``closing`` are always honored as-is. All other values "
             "are hints: the effective ET clock (--as-of / --as-of-time, or wall time) maps to "
             "a session via SESSION_WINDOWS (morning/early/afternoon/primary/late). If the hint "
             "differs, see [hybrid]. ``closing`` is never inferred from the clock — use it only "
             "for the dedicated closing confirmation brief. Required unless --sync-bet-ledger-only.",
    )
    p.add_argument(
        "--date", default=None,
        help="Game date to generate brief for (YYYY-MM-DD). Default: today.",
    )
    p.add_argument(
        "--force", action="store_true",
        help="Overwrite/rerun even if this session's brief was already generated today.",
    )
    p.add_argument(
        "--game-group-id",
        type=int,
        default=None,
        metavar="N",
        help="Pipeline: game_group_id for this run. Duplicate check is per (date, session, N); "
        "use with run_pipeline group_brief so each group can run and materialize bet_ledger.",
    )
    p.add_argument(
        "--dry-run", action="store_true",
        help="Print brief to console only. Do NOT write to brief_log or output file.",
    )
    p.add_argument(
        "--output", default=None,
        help="Write brief to this file path (appends). Default: outputs/briefs/brief-SLATE-DATE_RUN-STAMP_ET.txt",
    )
    p.add_argument(
        "--no-file", action="store_true",
        help="Suppress file output. Console only. Overrides --output.",
    )
    p.add_argument(
        "--warn-missing", action="store_true",
        help="Downgrade missing-data exits to warnings. Continue with partial data.",
    )
    p.add_argument(
        "--check-prereqs", action="store_true",
        help="Verify the required odds pull was completed before generating the brief.",
    )
    p.add_argument(
        "--verbose", action="store_true",
        help="Print extra DB diagnostic info during execution.",
    )
    p.add_argument(
        "--debug-wind", action="store_true",
        help="Per game, print wind classification debug: DB row vs dressed environment "
             "(direction, in/out gates, suppression, env_ceiling). Implies detailed stdout; "
             "pair with --verbose for tier lines too.",
    )
    p.add_argument(
        "--as-of",
        dest="as_of_dt",
        default=None,
        metavar="TS",
        type=_parse_as_of_arg,
        help='Full wall clock in America/New_York ("YYYY-MM-DD HH:MM") for a fixed instant — '
        "replays, tests, or scheduled runner. Overrides --as-of-time. For a normal live brief, "
        "omit; the script uses the current time in Eastern.",
    )
    p.add_argument(
        "--as-of-time",
        dest="as_of_time",
        default=None,
        metavar="HH:MM",
        type=_parse_as_of_time_arg,
        help="ET clock: HH:MM with --date, or full YYYY-MM-DD HH:MM (same as --as-of). "
        "Omit both --as-of and --as-of-time to use *current* Eastern Time (normal live run). "
        "Use a fixed time for replays, tests, or backtest. Session from SESSION_WINDOWS via ``now``.",
    )
    return p.parse_args()


def _maybe_email_report_docx(*, docx_path: str, slate_date: str, session: str) -> None:
    """
    After a Word brief is saved, optionally email it via ``delivery.email_sender``.
    Non-fatal: never raises; SMTP/import failures are logged only.
    """
    try:
        to_raw = (os.getenv("BRIEF_EMAIL_TO") or os.getenv("REPORT_EMAIL_TO") or "").strip()
        if not to_raw:
            return
        from delivery.email_sender import send_report_email

        subject = f"MLB brief {slate_date} ({session}) — {Path(docx_path).name}"
        body = (
            f"Slate date: {slate_date}\n"
            f"Session: {session}\n"
            f"Word report: {docx_path}\n"
        )
        ok, msg = send_report_email(docx_path, subject, to_raw, body=body)
        if ok:
            print(f"  ✓ [email] {msg}")
        else:
            print(f"  ⚠  [email] {msg}")
    except Exception as exc:
        print(f"  ⚠  [email] notification failed (non-fatal): {exc}")


def main():
    # ── Console encoding guard (Windows cp1252) ───────────────────────────
    # Some environments default to cp1252 and crash on unicode glyphs used
    # in headers (box drawing) and section icons. Prefer UTF-8; otherwise
    # replace unsupported characters instead of raising.
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass
    try:
        sys.stderr.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass

    args = parse_args()
    raw_session = args.session
    # Explicit clock: replay / test. If unset, "now" is current America/New_York (live run).
    explicit_as_of = args.as_of_dt is not None or getattr(args, "as_of_time", None) is not None

    # now: precedence --as-of > --as-of-time + date > wall clock (Eastern)
    if args.as_of_dt is not None:
        now = _now_et(args.as_of_dt)
    elif getattr(args, "as_of_time", None) is not None:
        tot = args.as_of_time
        if isinstance(tot, datetime.datetime):
            now = _now_et(tot)
        else:
            d_str = args.date or datetime.datetime.now(tz=_ET).date().isoformat()
            now = _et_datetime_from_date_and_time(d_str, tot)
    else:
        now = _now_et(None)

    today = args.date or now.date().isoformat()

    # In dry-run, the brief should not write to any DB-backed ledgers.
    global PERSIST_WRITES
    PERSIST_WRITES = not args.dry_run

    if explicit_as_of:
        print(f"[TIME] Running as-of: {now.strftime('%Y-%m-%d %H:%M ET')}")
    else:
        print(
            f"[TIME] no --as-of / --as-of-time: using current Eastern Time — "
            f"{now.strftime('%Y-%m-%d %H:%M ET')}"
        )
    _ggid: int | None = None
    if getattr(args, "game_group_id", None) is not None and int(args.game_group_id) > 0:
        _ggid = int(args.game_group_id)
        print(f"[TIME] game_group_id: {_ggid} (per-group duplicate scope)")

    # Validate date
    try:
        datetime.date.fromisoformat(today)
    except ValueError:
        print(f"✗  Invalid --date value: '{today}'. Use YYYY-MM-DD.")
        sys.exit(1)

    if args.sync_bet_ledger_only:
        if raw_session is not None and args.verbose:
            print("  [verbose] --session ignored when using --sync-bet-ledger-only")
        print(f"\n{'═'*72}")
        print(f"  MLB Betting Model · bet_ledger sync (signal_state) · {today}")
        print(f"{'═'*72}")
        conn = open_db(DB_PATH)
        try:
            if args.dry_run:
                print("\n  (dry-run: no DB writes; bet_ledger sync skipped.)\n")
            else:
                n_bets = generate_bets_from_signal_state(conn, today, now=now)
                if n_bets:
                    print(f"\n  ✓ bet_ledger: inserted {n_bets} row(s) from signal_state (pregame window).\n")
                else:
                    msg = (
                        "no new rows (outside 30m window, duplicate keys, "
                        "or no qualifying signal_state rows)."
                    )
                    if args.verbose:
                        print(f"\n  [verbose] bet_ledger: {msg}\n")
                    else:
                        print(f"\n  bet_ledger: {msg}\n")
        except Exception as e:
            print(f"\n  ⚠  bet_ledger sync failed: {e}\n")
            sys.exit(1)
        finally:
            conn.close()
        return

    if raw_session is None:
        print("✗  --session is required unless using --sync-bet-ledger-only.")
        sys.exit(1)

    if raw_session == "prior":
        session = "prior"
    elif raw_session == "closing":
        # Closing confirmation brief: never auto-selected from wall clock.
        session = "closing"
    else:
        session = _session_from_et_datetime(now)
        if raw_session != session:
            print(
                f"  [hybrid] resolved session={session!r} from clock "
                f"(ignoring --session {raw_session!r})"
            )

    # Slate: only pass a simulated instant into load_games when user/runner set an
    # explicit --as-of or --as-of-time (replays include Final rows, then we filter in main).
    # Live "no as-of" runs use as_of_dt=None in load_games so already-Final games stay out in SQL.
    as_of_for_slate = args.as_of_dt if args.as_of_dt is not None else (now if explicit_as_of else None)

    print(f"\n{'═'*72}")
    print(f"  MLB Betting Model · Daily Brief · {session.upper()} · {today}")
    print(f"{'═'*72}")
    if args.debug_wind:
        print(
            "\n  [debug-wind] ENABLED — printing DB vs dressed wind fields for each game "
            "(see evaluate_signals / fully_dressed_game.build_game_environment).\n"
        )

    conn = open_db(DB_PATH)

    # ── Prereq check ─────────────────────────────────────────────────────
    if args.check_prereqs:
        check_prereqs(conn, today, session)

    # ── Duplicate guard ──────────────────────────────────────────────────
    # brief_log.game_date matches the slate the brief is for: forward sessions use
    # ``today``; ``prior`` uses the previous calendar day (same key as log_brief()).
    key_date = (
        (datetime.date.fromisoformat(today) - datetime.timedelta(days=1)).isoformat()
        if session == "prior"
        else today
    )
    ensure_brief_log(conn)
    if not args.dry_run and not args.force:
        if already_ran(conn, key_date, session, game_group_id=(None if session == "prior" else _ggid)):
            gtxt = f" (game_group_id={_ggid})" if _ggid is not None else ""
            print(f"\n  ⚠  A {session} brief for {key_date}{gtxt} already exists in brief_log.")
            print(f"     Use --force to regenerate, or --dry-run to preview.\n")
            sys.exit(0)

    # ── Prior day report — uses yesterday, not today ─────────────────────
    if session == "prior":
        brief_text = build_prior_day_report(
            conn, key_date, args.verbose, now=now, debug_wind=args.debug_wind,
        )
        print(brief_text)

        # ── Save txt file ──────────────────────────────────────────────────
        output_file = None
        if not args.no_file and not args.dry_run:
            OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
            stem = _default_brief_file_stem(key_date, now, is_prior=True)
            output_file = str(OUTPUT_DIR / f"{stem}.txt")
            with open(output_file, "w", encoding="utf-8") as fh:
                fh.write(brief_text)
            print(f"\n  ✓ Prior day report saved to: {output_file}")

        # ── Word (.docx) output (default) ───────────────────────────────────
        if not args.dry_run and not args.no_file:
            if not DOCX_AVAILABLE:
                print("\n  ⚠  Word output skipped: python-docx is not installed.")
                print("     Install with: pip install python-docx")
            else:
                try:
                    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
                    docx_path = str(OUTPUT_DIR / f"{stem}.docx")
                    doc = build_docx_from_text("prior", key_date, brief_text)
                    doc.save(docx_path)
                    print(f"  ✓ Word brief saved to: {docx_path}")
                    _maybe_email_report_docx(docx_path=docx_path, slate_date=key_date, session="prior")
                except Exception as e:
                    print(f"\n  ⚠  Word output failed: {e}")
                    import traceback; traceback.print_exc()

        # ── Log ───────────────────────────────────────────────────────────
        if not args.dry_run:
            log_brief(conn, key_date, "prior", 0, 0, output_file, now=now)
        conn.close()
        print(f"\n  Done.\n")
        return

    # ── Load data — for all forward-looking sessions ──────────────────────
    games = load_games(conn, today, args.verbose, as_of_dt=as_of_for_slate)

    # Historical --as-of: DB may already mark games Final; keep slate as-of simulated time
    if as_of_for_slate is not None and games:
        now_utc = now.astimezone(datetime.timezone.utc).replace(tzinfo=None)
        grace = datetime.timedelta(minutes=10)
        games = [
            g for g in games
            if _game_start_utc_dt(g) is None
            or _game_start_utc_dt(g) > (now_utc - grace)
        ]

    if not games:
        print(f"\n  ⚠  No upcoming games found in DB for {today}.")
        print(f"     Possible causes:")
        print(f"     • Stats pull (6 AM) has not run yet")
        print(f"     • Odds pull has not run yet — run: python load_odds.py --pregame --markets game")
        print(f"     • All games for {today} are already marked Final")
        print(f"     • Games loaded as wrong game_type (e.g. 'S' for Spring Training on Opening Day)")
        print(f"       Check: SELECT game_date, game_type, COUNT(*) FROM games")
        print(f"               WHERE game_date='{today}' GROUP BY game_type;")
        print(f"       If type is 'S', re-run load_mlb_stats.py after the MLB API reclassifies (1-2 hrs).")
        print(f"     • Wrong date — check: python check_db.py")
        if as_of_for_slate is not None:
            print(
                "     • With --as-of / hybrid clock: no games were still unstarted at that simulated time, "
                "or the slate is empty for this date.\n"
            )
        else:
            print()
        sys.exit(0)

    # ── Session game filtering ────────────────────────────────────────────
    # Design: session = when you pull odds and act, NOT which games start
    # in that hour. Every action session (early/afternoon/primary) shows
    # ALL unplayed games for the date — the full pickable slate.
    #
    # The only games excluded are those that have already started or
    # finished by the time the brief is run. This means:
    #   · An all-afternoon slate (Opening Week) correctly appears in
    #     --session primary because those games haven't started yet.
    #   · West Coast late games (10 PM ET) appear in --session afternoon
    #     because they haven't started at 3:45 PM when you run it.
    #   · A 1 PM game is excluded from --session primary (run at 5:30 PM)
    #     because it already started 4+ hours ago.
    #
    # morning / closing = all games regardless (watch list / confirmation)

    if session in ("early", "afternoon", "primary") and as_of_for_slate is None:
        now_utc = now.astimezone(datetime.timezone.utc).replace(tzinfo=None)

        # Keep games that:
        #   a) have no start time (include — don't exclude on missing data), OR
        #   b) haven't started yet (start_utc > now_utc)
        # Exclude games already in progress or finished (start_utc <= now_utc)
        # Give a 10-minute grace buffer so a game that just tipped off isn't
        # immediately dropped (first pitch ≠ game over).
        GRACE_MINUTES = 10
        grace = datetime.timedelta(minutes=GRACE_MINUTES)

        filtered = [
            g for g in games
            if _game_start_utc_dt(g) is None
            or _game_start_utc_dt(g) > (now_utc - grace)
        ]

        if not filtered:
            print(f"\n  ℹ  No unplayed games remaining on today's slate for {session} session.")
            print(f"     All {len(games)} game(s) have already started or finished.")
            print(f"     Run --session closing to confirm lines, or check tomorrow's slate.\n")
            conn.close()
            return

        # Informational note if some games were filtered out as already started
        started = len(games) - len(filtered)
        if started > 0:
            started_teams = [
                f"{g['away_abbr']}@{g['home_abbr']}"
                for g in games
                if _game_start_utc_dt(g) is not None
                and _game_start_utc_dt(g) <= (now_utc - grace)
            ]
            print(f"  ℹ  {started} game(s) already started — excluded from {session} picks:")
            for t in started_teams:
                print(f"       {t}")
            print()

        games = filtered

    # Check for games missing odds — non-fatal, brief runs on games that have odds.
    missing_odds = [g for g in games if g.get("home_ml") is None]
    if missing_odds:
        missing_list = ', '.join(g['away_abbr'] + '@' + g['home_abbr']
                                 for g in missing_odds)
        with_odds = len(games) - len(missing_odds)
        print(f"\n  \u2139  {len(missing_odds)} of {len(games)} game(s) have no odds yet "
              f"\u2014 brief will run on the {with_odds} game(s) with odds.")
        print(f"     Pending odds: {missing_list}")
        print(f"     Re-run load_odds.py when lines are posted to update.\n")

    team_ids = list({g["home_team_id"] for g in games} | {g["away_team_id"] for g in games})
    streaks  = load_streaks(conn, today, team_ids, args.verbose)
    starters = load_starters(conn, today, args.verbose)
    movement = load_line_movement(conn, today, args.verbose) if session == "closing" else {}

    # ── S6 pitcher streak check (monitoring signal) ───────────────────────
    # Runs after streaks are loaded so we can pass S1-active game_pks.
    # S6 needs DB access for rolling streak computation — runs here, not
    # outside evaluate_signals() (dress + score live in enrich_game / score_game).
    s6_fires = {}
    if S6_AVAILABLE and session in ("early", "afternoon", "primary"):
        try:
            # Determine which games have S1 W7+ active (for double-confirmation flag)
            s1_w7_pks = set()
            for g in games:
                home_streak = streaks.get(g.get("home_team_id"), 0)
                if home_streak >= S6_WIN_STREAK_MIN:
                    s1_w7_pks.add(g["game_pk"])

            s6_fires = check_s6_pitcher_streak(
                conn, today,
                season=int(today[:4]),
                s1_active_game_pks=s1_w7_pks,
            )
            if s6_fires and not args.dry_run:
                for fire in s6_fires.values():
                    log_s6_fire_to_db(conn, fire, game_start_utc=None)
            if args.verbose and s6_fires:
                print(f"\n  [verbose] S6 W≥7 pitcher fade fires today: {len(s6_fires)}")
                for gk, f in s6_fires.items():
                    print(f"           game_pk={gk}  {f['pitcher_name']}  "
                          f"W{f['win_streak']}  bet={f['bet_side']}  "
                          f"stake={f['stake_recommendation']}")
        except Exception as e:
            if args.verbose:
                print(f"\n  [verbose] S6 check failed (non-fatal): {e}")

    # ── Generate brief ───────────────────────────────────────────────────
    if session == "morning":
        brief_text = build_morning_brief(
            games, streaks, starters, today,
            conn=conn, session=session, now=now,
            verbose=args.verbose, debug_wind=args.debug_wind,
        )
    elif session in ("early", "afternoon", "primary", "late"):
        label = {"early": "EARLY GAMES", "afternoon": "AFTERNOON",
                 "primary": "PRIMARY", "late": "LATE GAMES"}.get(session, "PRIMARY")
        brief_text = build_primary_brief(
            games, streaks, starters, today,
            session_label=label, s6_fires=s6_fires,
            conn=conn, session=session, now=now,
            verbose=args.verbose, debug_wind=args.debug_wind,
        )
    else:
        brief_text = build_closing_brief(
            games, streaks, starters, movement, today,
            conn=conn, session=session, now=now,
            verbose=args.verbose, debug_wind=args.debug_wind,
        )

    # ── Output ───────────────────────────────────────────────────────────
    # Console: allow ANSI. Files/Word: strip ANSI.
    print(brief_text)

    output_file = None
    default_stem = _default_brief_file_stem(today, now, is_prior=False, game_group_id=_ggid)
    if not args.no_file and not args.dry_run:
        if args.output:
            output_file = args.output
        else:
            OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
            output_file = str(OUTPUT_DIR / f"{default_stem}.txt")

        with open(output_file, "w", encoding="utf-8") as fh:
            fh.write(strip_ansi(brief_text))
        print(f"\n  ✓ Brief saved to: {output_file}")

    # ── Word output (default) ─────────────────────────────────────────────
    if not args.dry_run and not args.no_file:
        if not DOCX_AVAILABLE:
            print("\n  ⚠  Word output skipped: python-docx is not installed.")
            print("     Install with: pip install python-docx")
        else:
            try:
                OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
                docx_path = (
                    str(Path(output_file).with_suffix(".docx"))
                    if output_file
                    else str(OUTPUT_DIR / f"{default_stem}.docx")
                )
                doc = build_docx_from_text(session, today, strip_ansi(brief_text))
                doc.save(docx_path)
                print(f"  ✓ Word brief saved to: {docx_path}")
                _maybe_email_report_docx(docx_path=docx_path, slate_date=today, session=session)
            except Exception as e:
                print(f"\n  ⚠  Word output failed: {e}")
                import traceback; traceback.print_exc()

    # ── Log ──────────────────────────────────────────────────────────────
    if not args.dry_run:
        # Build pick_entries for action sessions so they can be saved
        # to brief_picks for confirmed prior-report grading.
        pick_entries_for_log = None
        if session in ("primary", "early", "afternoon", "late"):
            all_sig = []
            for g in games:
                sigs = evaluate_signals(
                    conn, g, streaks, session, starters,
                    verbose=args.verbose, debug_wind=False,
                )
                if sigs["picks"]:
                    all_sig.append({"game": g, "sigs": sigs})
            all_sig.sort(key=lambda e: min(p["priority"]
                                          for p in e["sigs"]["picks"]))
            pick_entries_for_log = all_sig
        picks_count = 0
        for g in games:
            picks_count += len(
                evaluate_signals(
                    conn, g, streaks, session, starters,
                    verbose=args.verbose, debug_wind=False,
                )["picks"]
            )
        log_brief(
            conn,
            today,
            session,
            len(games),
            picks_count,
            output_file,
            pick_entries=pick_entries_for_log,
            avoid_entries=None,
            now=now,
            game_group_id=_ggid,
        )
        if args.verbose:
            print(f"  [verbose] brief_log entry written: {today} / {session} / {picks_count} picks")

        # Materialize bet_ledger from signal_state inside the 30-minute pregame window.
        if session != "prior":
            try:
                n_bets = generate_bets_from_signal_state(conn, today, now=now)
                if n_bets:
                    print(f"  ✓ bet_ledger: inserted {n_bets} row(s) from signal_state (pregame window).")
                elif args.verbose:
                    print(
                        "  [verbose] bet_ledger: no new rows "
                        "(outside 30m window, duplicate game_pk+market_type, or no top/next signals)."
                    )
            except Exception as e:
                print(f"  ⚠  bet_ledger sync failed (non-fatal): {e}")

    conn.close()
    print(f"\n  Done.\n")


if __name__ == "__main__":
    main()
