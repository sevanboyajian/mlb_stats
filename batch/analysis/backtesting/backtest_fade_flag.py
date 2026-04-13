"""
backtest_fade_flag.py
=====================
MLB Betting Model · Fade the Flag Backtester

Tests the contrarian hypothesis: when the model issues a Watch or Avoid
signal, does betting the FLAGGED side (the team the model is warning about)
produce positive ROI?

Three sections:

  Section 1 — WATCH-Flag Games (no primary signal fires)
    A game receives a WATCH flag (home team on L4 streak — below the L5
    signal threshold). No primary signal fires. Does backing the home team
    ML (the "struggling" flagged side) outperform in these games?
    Hypothesis: the market has already over-discounted the home team's
    recent losses, and the flag catches the edge of that over-reaction zone.

  Section 2 — AVOID-Context Games (signal fires on away ML)
    A primary signal fires recommending the AWAY ML (away team).
    Does betting the HOME team ML — the team the model is avoiding — still
    produce positive ROI? And how does that compare to the model's own ROI
    on the same games?
    Hypothesis: if the home team still beats the signal on those games,
    it suggests the market over-discounts home teams in streak situations.

  Section 3 — No-Signal Slate (home dog sub-population)
    Games where no signal fires AND no watch/avoid flag is raised.
    Is there a home underdog sub-population (implied prob ≤ 40%) that
    outperforms in "quiet" games?
    Hypothesis: home dogs may be systematically underpriced when the market
    over-corrects toward the away team absent a home streak narrative.

USAGE
-----
    # Default: all seasons combined
    python backtest_fade_flag.py --from 2018-04-01 --to 2025-09-30

    # Single season
    python backtest_fade_flag.py --from 2025-04-01 --to 2025-09-30

    # Specific month
    python backtest_fade_flag.py --month 2025-09

    # Save Word report
    python backtest_fade_flag.py --from 2023-04-01 --to 2025-09-30 --docx

    # Save text report
    python backtest_fade_flag.py --from 2023-04-01 --to 2025-09-30 --output reports/fade_flag.txt

    # List available months
    python backtest_fade_flag.py --list-months

ROI MATH
--------
    Bet favourite -150 → WIN:  +$6.67 on $10 bet (100/150 × $10)
    Bet favourite -150 → LOSS: -$10.00
    Bet underdog  +140 → WIN:  +$14.00
    Bet underdog  +140 → LOSS: -$10.00
    PUSH: $0.00
"""

import argparse
import datetime
import importlib.util
import os
import sqlite3
import sys
import textwrap
from pathlib import Path

# ── Locate generate_daily_brief.py in the same folder ──────────────────────
_SCRIPT_DIR = Path(__file__).parent
_BRIEF_PATH = _SCRIPT_DIR / "generate_daily_brief.py"

DEFAULT_DB = os.path.join(os.path.dirname(os.path.abspath(__file__)), "mlb_stats.db")

# ── Constants ───────────────────────────────────────────────────────────────
WATCH_STREAK_THRESHOLD  = -4   # L4 or worse triggers Watch flag
SIGNAL_STREAK_THRESHOLD = -5   # L5 or worse triggers primary signal
STAKE                   = 10.0
PASS_ROI                = 4.5  # minimum ROI % to pass threshold
PASS_N                  = 50   # minimum fires to grade a result


# ═══════════════════════════════════════════════════════════════════════════
# Module loader
# ═══════════════════════════════════════════════════════════════════════════

def _load_brief_module():
    """Import generate_daily_brief as a module to reuse its signal engine."""
    if not _BRIEF_PATH.exists():
        print(f"\n✗  generate_daily_brief.py not found in {_SCRIPT_DIR}")
        print("   Both scripts must live in the same directory.")
        sys.exit(1)
    spec = importlib.util.spec_from_file_location("generate_daily_brief", _BRIEF_PATH)
    mod  = importlib.util.module_from_spec(spec)
    _orig    = sys.argv
    sys.argv = sys.argv[:1]
    try:
        spec.loader.exec_module(mod)
    except SystemExit:
        pass
    sys.argv = _orig
    return mod


# ═══════════════════════════════════════════════════════════════════════════
# Database helpers
# ═══════════════════════════════════════════════════════════════════════════

def get_connection(db_path: str) -> sqlite3.Connection:
    if not os.path.exists(db_path):
        print(f"\n✗  Database not found: {db_path}")
        sys.exit(1)
    con = sqlite3.connect(db_path, timeout=30)
    con.row_factory = sqlite3.Row
    return con


def get_game_dates(con, date_from: str, date_to: str) -> list:
    """Sorted list of dates with at least one Final regular-season game."""
    cur = con.execute(
        """
        SELECT DISTINCT game_date
        FROM   games
        WHERE  game_date BETWEEN ? AND ?
          AND  status    = 'Final'
          AND  game_type = 'R'
        ORDER  BY game_date
        """,
        (date_from, date_to),
    )
    return [r["game_date"] for r in cur.fetchall()]


def list_available_months(con):
    cur = con.execute(
        """
        SELECT SUBSTR(game_date,1,7) AS month,
               COUNT(DISTINCT game_date) AS game_days,
               COUNT(*) AS total_games
        FROM   games
        WHERE  status    = 'Final'
          AND  game_type = 'R'
        GROUP  BY month
        ORDER  BY month
        """
    )
    rows = cur.fetchall()
    if not rows:
        print("\n  No Final regular-season games found in DB.\n")
        return
    print(f"\n  {'MONTH':<10}  {'GAME DAYS':>10}  {'TOTAL GAMES':>12}")
    print(f"  {'─'*10}  {'─'*10}  {'─'*12}")
    for r in rows:
        print(f"  {r['month']:<10}  {r['game_days']:>10}  {r['total_games']:>12}")
    print()


def load_games_for_date(con, game_date: str) -> list:
    """Load Final regular-season games with closing-line odds for one date."""
    cur = con.execute(
        """
        SELECT
            g.game_pk,
            g.game_date,
            g.game_start_utc,
            v.name          AS venue_name,
            g.temp_f,
            g.wind_mph,
            g.wind_direction,
            g.sky_condition,
            v.wind_effect,
            v.wind_note,
            v.roof_type,
            v.elevation_ft,
            v.park_factor_runs,
            v.park_factor_hr,
            v.orientation_hp,
            th.team_id      AS home_team_id,
            th.abbreviation AS home_abbr,
            th.name         AS home_name,
            ta.team_id      AS away_team_id,
            ta.abbreviation AS away_abbr,
            ta.name         AS away_name,
            ml.home_ml,
            ml.away_ml,
            tot.total_line,
            tot.over_odds,
            tot.under_odds,
            rl.home_rl_line AS home_rl,
            rl.away_rl_line AS away_rl,
            rl.home_rl_odds,
            rl.away_rl_odds,
            g.home_score,
            g.away_score
        FROM   games g
        JOIN   teams  th ON th.team_id  = g.home_team_id
        JOIN   teams  ta ON ta.team_id  = g.away_team_id
        LEFT JOIN venues v   ON v.venue_id    = g.venue_id
        LEFT JOIN v_closing_game_odds ml  ON ml.game_pk   = g.game_pk
                                         AND ml.market_type = 'moneyline'
        LEFT JOIN v_closing_game_odds tot ON tot.game_pk  = g.game_pk
                                         AND tot.market_type = 'total'
        LEFT JOIN v_closing_game_odds rl  ON rl.game_pk   = g.game_pk
                                         AND rl.market_type = 'runline'
        WHERE  g.game_date = ?
          AND  g.status    = 'Final'
          AND  g.game_type = 'R'
        ORDER  BY g.game_start_utc
        """,
        (game_date,),
    )
    return [dict(r) for r in cur.fetchall()]


# ═══════════════════════════════════════════════════════════════════════════
# Grading + P&L helpers (mirrors backtest_top_pick.py exactly)
# ═══════════════════════════════════════════════════════════════════════════

def _implied(ml: float) -> float:
    """American odds → implied win probability (0–1)."""
    if ml is None:
        return 0.5
    ml = float(ml)
    return abs(ml) / (abs(ml) + 100) if ml < 0 else 100 / (ml + 100)


def grade_ml(game: dict, bet_home: bool) -> str:
    """Grade a moneyline bet. Returns WIN / LOSS / PUSH / NO_RESULT."""
    hs = game.get("home_score")
    as_ = game.get("away_score")
    if hs is None or as_ is None:
        return "NO_RESULT"
    if hs > as_:
        return "WIN" if bet_home else "LOSS"
    elif as_ > hs:
        return "LOSS" if bet_home else "WIN"
    else:
        return "PUSH"


def calc_pnl(odds: int | None, result: str, stake: float) -> float:
    """Calculate flat-bet P&L from American odds and result string."""
    if result in ("NO_RESULT", "PUSH"):
        return 0.0
    if result == "LOSS":
        return -stake
    # WIN
    if odds is None:
        odds = -110
    if odds > 0:
        return stake * (odds / 100.0)
    elif odds < 0:
        return stake * (100.0 / abs(odds))
    return 0.0


def streak_home(game: dict, streaks: dict) -> int:
    """
    Return the home team's current overall streak entering this game.
    Positive = win streak, negative = loss streak. Returns 0 if unavailable.

    NOTE: load_streaks() returns {team_id: int} — a flat dict of plain
    integers, NOT nested dicts. The streak is an overall (home + away)
    streak, not a home-game-only streak. There is no separate home-specific
    streak field in the signal engine.
    """
    team_id = game.get("home_team_id")
    if team_id is None or streaks is None:
        return 0
    val = streaks.get(team_id, 0)
    return int(val) if isinstance(val, (int, float)) else 0


def streak_away(game: dict, streaks: dict) -> int:
    """
    Return the away team's current overall streak entering this game.
    load_streaks() returns {team_id: int} — read directly as integer.
    """
    team_id = game.get("away_team_id")
    if team_id is None or streaks is None:
        return 0
    val = streaks.get(team_id, 0)
    return int(val) if isinstance(val, (int, float)) else 0


# ═══════════════════════════════════════════════════════════════════════════
# Section classifiers
# ═══════════════════════════════════════════════════════════════════════════

def classify_game(game: dict, streaks: dict, mod) -> str:
    """
    Classify one game into one of four buckets:

      "WATCH"    — home team on L4 streak, no primary signal fires on either side
      "AVOID"    — primary signal fires recommending AWAY ML (model avoids home)
      "QUIET"    — no signal, no watch/avoid flag
      "OTHER"    — signal fires but not an AVOID scenario (e.g. OVER pick);
                   excluded from all three sections

    Returns one of the above strings.
    """
    home_streak = streak_home(game, streaks)
    has_odds    = game.get("home_ml") is not None

    if not has_odds:
        return "NO_ODDS"

    # Evaluate signals (mirrors backtest_top_pick.py)
    try:
        sigs = mod.evaluate_signals(game, streaks, "primary")
        picks = sigs.get("picks", [])
    except Exception:
        picks = []

    if picks:
        # Check if the top pick is betting the AWAY ML (Avoid scenario)
        top_pick = sorted(picks, key=lambda p: p.get("priority", 99))[0]
        market   = top_pick.get("market", "").upper()
        bet_str  = top_pick.get("bet", "").upper()
        away_abbr = (game.get("away_abbr") or "").upper()
        if market == "ML" and away_abbr in bet_str:
            return "AVOID"
        else:
            return "OTHER"

    # No signal fires — check for Watch flag.
    # Streak is overall (home + away games) — load_streaks() has no
    # home-game-only streak. L4 overall = 4 consecutive losses any venue.
    if WATCH_STREAK_THRESHOLD <= home_streak <= -1:
        return "WATCH"

    return "QUIET"


# ═══════════════════════════════════════════════════════════════════════════
# Core backtest loop
# ═══════════════════════════════════════════════════════════════════════════

def run_backtest(con, mod, date_from: str, date_to: str, stake: float) -> dict:
    """
    Main loop. Classifies every Final regular-season game and routes it
    into one of three sections. Returns a comprehensive results dict.
    """
    game_dates = get_game_dates(con, date_from, date_to)

    if not game_dates:
        print(f"\n  ⚠  No Final regular-season games found between {date_from} and {date_to}.")
        print("     Run: python backtest_fade_flag.py --list-months\n")
        sys.exit(0)

    # Result containers for each section
    s1 = []   # Section 1: Watch-flag games, bet home ML
    s2 = []   # Section 2: Avoid-context games — track home ML AND the model's away pick
    s3 = []   # Section 3: Quiet games with home dog (implied ≤ 40%)

    total_games_seen = 0
    no_odds_count    = 0

    for game_date in game_dates:
        games = load_games_for_date(con, game_date)
        if not games:
            continue

        # Load streaks entering this date
        team_ids = list({g["home_team_id"] for g in games}
                      | {g["away_team_id"] for g in games})
        try:
            streaks = mod.load_streaks(con, game_date, team_ids, False)
        except Exception:
            streaks = {}

        for game in games:
            total_games_seen += 1
            bucket = classify_game(game, streaks, mod)

            hs      = game.get("home_score")
            as_     = game.get("away_score")
            home_ml = game.get("home_ml")
            away_ml = game.get("away_ml")
            home_streak = streak_home(game, streaks)
            total   = game.get("total_line")

            base = {
                "date":         game_date,
                "home_abbr":    game.get("home_abbr", ""),
                "away_abbr":    game.get("away_abbr", ""),
                "home_ml":      home_ml,
                "away_ml":      away_ml,
                "total_line":   total,
                "home_score":   hs,
                "away_score":   as_,
                "home_streak":  home_streak,
            }

            if bucket == "NO_ODDS":
                no_odds_count += 1
                continue

            # ── Section 1: Watch-flag games — bet the flagged home team ───
            if bucket == "WATCH":
                result = grade_ml(game, bet_home=True)
                pnl    = calc_pnl(int(home_ml) if home_ml else None, result, stake)
                s1.append({**base, "result": result, "pnl": pnl,
                            "home_impl": _implied(home_ml)})

            # ── Section 2: Avoid-context — track BOTH sides ──────────────
            elif bucket == "AVOID":
                # Contrarian bet: HOME ML (the team the model avoids)
                result_home = grade_ml(game, bet_home=True)
                pnl_home    = calc_pnl(int(home_ml) if home_ml else None, result_home, stake)
                # Model's pick: AWAY ML
                result_away = grade_ml(game, bet_home=False)
                pnl_away    = calc_pnl(int(away_ml) if away_ml else None, result_away, stake)
                s2.append({**base,
                            "result_home": result_home, "pnl_home": pnl_home,
                            "result_away": result_away, "pnl_away": pnl_away,
                            "home_impl": _implied(home_ml),
                            "away_impl": _implied(away_ml)})

            # ── Section 3: Quiet games — home dog sub-population ─────────
            elif bucket == "QUIET":
                if home_ml is None:
                    continue
                home_impl = _implied(home_ml)
                # Only capture home underdogs (positive ML, implied ≤ 40%)
                if home_ml > 0 and home_impl <= 0.40:
                    result = grade_ml(game, bet_home=True)
                    pnl    = calc_pnl(int(home_ml), result, stake)
                    s3.append({**base, "result": result, "pnl": pnl,
                                "home_impl": home_impl})

    return {
        "date_from":          date_from,
        "date_to":            date_to,
        "stake":              stake,
        "total_games_seen":   total_games_seen,
        "no_odds_count":      no_odds_count,
        "game_dates_count":   len(game_dates),
        "s1":                 s1,
        "s2":                 s2,
        "s3":                 s3,
    }


# ═══════════════════════════════════════════════════════════════════════════
# Statistics helpers
# ═══════════════════════════════════════════════════════════════════════════

def _stats(bets: list, pnl_key: str = "pnl", result_key: str = "result",
           stake: float = STAKE) -> dict:
    """Compute W/L/P/N, hit rate, staked, P&L, ROI for a list of bet records."""
    graded = [b for b in bets if b.get(result_key) not in ("NO_RESULT", None)]
    wins   = sum(1 for b in graded if b[result_key] == "WIN")
    losses = sum(1 for b in graded if b[result_key] == "LOSS")
    pushes = sum(1 for b in graded if b[result_key] == "PUSH")
    n      = wins + losses + pushes
    staked = (wins + losses) * stake   # pushes don't consume stake in ROI math
    pnl    = sum(b.get(pnl_key, 0.0) for b in graded)
    hit    = (wins / n * 100) if n > 0 else 0.0
    roi    = (pnl / staked * 100) if staked > 0 else 0.0
    return {"wins": wins, "losses": losses, "pushes": pushes, "n": n,
            "hit": hit, "staked": staked, "pnl": pnl, "roi": roi}


def _season_breakdown(bets: list, pnl_key: str = "pnl",
                      result_key: str = "result", stake: float = STAKE) -> dict:
    """Group bets by season year and return per-season stats."""
    by_season: dict = {}
    for b in bets:
        yr = b["date"][:4]
        by_season.setdefault(yr, []).append(b)
    out = {}
    for yr, group in sorted(by_season.items()):
        out[yr] = _stats(group, pnl_key, result_key, stake)
    return out


def _streak_breakdown(bets: list, pnl_key: str = "pnl",
                      result_key: str = "result", stake: float = STAKE) -> dict:
    """
    For Section 1 / Section 2, group by home streak bucket.
    Buckets: L4 only, L5, L6, L7+
    """
    buckets = {"L4": [], "L5": [], "L6": [], "L7+": []}
    for b in bets:
        s = b.get("home_streak", 0)
        if s == -4:
            buckets["L4"].append(b)
        elif s == -5:
            buckets["L5"].append(b)
        elif s == -6:
            buckets["L6"].append(b)
        elif s <= -7:
            buckets["L7+"].append(b)
    return {k: _stats(v, pnl_key, result_key, stake) for k, v in buckets.items()}


def _odds_breakdown(bets: list, home_ml_key: str = "home_ml",
                    pnl_key: str = "pnl", result_key: str = "result",
                    stake: float = STAKE) -> dict:
    """
    Group bets by home ML odds bucket:
      Heavy fav  (≤ -200)
      Fav        (-199 to -131)
      Slight fav (-130 to -101)
      Pick       (-100 to +100)
      Slight dog (+101 to +130)
      Dog        (+131 to +200)
      Big dog    (+201 and above)
    """
    def bucket(ml):
        if ml is None:
            return "Unknown"
        ml = int(ml)
        if ml <= -200: return "Heavy fav (≤-200)"
        if ml <= -131: return "Fav (-199/-131)"
        if ml <= -101: return "Slight fav (-130/-101)"
        if ml <=  100: return "Pick-em (-100/+100)"
        if ml <=  130: return "Slight dog (+101/+130)"
        if ml <=  200: return "Dog (+131/+200)"
        return "Big dog (>+200)"

    order = ["Heavy fav (≤-200)", "Fav (-199/-131)", "Slight fav (-130/-101)",
             "Pick-em (-100/+100)", "Slight dog (+101/+130)",
             "Dog (+131/+200)", "Big dog (>+200)"]

    by_bucket: dict = {k: [] for k in order}
    for b in bets:
        bk = bucket(b.get(home_ml_key))
        if bk in by_bucket:
            by_bucket[bk].append(b)

    return {k: _stats(v, pnl_key, result_key, stake)
            for k, v in by_bucket.items() if v}


# ═══════════════════════════════════════════════════════════════════════════
# Text report formatter
# ═══════════════════════════════════════════════════════════════════════════

def _verdict(st: dict) -> str:
    if st["n"] < PASS_N:
        return f"INSUFFICIENT DATA  (N={st['n']}, need ≥{PASS_N})"
    if st["roi"] >= PASS_ROI:
        return f"✓ PASS  ROI {st['roi']:+.1f}% — above +{PASS_ROI}% threshold"
    return f"✗ FAIL  ROI {st['roi']:+.1f}% — below +{PASS_ROI}% threshold"


def _stat_row(label: str, st: dict) -> str:
    pnl_str = f"+${st['pnl']:.2f}" if st['pnl'] >= 0 else f"-${abs(st['pnl']):.2f}"
    return (f"  {label:<26}  N={st['n']:>4}  W={st['wins']:>3}  L={st['losses']:>3}"
            f"  Hit={st['hit']:>5.1f}%  ROI={st['roi']:>+7.2f}%  P&L={pnl_str}")


def format_report(results: dict) -> str:
    W     = 74
    stake = results["stake"]

    def bar(ch="═"):
        return ch * W

    lines = []
    lines.append(f"\n{bar()}")
    lines.append(f"  MLB SCOUT · FADE THE FLAG BACKTESTER")
    lines.append(f"  {results['date_from']}  →  {results['date_to']}"
                 f"     Stake: ${stake:.2f}/bet")
    lines.append(bar())
    lines.append(f"  Games evaluated : {results['total_games_seen']:,}"
                 f"     Game dates : {results['game_dates_count']:,}"
                 f"     No-odds skipped : {results['no_odds_count']:,}")
    lines.append(f"  Watch-flag fires : {len(results['s1']):,}"
                 f"    Avoid-context fires : {len(results['s2']):,}"
                 f"    Quiet dog fires : {len(results['s3']):,}")
    lines.append(bar())

    # ── Section 1: Watch-flag games ────────────────────────────────────────
    s1 = results["s1"]
    st1 = _stats(s1, stake=stake)
    lines.append(f"\n{'─'*W}")
    lines.append(f"  SECTION 1 — WATCH-FLAG GAMES  (bet the flagged home team ML)")
    lines.append(f"{'─'*W}")
    lines.append(f"  Trigger: home team on L4 overall streak entering game. No primary signal.")
    lines.append(f"  Contrarian bet: HOME ML (the market-discounted, struggling home side).")
    lines.append(f"  Hypothesis: market has over-discounted the home team's recent losses.")
    lines.append(f"  Note: streak is overall (home + away games) — no home-only streak.")
    lines.append(f"")
    lines.append(f"  OVERALL")
    lines.append(_stat_row("All Watch-flag games", st1))
    lines.append(f"  Verdict: {_verdict(st1)}")
    lines.append(f"")

    # Season breakdown
    sb1 = _season_breakdown(s1, stake=stake)
    if sb1:
        lines.append(f"  SEASON-BY-SEASON STABILITY")
        lines.append(f"  {'Season':<10}  {'N':>4}  {'W':>3}  {'L':>3}  {'Hit%':>6}  {'ROI':>8}  {'P&L':>10}")
        lines.append(f"  {'─'*10}  {'─'*4}  {'─'*3}  {'─'*3}  {'─'*6}  {'─'*8}  {'─'*10}")
        for yr, st in sb1.items():
            pnl_str = f"+${st['pnl']:.2f}" if st['pnl'] >= 0 else f"-${abs(st['pnl']):.2f}"
            lines.append(f"  {yr:<10}  {st['n']:>4}  {st['wins']:>3}  {st['losses']:>3}"
                         f"  {st['hit']:>5.1f}%  {st['roi']:>+7.2f}%  {pnl_str:>10}")
        lines.append(f"")

    # Streak bucket breakdown
    stk1 = _streak_breakdown(s1, stake=stake)
    non_empty = {k: v for k, v in stk1.items() if v["n"] > 0}
    if non_empty:
        lines.append(f"  STREAK DEPTH — does a longer overall losing streak improve the contrarian edge?")
        lines.append(f"  {'Bucket':<10}  {'N':>4}  {'W':>3}  {'L':>3}  {'Hit%':>6}  {'ROI':>8}  {'P&L':>10}")
        lines.append(f"  {'─'*10}  {'─'*4}  {'─'*3}  {'─'*3}  {'─'*6}  {'─'*8}  {'─'*10}")
        for bkt, st in non_empty.items():
            pnl_str = f"+${st['pnl']:.2f}" if st['pnl'] >= 0 else f"-${abs(st['pnl']):.2f}"
            lines.append(f"  {bkt:<10}  {st['n']:>4}  {st['wins']:>3}  {st['losses']:>3}"
                         f"  {st['hit']:>5.1f}%  {st['roi']:>+7.2f}%  {pnl_str:>10}")
        lines.append(f"")

    # Odds profile
    op1 = _odds_breakdown(s1, stake=stake)
    if op1:
        lines.append(f"  ODDS PROFILE — which price range drives results?")
        lines.append(f"  {'Bucket':<28}  {'N':>4}  {'Hit%':>6}  {'ROI':>8}  {'P&L':>10}")
        lines.append(f"  {'─'*28}  {'─'*4}  {'─'*6}  {'─'*8}  {'─'*10}")
        for bkt, st in op1.items():
            pnl_str = f"+${st['pnl']:.2f}" if st['pnl'] >= 0 else f"-${abs(st['pnl']):.2f}"
            lines.append(f"  {bkt:<28}  {st['n']:>4}  {st['hit']:>5.1f}%  "
                         f"{st['roi']:>+7.2f}%  {pnl_str:>10}")
        lines.append(f"")

    lines.append(f"  INTERPRETATION GUIDE — SECTION 1")
    lines.append(f"  ROI ≥ +{PASS_ROI}% on N ≥ {PASS_N}: contrarian Watch-flag home bet is a real edge.")
    lines.append(f"  Streak depth shows longer streak = stronger signal if ROI increases with depth.")
    lines.append(f"  Odds profile shows if edge concentrates in a specific pricing band.")

    # ── Section 2: Avoid-context games ────────────────────────────────────
    s2      = results["s2"]
    st2h    = _stats(s2, pnl_key="pnl_home", result_key="result_home", stake=stake)
    st2a    = _stats(s2, pnl_key="pnl_away", result_key="result_away", stake=stake)

    lines.append(f"\n{'─'*W}")
    lines.append(f"  SECTION 2 — AVOID-CONTEXT GAMES  (signal fires: away ML recommended)")
    lines.append(f"{'─'*W}")
    lines.append(f"  Trigger: primary model signal fires recommending AWAY ML.")
    lines.append(f"  Two parallel tracks: (A) contrarian home ML  vs  (B) model away ML.")
    lines.append(f"  Hypothesis: model identifies over-priced home teams. Does backing")
    lines.append(f"  the home side still show positive ROI in this population?")
    lines.append(f"")
    lines.append(f"  THREE-WAY COMPARISON  (same {len(s2):,} games)")
    lines.append(_stat_row("A — Contrarian (home ML)", st2h))
    lines.append(f"      Verdict: {_verdict(st2h)}")
    lines.append(_stat_row("B — Model pick (away ML)", st2a))
    lines.append(f"      Verdict: {_verdict(st2a)}")
    if st2a["n"] > 0 and st2h["n"] > 0:
        diff = st2h["roi"] - st2a["roi"]
        lines.append(f"")
        lines.append(f"  ROI gap (contrarian minus model): {diff:+.2f}pp")
        if diff >= 0:
            lines.append(f"  → Contrarian matches or beats the model on these games.")
            lines.append(f"    Signals may be identifying over-priced home teams but")
            lines.append(f"    both sides show positive expected value.")
        else:
            lines.append(f"  → Model outperforms contrarian by {abs(diff):.2f}pp ROI.")
            lines.append(f"    The signal is correctly identifying value on the away side.")
    lines.append(f"")

    # Season breakdown for Avoid (model side only — most informative)
    sb2 = _season_breakdown(s2, pnl_key="pnl_away", result_key="result_away", stake=stake)
    if sb2:
        lines.append(f"  MODEL PICK (AWAY ML) — SEASON-BY-SEASON STABILITY")
        lines.append(f"  {'Season':<10}  {'N':>4}  {'W':>3}  {'L':>3}  {'Hit%':>6}  {'ROI':>8}  {'P&L':>10}")
        lines.append(f"  {'─'*10}  {'─'*4}  {'─'*3}  {'─'*3}  {'─'*6}  {'─'*8}  {'─'*10}")
        for yr, st in sb2.items():
            pnl_str = f"+${st['pnl']:.2f}" if st['pnl'] >= 0 else f"-${abs(st['pnl']):.2f}"
            lines.append(f"  {yr:<10}  {st['n']:>4}  {st['wins']:>3}  {st['losses']:>3}"
                         f"  {st['hit']:>5.1f}%  {st['roi']:>+7.2f}%  {pnl_str:>10}")
        lines.append(f"")

    # Streak breakdown for Avoid (home side)
    stk2 = _streak_breakdown(s2, pnl_key="pnl_home", result_key="result_home", stake=stake)
    non_empty2 = {k: v for k, v in stk2.items() if v["n"] > 0}
    if non_empty2:
        lines.append(f"  CONTRARIAN (HOME ML) — STREAK DEPTH BREAKDOWN")
        lines.append(f"  {'Bucket':<10}  {'N':>4}  {'W':>3}  {'L':>3}  {'Hit%':>6}  {'ROI':>8}")
        lines.append(f"  {'─'*10}  {'─'*4}  {'─'*3}  {'─'*3}  {'─'*6}  {'─'*8}")
        for bkt, st in non_empty2.items():
            lines.append(f"  {bkt:<10}  {st['n']:>4}  {st['wins']:>3}  {st['losses']:>3}"
                         f"  {st['hit']:>5.1f}%  {st['roi']:>+7.2f}%")
        lines.append(f"")

    lines.append(f"  INTERPRETATION GUIDE — SECTION 2")
    lines.append(f"  If contrarian ROI > 0 but < model ROI: home still has some residual value,")
    lines.append(f"  but the signal correctly identifies the away team as the better bet.")
    lines.append(f"  If contrarian ROI > model ROI: signals may be too aggressive fading home.")
    lines.append(f"  If contrarian ROI < 0: model is correctly identifying overpriced home teams.")

    # ── Section 3: Quiet home dogs ─────────────────────────────────────────
    s3  = results["s3"]
    st3 = _stats(s3, stake=stake)
    lines.append(f"\n{'─'*W}")
    lines.append(f"  SECTION 3 — QUIET SLATE HOME UNDERDOGS  (no signal, no flag)")
    lines.append(f"{'─'*W}")
    lines.append(f"  Trigger: no signal fires, no Watch/Avoid flag raised.")
    lines.append(f"  Bet: home team ML where home_ml > 0 AND implied prob ≤ 40%.")
    lines.append(f"  Hypothesis: home dogs are systematically underpriced when the market")
    lines.append(f"  has no home-team narrative to rally behind.")
    lines.append(f"")
    lines.append(f"  OVERALL")
    lines.append(_stat_row("Quiet home dogs (impl ≤40%)", st3))
    lines.append(f"  Verdict: {_verdict(st3)}")
    lines.append(f"")

    sb3 = _season_breakdown(s3, stake=stake)
    if sb3:
        lines.append(f"  SEASON-BY-SEASON STABILITY")
        lines.append(f"  {'Season':<10}  {'N':>4}  {'W':>3}  {'L':>3}  {'Hit%':>6}  {'ROI':>8}  {'P&L':>10}")
        lines.append(f"  {'─'*10}  {'─'*4}  {'─'*3}  {'─'*3}  {'─'*6}  {'─'*8}  {'─'*10}")
        for yr, st in sb3.items():
            pnl_str = f"+${st['pnl']:.2f}" if st['pnl'] >= 0 else f"-${abs(st['pnl']):.2f}"
            lines.append(f"  {yr:<10}  {st['n']:>4}  {st['wins']:>3}  {st['losses']:>3}"
                         f"  {st['hit']:>5.1f}%  {st['roi']:>+7.2f}%  {pnl_str:>10}")
        lines.append(f"")

    op3 = _odds_breakdown(s3, stake=stake)
    if op3:
        lines.append(f"  ODDS PROFILE — home ML price distribution in quiet dog population")
        lines.append(f"  {'Bucket':<28}  {'N':>4}  {'Hit%':>6}  {'ROI':>8}  {'P&L':>10}")
        lines.append(f"  {'─'*28}  {'─'*4}  {'─'*6}  {'─'*8}  {'─'*10}")
        for bkt, st in op3.items():
            pnl_str = f"+${st['pnl']:.2f}" if st['pnl'] >= 0 else f"-${abs(st['pnl']):.2f}"
            lines.append(f"  {bkt:<28}  {st['n']:>4}  {st['hit']:>5.1f}%  "
                         f"{st['roi']:>+7.2f}%  {pnl_str:>10}")
        lines.append(f"")

    lines.append(f"  INTERPRETATION GUIDE — SECTION 3")
    lines.append(f"  ROI ≥ +{PASS_ROI}% here suggests that the 'nothing fires' population still")
    lines.append(f"  contains exploitable home-underdog pricing. If confirmed, could become")
    lines.append(f"  an H1a-style educational pick filter that earns positive ROI.")

    # ── Decision matrix ────────────────────────────────────────────────────
    lines.append(f"\n{'─'*W}")
    lines.append(f"  OVERALL DECISION MATRIX")
    lines.append(f"{'─'*W}")
    lines.append(f"  {'Section':<32}  {'ROI':>8}  {'N':>5}  {'Verdict'}")
    lines.append(f"  {'─'*32}  {'─'*8}  {'─'*5}  {'─'*30}")

    for label, st in [
        ("S1 — Watch-flag home bet",        st1),
        ("S2 — Contrarian avoid-ctx home",  st2h),
        ("S2 — Model pick (away ML)",       st2a),
        ("S3 — Quiet home dog (impl ≤40%)", st3),
    ]:
        verdict = ("PASS ✓" if st["roi"] >= PASS_ROI and st["n"] >= PASS_N
                   else ("LOW N" if st["n"] < PASS_N else "FAIL ✗"))
        lines.append(f"  {label:<32}  {st['roi']:>+7.2f}%  {st['n']:>5}  {verdict}")

    lines.append(f"")
    lines.append(f"  Pass threshold: ROI ≥ +{PASS_ROI}% on N ≥ {PASS_N} graded bets.")
    lines.append(f"  Confirm with OOS validation before adding any signal to live model.")

    # ── Notes ──────────────────────────────────────────────────────────────
    lines.append(f"\n{'─'*W}")
    lines.append(f"  NOTES")
    lines.append(f"{'─'*W}")
    lines.append(f"  · All odds are closing-line prices from mlb_stats.db.")
    lines.append(f"  · No opening odds comparison (CLV) included in this backtest.")
    lines.append(f"    Run backtest_top_pick.py to evaluate CLV on individual signal fires.")
    lines.append(f"  · ROI is after-vig (no vig removal). +{PASS_ROI}% clears vig on flat -110.")
    lines.append(f"  · Watch threshold: overall team streak ≤ L4 (signal fires at L5+).")
    lines.append(f"    Streak counts all games (home + away) — no home-only streak.")
    lines.append(f"  · Avoid threshold: any fire where model recommends away ML.")
    lines.append(f"  · Quiet home dog: implied prob ≤ 40% (home_ml ≥ +150 approx).")
    lines.append(f"  · Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')} ET")
    lines.append(f"{'─'*W}\n")

    return "\n".join(lines)


# ═══════════════════════════════════════════════════════════════════════════
# Word (.docx) report builder
# ═══════════════════════════════════════════════════════════════════════════

def build_docx_report(results: dict, out_path: str, stake: float):
    """
    Generate a formatted Word report for the Fade the Flag backtest results.
    Requires: pip install python-docx
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        print("  ⚠  python-docx not installed. Run: pip install python-docx")
        print("     Word report skipped — text report still written.")
        return False

    NAVY   = RGBColor(0x1F, 0x38, 0x64)
    WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
    GREEN  = RGBColor(0x00, 0x70, 0x00)
    RED    = RGBColor(0xC0, 0x00, 0x00)
    AMBER  = RGBColor(0xBF, 0x8F, 0x00)
    GRAY   = RGBColor(0x40, 0x40, 0x40)

    NAVY_HEX  = "1F3864"
    LTBLUE    = "D9E2F3"
    LTGRAY    = "F2F2F2"
    PASS_CLR  = "E2EFDA"   # light green
    FAIL_CLR  = "FCE4D6"   # light red / salmon
    WARN_CLR  = "FFF2CC"   # amber

    def _set_cell_bg(cell, hex_color: str):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color)
        tcPr.append(shd)

    def _cell_para(cell, text: str, bold: bool = False, italic: bool = False,
                   align: str = "left", color: RGBColor = None, size_pt: int = 10):
        p   = cell.paragraphs[0]
        p.alignment = (WD_ALIGN_PARAGRAPH.CENTER if align == "center"
                       else WD_ALIGN_PARAGRAPH.RIGHT if align == "right"
                       else WD_ALIGN_PARAGRAPH.LEFT)
        run = p.add_run(str(text))
        run.bold   = bold
        run.italic = italic
        run.font.size = Pt(size_pt)
        if color:
            run.font.color.rgb = color
        return run

    def _add_heading(doc, text: str, level: int = 1):
        p   = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        if level == 1:
            run.font.size = Pt(14)
            run.font.color.rgb = WHITE
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            # Navy shaded banner via paragraph shading
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"),   "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"),  NAVY_HEX)
            pPr.append(shd)
            p.paragraph_format.left_indent  = Inches(0.1)
        else:
            run.font.size = Pt(11)
            run.font.color.rgb = NAVY
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(2)
            # Navy bottom border
            pPr  = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bot  = OxmlElement("w:bottom")
            bot.set(qn("w:val"),   "single")
            bot.set(qn("w:sz"),    "6")
            bot.set(qn("w:color"), NAVY_HEX)
            pBdr.append(bot)
            pPr.append(pBdr)
        return p

    def _add_body(doc, text: str, color: RGBColor = None, italic: bool = False,
                  size_pt: int = 10, space_after: int = 2):
        p   = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(size_pt)
        run.italic    = italic
        if color:
            run.font.color.rgb = color
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(space_after)
        return p

    def _add_stat_table(doc, headers: list, rows: list,
                        col_widths_inches: list, header_bg: str = NAVY_HEX):
        """
        Add a data table. headers = list of strings. rows = list of lists.
        col_widths_inches: width of each column in inches.
        Alternating row shading for readability.
        """
        n_cols = len(headers)
        table  = doc.add_table(rows=1 + len(rows), cols=n_cols)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        # Set column widths
        for i, cell in enumerate(table.rows[0].cells):
            cell.width = Inches(col_widths_inches[i] if i < len(col_widths_inches)
                                else col_widths_inches[-1])

        # Header row
        for i, hdr in enumerate(headers):
            cell = table.rows[0].cells[i]
            _set_cell_bg(cell, header_bg)
            _cell_para(cell, hdr, bold=True, color=WHITE,
                       align="center" if i > 0 else "left", size_pt=9)

        # Data rows
        for r_idx, row_data in enumerate(rows):
            bg = LTGRAY if r_idx % 2 == 0 else "FFFFFF"
            for c_idx, val in enumerate(row_data):
                cell = table.rows[r_idx + 1].cells[c_idx]
                _set_cell_bg(cell, bg)
                # Color-code ROI and verdict columns
                col_color = None
                val_str   = str(val)
                if c_idx > 0 and isinstance(val, str):
                    if val.startswith("+") and "%" in val:
                        col_color = GREEN
                    elif val.startswith("-") and "%" in val:
                        col_color = RED
                    elif "PASS" in val:
                        col_color = GREEN
                    elif "FAIL" in val:
                        col_color = RED
                _cell_para(cell, val_str, color=col_color,
                           align="center" if c_idx > 0 else "left", size_pt=9)

        doc.add_paragraph()  # spacing after table

    def _verdict_color(st: dict) -> str:
        if st["n"] < PASS_N:
            return WARN_CLR
        return PASS_CLR if st["roi"] >= PASS_ROI else FAIL_CLR

    def _verdict_text(st: dict) -> str:
        if st["n"] < PASS_N:
            return f"INSUFFICIENT DATA (N={st['n']}, need ≥{PASS_N})"
        if st["roi"] >= PASS_ROI:
            return f"PASS — ROI {st['roi']:+.1f}% exceeds +{PASS_ROI}% threshold"
        return f"FAIL — ROI {st['roi']:+.1f}% below +{PASS_ROI}% threshold"

    def _add_verdict_box(doc, st: dict, label: str):
        """Add a shaded verdict paragraph."""
        bg   = _verdict_color(st)
        text = f"  {label}: {_verdict_text(st)}"
        p    = doc.add_paragraph()
        run  = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10)
        pPr  = p._p.get_or_add_pPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  bg)
        pPr.append(shd)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(8)

    # ── Build document ─────────────────────────────────────────────────────
    doc = Document()

    # Page setup: US Letter, 0.75" margins
    from docx.oxml import OxmlElement as oxe
    section_props = doc.sections[0]
    section_props.page_width  = Inches(8.5)
    section_props.page_height = Inches(11.0)
    section_props.left_margin   = Inches(0.75)
    section_props.right_margin  = Inches(0.75)
    section_props.top_margin    = Inches(0.75)
    section_props.bottom_margin = Inches(0.75)

    # Running header
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    header = doc.sections[0].header
    hdr_p  = header.paragraphs[0]
    hdr_p.clear()
    hdr_run = hdr_p.add_run(
        f"MLB Scout · Fade the Flag Backtester  ·  "
        f"{results['date_from']} – {results['date_to']}"
    )
    hdr_run.font.size = Pt(8)
    hdr_run.font.color.rgb = GRAY
    hdr_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # ── Title block ────────────────────────────────────────────────────────
    _add_heading(doc, "  MLB SCOUT — FADE THE FLAG BACKTESTER", level=1)
    _add_body(doc, f"Date range: {results['date_from']}  →  {results['date_to']}   |   "
                   f"Stake: ${stake:.2f}/bet   |   Pass threshold: ROI ≥ +{PASS_ROI}% on N ≥ {PASS_N}",
              color=GRAY, size_pt=9, space_after=4)
    _add_body(doc, f"Games evaluated: {results['total_games_seen']:,}   |   "
                   f"Game dates: {results['game_dates_count']:,}   |   "
                   f"No-odds skipped: {results['no_odds_count']:,}",
              color=GRAY, size_pt=9, space_after=4)
    _add_body(doc, f"Watch-flag fires: {len(results['s1']):,}   |   "
                   f"Avoid-context fires: {len(results['s2']):,}   |   "
                   f"Quiet dog fires: {len(results['s3']):,}",
              color=GRAY, size_pt=9, space_after=8)

    # ── Decision matrix (at top for quick reference) ───────────────────────
    s1, s2, s3 = results["s1"], results["s2"], results["s3"]
    st1  = _stats(s1, stake=stake)
    st2h = _stats(s2, pnl_key="pnl_home", result_key="result_home", stake=stake)
    st2a = _stats(s2, pnl_key="pnl_away", result_key="result_away", stake=stake)
    st3  = _stats(s3, stake=stake)

    _add_heading(doc, "Decision Matrix — Quick Reference", level=2)
    _add_stat_table(
        doc,
        headers=["Section", "N", "W", "L", "Hit %", "ROI", "Verdict"],
        rows=[
            ["S1 — Watch-flag home bet",
             st1["n"], st1["wins"], st1["losses"],
             f"{st1['hit']:.1f}%", f"{st1['roi']:+.2f}%",
             "PASS ✓" if st1["roi"] >= PASS_ROI and st1["n"] >= PASS_N
             else ("LOW N" if st1["n"] < PASS_N else "FAIL ✗")],
            ["S2 — Contrarian avoid-ctx (home)",
             st2h["n"], st2h["wins"], st2h["losses"],
             f"{st2h['hit']:.1f}%", f"{st2h['roi']:+.2f}%",
             "PASS ✓" if st2h["roi"] >= PASS_ROI and st2h["n"] >= PASS_N
             else ("LOW N" if st2h["n"] < PASS_N else "FAIL ✗")],
            ["S2 — Model pick (away ML)",
             st2a["n"], st2a["wins"], st2a["losses"],
             f"{st2a['hit']:.1f}%", f"{st2a['roi']:+.2f}%",
             "PASS ✓" if st2a["roi"] >= PASS_ROI and st2a["n"] >= PASS_N
             else ("LOW N" if st2a["n"] < PASS_N else "FAIL ✗")],
            ["S3 — Quiet home dog (impl ≤40%)",
             st3["n"], st3["wins"], st3["losses"],
             f"{st3['hit']:.1f}%", f"{st3['roi']:+.2f}%",
             "PASS ✓" if st3["roi"] >= PASS_ROI and st3["n"] >= PASS_N
             else ("LOW N" if st3["n"] < PASS_N else "FAIL ✗")],
        ],
        col_widths_inches=[2.8, 0.5, 0.5, 0.5, 0.7, 0.7, 1.1],
    )

    # ── Section 1 ──────────────────────────────────────────────────────────
    _add_heading(doc, "Section 1 — Watch-Flag Games (Bet the Flagged Home Team ML)", level=2)
    _add_body(doc,
              "Trigger: home team on L4 overall losing streak entering the game (home + away games), "
              "no primary signal fires. Contrarian bet: HOME ML. Hypothesis: the market has over-discounted "
              "the home team's recent losses. Note: streak is overall, not home-game-only.",
              color=GRAY, size_pt=9, space_after=6)
    _add_verdict_box(doc, st1, "Section 1 — Watch-flag home bet")

    # Overall row
    _add_heading(doc, "Overall Performance", level=2)
    _add_stat_table(
        doc,
        headers=["Population", "N", "W", "L", "Hit %", "ROI", "Net P&L"],
        rows=[[
            "All Watch-flag games (home ML)",
            st1["n"], st1["wins"], st1["losses"],
            f"{st1['hit']:.1f}%", f"{st1['roi']:+.2f}%",
            f"+${st1['pnl']:.2f}" if st1["pnl"] >= 0 else f"-${abs(st1['pnl']):.2f}",
        ]],
        col_widths_inches=[2.8, 0.5, 0.5, 0.5, 0.7, 0.7, 1.1],
    )

    # Season breakdown
    sb1 = _season_breakdown(s1, stake=stake)
    if sb1:
        _add_heading(doc, "Season-by-Season Stability", level=2)
        rows = []
        for yr, st in sb1.items():
            pnl_str = f"+${st['pnl']:.2f}" if st["pnl"] >= 0 else f"-${abs(st['pnl']):.2f}"
            rows.append([yr, st["n"], st["wins"], st["losses"],
                         f"{st['hit']:.1f}%", f"{st['roi']:+.2f}%", pnl_str])
        _add_stat_table(doc,
            headers=["Season", "N", "W", "L", "Hit %", "ROI", "Net P&L"],
            rows=rows,
            col_widths_inches=[1.0, 0.5, 0.5, 0.5, 0.7, 0.8, 1.0])

    # Streak bucket
    stk1 = _streak_breakdown(s1, stake=stake)
    non_empty1 = {k: v for k, v in stk1.items() if v["n"] > 0}
    if non_empty1:
        _add_heading(doc, "Streak Depth — Does a Longer Losing Streak Improve the Edge?", level=2)
        rows = []
        for bkt, st in non_empty1.items():
            pnl_str = f"+${st['pnl']:.2f}" if st["pnl"] >= 0 else f"-${abs(st['pnl']):.2f}"
            rows.append([bkt, st["n"], st["wins"], st["losses"],
                         f"{st['hit']:.1f}%", f"{st['roi']:+.2f}%", pnl_str])
        _add_stat_table(doc,
            headers=["Streak Bucket", "N", "W", "L", "Hit %", "ROI", "Net P&L"],
            rows=rows,
            col_widths_inches=[1.2, 0.5, 0.5, 0.5, 0.7, 0.8, 1.0])

    # Odds profile
    op1 = _odds_breakdown(s1, stake=stake)
    if op1:
        _add_heading(doc, "Odds Profile — Which Price Range Drives Results?", level=2)
        rows = []
        for bkt, st in op1.items():
            pnl_str = f"+${st['pnl']:.2f}" if st["pnl"] >= 0 else f"-${abs(st['pnl']):.2f}"
            rows.append([bkt, st["n"], f"{st['hit']:.1f}%", f"{st['roi']:+.2f}%", pnl_str])
        _add_stat_table(doc,
            headers=["Home ML Bucket", "N", "Hit %", "ROI", "Net P&L"],
            rows=rows,
            col_widths_inches=[2.8, 0.6, 0.8, 0.8, 1.2])

    # ── Section 2 ──────────────────────────────────────────────────────────
    doc.add_page_break()
    _add_heading(doc, "Section 2 — Avoid-Context Games (Signal Fires: Away ML Recommended)", level=2)
    _add_body(doc,
              "Trigger: a primary model signal fires recommending the AWAY ML. "
              "Two parallel tracks are tracked on the same set of games: "
              "(A) contrarian — bet the HOME ML anyway; "
              "(B) model — bet the AWAY ML as recommended. "
              "If the contrarian ROI is also positive, both sides may have value. "
              "If the model beats the contrarian by a wide margin, the signal is correctly "
              "identifying the better side.",
              color=GRAY, size_pt=9, space_after=6)

    _add_verdict_box(doc, st2h, "S2-A — Contrarian (home ML)")
    _add_verdict_box(doc, st2a, "S2-B — Model pick (away ML)")

    # Comparison table
    _add_heading(doc, "Three-Way Comparison (Same Games)", level=2)
    if st2a["n"] > 0 and st2h["n"] > 0:
        diff = st2h["roi"] - st2a["roi"]
        gap_note = (f"Contrarian leads model by {abs(diff):.2f}pp ROI"
                    if diff >= 0 else f"Model leads contrarian by {abs(diff):.2f}pp ROI")
    else:
        gap_note = "N/A"

    _add_stat_table(doc,
        headers=["Track", "N", "W", "L", "Hit %", "ROI", "Net P&L"],
        rows=[
            ["A — Contrarian (home ML)",
             st2h["n"], st2h["wins"], st2h["losses"],
             f"{st2h['hit']:.1f}%", f"{st2h['roi']:+.2f}%",
             f"+${st2h['pnl']:.2f}" if st2h["pnl"] >= 0 else f"-${abs(st2h['pnl']):.2f}"],
            ["B — Model pick (away ML)",
             st2a["n"], st2a["wins"], st2a["losses"],
             f"{st2a['hit']:.1f}%", f"{st2a['roi']:+.2f}%",
             f"+${st2a['pnl']:.2f}" if st2a["pnl"] >= 0 else f"-${abs(st2a['pnl']):.2f}"],
            ["ROI gap (A minus B)", "", "", "", "", gap_note, ""],
        ],
        col_widths_inches=[2.4, 0.5, 0.5, 0.5, 0.7, 0.9, 1.1])

    # Season breakdown (model side)
    sb2 = _season_breakdown(s2, pnl_key="pnl_away", result_key="result_away", stake=stake)
    if sb2:
        _add_heading(doc, "Model Pick (Away ML) — Season-by-Season Stability", level=2)
        rows = []
        for yr, st in sb2.items():
            pnl_str = f"+${st['pnl']:.2f}" if st["pnl"] >= 0 else f"-${abs(st['pnl']):.2f}"
            rows.append([yr, st["n"], st["wins"], st["losses"],
                         f"{st['hit']:.1f}%", f"{st['roi']:+.2f}%", pnl_str])
        _add_stat_table(doc,
            headers=["Season", "N", "W", "L", "Hit %", "ROI", "Net P&L"],
            rows=rows,
            col_widths_inches=[1.0, 0.5, 0.5, 0.5, 0.7, 0.8, 1.0])

    # Streak breakdown (contrarian side)
    stk2 = _streak_breakdown(s2, pnl_key="pnl_home", result_key="result_home", stake=stake)
    non_empty2 = {k: v for k, v in stk2.items() if v["n"] > 0}
    if non_empty2:
        _add_heading(doc, "Contrarian (Home ML) — Streak Depth in Avoid Games", level=2)
        rows = []
        for bkt, st in non_empty2.items():
            rows.append([bkt, st["n"], st["wins"], st["losses"],
                         f"{st['hit']:.1f}%", f"{st['roi']:+.2f}%"])
        _add_stat_table(doc,
            headers=["Streak Bucket", "N", "W", "L", "Hit %", "ROI"],
            rows=rows,
            col_widths_inches=[1.2, 0.6, 0.6, 0.6, 0.8, 0.8])

    # ── Section 3 ──────────────────────────────────────────────────────────
    doc.add_page_break()
    _add_heading(doc, "Section 3 — Quiet Slate Home Underdogs (No Signal, No Flag)", level=2)
    _add_body(doc,
              "Trigger: no model signal fires and no Watch/Avoid flag is raised. "
              "Bet: home team ML where the home team is a genuine underdog (home_ml > 0) "
              "AND implied win probability ≤ 40% (approx. +150 or longer). "
              "Hypothesis: home dogs are systematically underpriced when the market has "
              "no home-team narrative to rally behind.",
              color=GRAY, size_pt=9, space_after=6)
    _add_verdict_box(doc, st3, "Section 3 — Quiet home dog bet")

    _add_heading(doc, "Overall Performance", level=2)
    _add_stat_table(doc,
        headers=["Population", "N", "W", "L", "Hit %", "ROI", "Net P&L"],
        rows=[[
            "Quiet home dogs (impl ≤ 40%)",
            st3["n"], st3["wins"], st3["losses"],
            f"{st3['hit']:.1f}%", f"{st3['roi']:+.2f}%",
            f"+${st3['pnl']:.2f}" if st3["pnl"] >= 0 else f"-${abs(st3['pnl']):.2f}",
        ]],
        col_widths_inches=[2.8, 0.5, 0.5, 0.5, 0.7, 0.7, 1.1])

    sb3 = _season_breakdown(s3, stake=stake)
    if sb3:
        _add_heading(doc, "Season-by-Season Stability", level=2)
        rows = []
        for yr, st in sb3.items():
            pnl_str = f"+${st['pnl']:.2f}" if st["pnl"] >= 0 else f"-${abs(st['pnl']):.2f}"
            rows.append([yr, st["n"], st["wins"], st["losses"],
                         f"{st['hit']:.1f}%", f"{st['roi']:+.2f}%", pnl_str])
        _add_stat_table(doc,
            headers=["Season", "N", "W", "L", "Hit %", "ROI", "Net P&L"],
            rows=rows,
            col_widths_inches=[1.0, 0.5, 0.5, 0.5, 0.7, 0.8, 1.0])

    op3 = _odds_breakdown(s3, stake=stake)
    if op3:
        _add_heading(doc, "Odds Profile — Home ML Price Distribution in Quiet Dog Population", level=2)
        rows = []
        for bkt, st in op3.items():
            pnl_str = f"+${st['pnl']:.2f}" if st["pnl"] >= 0 else f"-${abs(st['pnl']):.2f}"
            rows.append([bkt, st["n"], f"{st['hit']:.1f}%", f"{st['roi']:+.2f}%", pnl_str])
        _add_stat_table(doc,
            headers=["Home ML Bucket", "N", "Hit %", "ROI", "Net P&L"],
            rows=rows,
            col_widths_inches=[2.8, 0.6, 0.8, 0.8, 1.2])

    # ── Notes page ─────────────────────────────────────────────────────────
    doc.add_page_break()
    _add_heading(doc, "Methodology Notes", level=2)

    notes = [
        "All odds are closing-line prices from mlb_stats.db. No opening odds (CLV) comparison is included here — use backtest_top_pick.py for CLV analysis on primary signal fires.",
        f"ROI is after-vig (no vig removal applied). The +{PASS_ROI}% pass threshold is set to clear the effective vig on a flat -110 betting structure.",
        f"Watch threshold: overall team streak ≤ L4 (primary signals fire at L5 or worse). Streak counts all games (home + away) — load_streaks() has no home-game-only streak. Section 1 captures L4 overall-streak games where no signal fires.",
        "Avoid threshold: any game where the top-priority model pick is the AWAY ML. All primary signals (S1, S1+H2, MV-F, MV-B, H3b) can trigger this.",
        "Quiet home dog: implied win probability ≤ 40% (approximately home_ml ≥ +150). Only games with no signal and no Watch/Avoid flag.",
        f"Pass threshold: ROI ≥ +{PASS_ROI}% on N ≥ {PASS_N} graded bets per period. Both criteria must be met. Confirm with OOS validation before adding any signal to live model.",
        "Season-by-season stability: no single season should account for more than 60% of total P&L. Consistent positive ROI across multiple seasons is required for inclusion.",
        f"Generated: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')} ET",
    ]

    for note in notes:
        p   = doc.add_paragraph(style="List Bullet" if False else "Normal")
        run = p.add_run(f"·  {note}")
        run.font.size = Pt(9)
        run.font.color.rgb = GRAY
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(4)
        p.paragraph_format.left_indent  = Inches(0.15)

    # Save
    try:
        doc.save(out_path)
        return True
    except Exception as e:
        print(f"  ✗  Could not save Word document: {e}")
        return False


# ═══════════════════════════════════════════════════════════════════════════
# CLI
# ═══════════════════════════════════════════════════════════════════════════

def parse_args():
    p = argparse.ArgumentParser(
        description="MLB Betting Model — Fade the Flag Backtester",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=textwrap.dedent("""\
            EXAMPLES
            --------
            Full OOS period (OddsWarehouse era):
              python backtest_fade_flag.py --from 2022-04-01 --to 2025-09-30

            Full historical (SBRO + OW combined):
              python backtest_fade_flag.py --from 2018-04-01 --to 2025-09-30

            Single season:
              python backtest_fade_flag.py --from 2025-04-01 --to 2025-09-30

            Single month:
              python backtest_fade_flag.py --month 2025-09

            Save Word report (recommended):
              python backtest_fade_flag.py --from 2018-04-01 --to 2025-09-30 --docx

            Save text report:
              python backtest_fade_flag.py --from 2018-04-01 --to 2025-09-30 --output reports/fade_flag.txt

            List available months:
              python backtest_fade_flag.py --list-months
        """),
    )
    p.add_argument("--month", default=None,
                   help="Month as YYYY-MM. Sets --from to first and --to to last day.")
    p.add_argument("--from",  dest="date_from", default=None,
                   help="Start date YYYY-MM-DD (inclusive).")
    p.add_argument("--to",    dest="date_to",   default=None,
                   help="End date YYYY-MM-DD (inclusive).")
    p.add_argument("--stake", type=float, default=STAKE,
                   help=f"Flat bet stake per game in dollars (default: {STAKE}).")
    p.add_argument("--db",    default=DEFAULT_DB,
                   help="Path to mlb_stats.db (default: same folder as script).")
    p.add_argument("--output", default=None,
                   help="Save text report to this path.")
    p.add_argument("--docx",  action="store_true",
                   help="Generate a formatted Word (.docx) report in reports/ folder.")
    p.add_argument("--list-months", action="store_true",
                   help="List available months in DB and exit.")
    return p.parse_args()


def main():
    args = parse_args()

    print(f"\n{'═'*74}")
    print(f"  MLB Scout · Fade the Flag Backtester")
    print(f"  Loading signal engine from generate_daily_brief.py ...")

    mod = _load_brief_module()
    con = get_connection(args.db)

    if args.list_months:
        list_available_months(con)
        con.close()
        return

    # Resolve date range
    if args.month:
        try:
            y, m   = int(args.month[:4]), int(args.month[5:7])
            d_from = f"{y:04d}-{m:02d}-01"
            d_to   = (datetime.date(y + (m // 12), (m % 12) + 1, 1)
                      - datetime.timedelta(days=1)).isoformat()
        except (ValueError, IndexError):
            print(f"✗  Invalid --month '{args.month}'. Use YYYY-MM.")
            sys.exit(1)
    elif args.date_from and args.date_to:
        d_from = args.date_from
        d_to   = args.date_to
    else:
        print("✗  Specify either --month YYYY-MM  or both --from and --to.")
        sys.exit(1)

    try:
        datetime.date.fromisoformat(d_from)
        datetime.date.fromisoformat(d_to)
    except ValueError as e:
        print(f"✗  Invalid date: {e}")
        sys.exit(1)

    print(f"  Range: {d_from}  →  {d_to}   Stake: ${args.stake:.2f}/bet")
    print(f"{'═'*74}")
    print(f"  Running backtest ...\n")

    results = run_backtest(con, mod, d_from, d_to, args.stake)
    report  = format_report(results)
    print(report)

    # Save text report
    if args.output:
        out_path = Path(args.output)
        out_path.parent.mkdir(parents=True, exist_ok=True)
        with open(out_path, "w", encoding="utf-8") as fh:
            fh.write(report)
        print(f"  ✓ Text report saved: {args.output}\n")

    # Save Word report
    if args.docx:
        reports_dir = _SCRIPT_DIR / "reports"
        reports_dir.mkdir(exist_ok=True)
        ts       = datetime.datetime.now().strftime("%Y%m%d_%H%M")
        docx_path = reports_dir / f"backtest_fade_flag_{ts}.docx"
        print(f"  Building Word report ...")
        ok = build_docx_report(results, str(docx_path), args.stake)
        if ok:
            print(f"  ✓ Word report saved: {docx_path}\n")

    con.close()


if __name__ == "__main__":
    main()
