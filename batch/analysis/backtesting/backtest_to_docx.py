#!/usr/bin/env python3
"""
backtest_to_docx.py
Converts a backtest_report.md file to a formatted Word document.
Saves to reports/ subfolder with a datetime stamp in the filename.

Usage:
    python backtest_to_docx.py
    python backtest_to_docx.py --input backtest_report.md
    python backtest_to_docx.py --input h4_rolling_era_report.md

Output:
    reports\backtest_report_2026-03-15_21-30.docx
"""

import argparse
import re
import sys
from datetime import datetime
from pathlib import Path

# ── Dependency check ──────────────────────────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

# ── Paths ─────────────────────────────────────────────────────────────────────
PROJ_DIR    = Path(r"C:\Users\sevan\OneDrive\Documents\Python\mlb_stats")
REPORTS_DIR = PROJ_DIR / "reports"
DEFAULT_IN  = PROJ_DIR / "backtest_report.md"

# ── Colours ───────────────────────────────────────────────────────────────────
NAVY  = RGBColor(0x1B, 0x3A, 0x5C)
TEAL  = RGBColor(0x0D, 0x7E, 0xA8)
GRAY  = RGBColor(0x64, 0x74, 0x8B)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK  = RGBColor(0x1E, 0x29, 0x3B)
CODE_BG = 0x1E293B   # hex int for table shading


# ── Style helpers ─────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Set table cell background colour via XML (python-docx doesn't expose this directly)."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color.upper())
    tcPr.append(shd)


def heading_paragraph(doc, text: str, level: int):
    """Add a styled heading that avoids built-in heading styles."""
    p    = doc.add_paragraph()
    run  = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size  = Pt(18)
        run.font.color.rgb = NAVY
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after  = Pt(8)
    elif level == 2:
        run.font.size  = Pt(14)
        run.font.color.rgb = TEAL
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(6)
    else:
        run.font.size  = Pt(12)
        run.font.color.rgb = NAVY
        run.font.italic = True
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(4)
    run.font.name = "Arial"
    return p


def body_paragraph(doc, text: str, italic=False, indent=False):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name  = "Arial"
    run.font.size  = Pt(11)
    run.font.color.rgb = DARK
    run.italic     = italic
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    return p


def blockquote_paragraph(doc, text: str):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name      = "Arial"
    run.font.size      = Pt(10)
    run.font.color.rgb = GRAY
    run.italic         = True
    p.paragraph_format.left_indent  = Inches(0.4)
    p.paragraph_format.space_after  = Pt(6)
    return p


def code_paragraph(doc, text: str):
    """Monospace dark-bg style for code blocks."""
    p   = doc.add_paragraph()
    run = p.add_run(text if text.strip() else " ")
    run.font.name      = "Courier New"
    run.font.size      = Pt(9)
    run.font.color.rgb = RGBColor(0xE2, 0xE8, 0xF0)
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    # Set paragraph background
    pPr  = p._p.get_or_add_pPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "1E293B")
    pPr.append(shd)
    return p


def add_markdown_table(doc, lines: list):
    """Parse and render a markdown table."""
    # lines[0] = header row, lines[1] = separator, lines[2:] = data rows
    def parse_row(line):
        return [c.strip() for c in line.strip().strip("|").split("|")]

    headers   = parse_row(lines[0])
    data_rows = [parse_row(l) for l in lines[2:] if l.strip() and "---" not in l]

    if not headers or not data_rows:
        return

    col_count = len(headers)
    tbl = doc.add_table(rows=1 + len(data_rows), cols=col_count)
    tbl.style = "Table Grid"

    # Header row
    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold           = True
        run.font.name      = "Arial"
        run.font.size      = Pt(9)
        run.font.color.rgb = WHITE
        set_cell_bg(cell, "1B3A5C")

    # Data rows
    for ri, row_data in enumerate(data_rows):
        row = tbl.rows[ri + 1]
        bg  = "F8FBFD" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row_data[:col_count]):
            cell = row.cells[ci]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            run.font.name      = "Arial"
            run.font.size      = Pt(9)
            run.font.color.rgb = DARK
            set_cell_bg(cell, bg)

    doc.add_paragraph()   # spacing after table


# ── Markdown parser ───────────────────────────────────────────────────────────

def md_to_docx(md_text: str, doc: Document):
    """Parse markdown and write into a python-docx Document."""
    lines = md_text.splitlines()
    i     = 0

    while i < len(lines):
        line = lines[i]

        # ── Headings ──────────────────────────────────────────────────────────
        if line.startswith("### "):
            heading_paragraph(doc, line[4:].strip(), 3)
            i += 1; continue
        if line.startswith("## "):
            heading_paragraph(doc, line[3:].strip(), 2)
            i += 1; continue
        if line.startswith("# "):
            heading_paragraph(doc, line[2:].strip(), 1)
            i += 1; continue

        # ── Blockquote ────────────────────────────────────────────────────────
        if line.startswith("> "):
            blockquote_paragraph(doc, line[2:].strip())
            i += 1; continue

        # ── Horizontal rule ───────────────────────────────────────────────────
        if line.strip() in ("---", "***", "___"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(6)
            i += 1; continue

        # ── Code block ────────────────────────────────────────────────────────
        if line.strip().startswith("```"):
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_paragraph(doc, lines[i])
                i += 1
            doc.add_paragraph()  # space after code block
            i += 1; continue

        # ── Markdown table ────────────────────────────────────────────────────
        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            if len(table_lines) >= 2:
                add_markdown_table(doc, table_lines)
            continue

        # ── Bullet ────────────────────────────────────────────────────────────
        if re.match(r"^[-*] ", line):
            text = line[2:].strip()
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)   # strip bold markers
            text = re.sub(r"`(.+?)`",        r"\1", text)   # strip code markers
            body_paragraph(doc, f"• {text}", indent=True)
            i += 1; continue

        # ── Italic / emphasis line ─────────────────────────────────────────────
        if line.startswith("*") and line.endswith("*") and len(line) > 2:
            body_paragraph(doc, line.strip("*"), italic=True)
            i += 1; continue

        # ── Blank line ────────────────────────────────────────────────────────
        if not line.strip():
            i += 1; continue

        # ── Normal paragraph ──────────────────────────────────────────────────
        # Strip inline markdown: **bold** and `code`
        clean = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
        clean = re.sub(r"`(.+?)`",        r"\1", clean)
        body_paragraph(doc, clean)
        i += 1


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    p = argparse.ArgumentParser(description="Convert backtest markdown report to Word doc")
    p.add_argument("--input", "-i", default=str(DEFAULT_IN),
                   help=f"Input .md file (default: {DEFAULT_IN})")
    p.add_argument("--out-dir", default=str(REPORTS_DIR),
                   help=f"Output directory (default: {REPORTS_DIR})")
    args = p.parse_args()

    in_path  = Path(args.input)
    out_dir  = Path(args.out_dir)

    if not in_path.exists():
        print(f"ERROR: Input file not found: {in_path}")
        sys.exit(1)

    # Create reports\ directory if it doesn't exist
    out_dir.mkdir(parents=True, exist_ok=True)

    # Build output filename with datetime stamp
    stem      = in_path.stem                                    # e.g. backtest_report
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M")       # e.g. 2026-03-15_21-30
    out_name  = f"{stem}_{timestamp}.docx"
    out_path  = out_dir / out_name

    # Read markdown
    md_text = in_path.read_text(encoding="utf-8")

    # Build document
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    # Cover line
    cover = doc.add_paragraph()
    r1 = cover.add_run("MLB Backtesting Report  ")
    r1.font.name  = "Arial"
    r1.font.size  = Pt(10)
    r1.font.color.rgb = GRAY
    r2 = cover.add_run(f"Generated {datetime.now().strftime('%B %d, %Y  %H:%M')}")
    r2.font.name  = "Arial"
    r2.font.size  = Pt(10)
    r2.font.color.rgb = TEAL
    r2.bold = True
    cover.paragraph_format.space_after = Pt(12)

    # Parse and render the markdown
    md_to_docx(md_text, doc)

    # Save
    doc.save(out_path)
    print(f"Saved: {out_path}")
    print(f"  Source: {in_path.name}  ({len(md_text):,} chars)")
    print(f"  Output: {out_name}")


if __name__ == "__main__":
    main()
